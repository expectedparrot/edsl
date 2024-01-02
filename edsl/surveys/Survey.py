from __future__ import annotations
import re
import uuid
from docx import Document
from rich import print
from typing import Any, Generator, Optional, Union
from edsl.exceptions import SurveyCreationError, SurveyHasNoRulesError
from edsl.questions.Question import Question
from edsl.surveys.base import RulePriority, EndOfSurvey
from edsl.surveys.Rule import Rule, RuleCollection


class Survey:
    """
    A collection of questions that supports skip logic.

    Initalization:
    - `questions`: the questions in the survey (optional)
    - `question_names`: the names of the questions (optional)
    - `name`: the name of the survey (optional)

    Methods:
    -

    Notes:
    - The presumed order of the survey is the order in which questions are added.
    """

    def __init__(
        self,
        questions: list[Question] = None,
        question_names: list[str] = None,
        name: str = None,
    ):
        """Creates a new survey."""
        self.question_uuid_to_index: dict[str, int] = {}
        self.question_name_to_uuid: dict[str, str] = {}
        self.question_name_to_index: dict[str, int] = {}
        self._questions: list[Question] = []
        self.rule_collection = RuleCollection()
        self.name = name
        self.uuid = "survey_" + str(uuid.uuid4())
        question_names = question_names or []
        questions = questions or []
        if question_names:
            if len(questions) != len(set(question_names)):
                raise SurveyCreationError(
                    "Question names must be unique and equal to the number of questions."
                )
        else:
            for i, question in enumerate(questions):
                question_name = question.question_name or f"q{i}"
                if question_name in question_names:
                    raise SurveyCreationError(
                        "Failed to automatically create names for questions without names."
                    )
                question_names.append(question_name)

        # adds all questions to the survey
        for question, question_name in zip(questions, question_names):
            self.add_question(question, question_name)

    def get_question(self, question_name) -> Question:
        """Returns the question object given the question name"""
        if question_name not in self.question_name_to_index:
            raise KeyError(f"Question name {question_name} not found in survey.")
        index = self.question_name_to_index[question_name]
        return self._questions[index]

    @property
    def question_names(self) -> list[str]:
        """Returns a list of question names in the survey"""
        return list(self.question_name_to_uuid.keys())

    def add_question(
        self, question: Question, question_name: Optional[str] = None
    ) -> Survey:
        """
        Adds a question to survey.
        - The question is appended at the end of the self.questions list
        - A default rule is created that the next index is the next question.
        """
        if question_name is None:
            if hasattr(question, "question_name"):
                question_name = question.question_name
            else:
                question_name = "q" + str(len(self._questions))

        index = len(self._questions)

        # question name needs to be unique
        assert question_name not in self.question_name_to_uuid
        self._questions.append(question)
        self.question_name_to_uuid[question_name] = question.uuid
        self.question_name_to_index[question_name] = index
        self.question_uuid_to_index[question.uuid] = index

        # using index + 1 presumes there is a next question
        self.rule_collection.add_rule(
            Rule(
                current_q=index,
                expression="True",
                next_q=index + 1,
                question_name_to_index=self.question_name_to_index,
                priority=RulePriority.DEFAULT.value,
            )
        )
        return self

    def add_stop_rule(self, question: Question, expression: str) -> Survey:
        """Adds a rule that stops the survey."""
        self.add_rule(question, expression, EndOfSurvey())
        return self

    def add_rule(
        self, question: Question, expression: str, next_question: Question
    ) -> Survey:
        """
        Adds a rule to a Question of the Survey with the appropriate priority.
        - The last rule added for the question will have the highest priority.
        - If there are no rules, the rule added gets priority -1.
        """

        # we let users refer to questions by name or pass the questions themselves.
        # if they pass the questions themselves, we use the uuid of the question to get the index.
        if isinstance(question, str):
            question_index = self.question_name_to_index[question]
        else:
            question_index = self.question_uuid_to_index[question.uuid]

        if isinstance(next_question, str):
            next_question_index = self.question_name_to_index[next_question]
        else:
            # we need to account for possibility that 'next_q' is actually an EndOfSurvey object
            if isinstance(next_question, EndOfSurvey):
                next_question_index = EndOfSurvey()
            else:
                next_question_index = self.question_uuid_to_index[next_question.uuid]

        # finds the priorities of existing rules that apply to the question
        priorities = [
            rule.priority for rule in self.rule_collection.which_rules(question_index)
        ]

        if len(priorities) == 0:
            priority = RulePriority.DEFAULT.value
        else:
            priority = max(priorities) + 1  # newer rules take priority over older rules

        self.rule_collection.add_rule(
            Rule(
                question_index,
                expression,
                next_question_index,
                self.question_name_to_index,
                priority=priority,
            )
        )

        return self

    ###################
    # FORWARD METHODS
    ###################
    def by(self, *args: Union[Agent, Scenario, LanguageModel]) -> Jobs:
        """Adds Agents, Scenarios, and LanguageModels to a survey and returns a runnable Jobs object"""
        from edsl.jobs.Jobs import Jobs

        job = Jobs(survey=self)
        return job.by(*args)

    def run(self, *args, **kwargs) -> Jobs:
        "Turns the survey into a Job and runs it"
        from edsl.jobs.Jobs import Jobs

        return Jobs(survey=self).run(*args, **kwargs)

    def next_question(
        self, current_question_name: str = None, answers=None
    ) -> Union[Question, EndOfSurvey]:
        """
        Returns the next question in a survey.
        - If called with no arguments, it returns the first question in the survey.
        - If no answers are provided for a question with a rule, the next question is returned. If answers are provided, the next question is determined by the rules and the answers.
        - If the next question is the last question in the survey, an EndOfSurvey object is returned.
        """
        if current_question_name is None:
            return self._questions[0]

        answers = answers or {}
        question_index = self.question_name_to_index[current_question_name]
        next_question_object = self.rule_collection.next_question(
            question_index, answers
        )

        if next_question_object.num_rules_found == 0:
            raise SurveyHasNoRulesError

        if isinstance(next_question_object.next_q, EndOfSurvey):
            return EndOfSurvey()
        else:
            if next_question_object.next_q >= len(self._questions):
                return EndOfSurvey()
            else:
                return self._questions[next_question_object.next_q]

    def gen_path_through_survey(self) -> Generator[Question, dict, None]:
        """
        Generates a coroutine that can be used to conduct an Interview.
        - The coroutine is a generator that yields a question and receives answers.
        - The coroutine starts with the first question in the survey.
        - The coroutine ends when an EndOfSurvey object is returned.
        """
        question = self.next_question()
        question_name = self.question_names[0]
        question.question_name = question_name
        ## We send to the agent the first question and it's name --- say q1 and 'q1'
        ## it sends back a response, which is answers = {'q1': "yes"}
        answers = yield question
        self.answers = answers
        while True:
            question = self.next_question(question.question_name, answers)
            if isinstance(question, EndOfSurvey):
                break
            index = self.question_uuid_to_index[question.uuid]
            question.question_name = self.question_names[index]
            answers = yield question
            self.answers = answers

    @property
    def scenario_attributes(self) -> list[str]:
        """Returns a list of attributes that admissible Scenarios should have"""
        temp = []
        for question in self._questions:
            question_text = question.question_text
            # extract the contents of all {{ }} in the question text using regex
            matches = re.findall(r"\{\{(.+?)\}\}", question_text)
            # remove whitespace
            matches = [match.strip() for match in matches]
            # add them to the temp list
            temp.extend(matches)
        return temp

    ###################
    # DUNDER METHODS
    ###################
    def __len__(self) -> int:
        """Returns the number of questions in the survey"""
        return len(self._questions)

    def __getitem__(self, index) -> Question:
        """Returns the question object given the question index"""
        return self._questions[index]

    ###################
    # SERIALIZATION METHODS
    ###################
    def to_dict(self) -> dict[str, Any]:
        """Serializes the Survey object to a dictionary."""
        return {
            "questions": [q.to_dict() for q in self._questions],
            "name": self.name,
            "rule_collection": self.rule_collection.to_dict(),
        }

    @classmethod
    def from_dict(cls, data: dict) -> Survey:
        """Deserializes the dictionary back to a Survey object."""
        questions = [Question.from_dict(q_dict) for q_dict in data["questions"]]
        survey = cls(questions=questions, name=data["name"])
        survey.rule_collection = RuleCollection.from_dict(data["rule_collection"])
        return survey

    ###################
    # DISPLAY METHODS
    ###################
    def __repr__(self) -> str:
        """Returns a string representation of the survey"""
        questions_string = ", ".join([repr(q) for q in self._questions])
        question_names_string = ", ".join([repr(name) for name in self.question_names])
        return f"Survey(questions=[{questions_string}], question_names=[{question_names_string}], name={repr(self.name)})"

    def _repr_html_(self) -> str:
        return self.html()

    def show_rules(self) -> None:
        "Prints out the rules in the survey"
        self.rule_collection.show_rules()

    def print(self) -> None:
        "Prints out the survey"
        self.show_questions()

    def show_questions(self):
        "Prints out the questions in the survey"
        for name, question in zip(self.question_names, self._questions):
            print(f"Question:{name},{question}")

    def html(self) -> str:
        "Generates the html for the survey"
        html_text = []
        for question in self._questions:
            html_text.append(
                f"<p><b>{question.question_name}</b> ({question.question_type}): {question.question_text}</p>"
            )
            html_text.append("<ul>")
            for option in getattr(question, "question_options", []):
                html_text.append(f"<li>{option}</li>")
            html_text.append("</ul>")
        return "\n".join(html_text)

    def docx(self) -> Document:
        "Generates a docx document for the survey"
        doc = Document()
        doc.add_heading("EDSL Auto-Generated Survey")
        doc.add_paragraph(f"\n")
        for index, question in enumerate(self._questions):
            h = doc.add_paragraph()  # Add question as a paragraph
            h.add_run(f"Question {index + 1} ({question.question_name})").bold = True
            h.add_run(f"; {question.question_type}").italic = True
            p = doc.add_paragraph()
            p.add_run(question.question_text)
            if question.question_type == "linear_scale":
                for key, value in getattr(question, "option_labels", {}).items():
                    doc.add_paragraph(str(key) + ": " + str(value), style="ListBullet")
            else:
                if hasattr(question, "question_options"):
                    for option in getattr(question, "question_options", []):
                        doc.add_paragraph(str(option), style="ListBullet")
        return doc

    def code(self) -> list[str]:
        ## TODO: Refactor to only use the questions actually in the survey
        "Creates the Python code representation of a survey"
        header_lines = ["from edsl.surveys.Survey import Survey"]
        header_lines.append(
            "from edsl.questions.QuestionMultipleChoice import QuestionMultipleChoice"
        )
        header_lines.append(
            "from edsl.questions.QuestionFreeText import QuestionFreeText"
        )
        header_lines.append(
            "from edsl.questions.derived.QuestionLinearScale import QuestionLinearScale"
        )
        header_lines.append(
            "from edsl.questions.QuestionNumerical import QuestionNumerical"
        )
        header_lines.append(
            "from edsl.questions.QuestionCheckBox import QuestionCheckBox"
        )
        header_lines.append(
            "from edsl.questions.derived.QuestionYesNo import QuestionYesNo"
        )
        lines = ["\n".join(header_lines)]
        for question in self._questions:
            lines.append(f"{question.question_name} = " + repr(question))
        lines.append(f"survey = Survey(questions = [{', '.join(self.question_names)}])")
        return lines

    def codebook(self) -> dict[str, str]:
        "Creates a codebook for the survey, mapping question names to question text"
        codebook = {}
        for question in self._questions:
            codebook[question.question_name] = question.question_text
        return codebook


def main():
    def example_survey():
        from edsl.questions.QuestionMultipleChoice import QuestionMultipleChoice
        from edsl.surveys.Survey import Survey

        q0 = QuestionMultipleChoice(
            question_text="Do you like school?",
            question_options=["yes", "no"],
            question_name="q0",
        )
        q1 = QuestionMultipleChoice(
            question_text="Why not?",
            question_options=["killer bees in cafeteria", "other"],
            question_name="q1",
        )
        q2 = QuestionMultipleChoice(
            question_text="Why?",
            question_options=["**lack*** of killer bees in cafeteria", "other"],
            question_name="q2",
        )
        s = Survey(questions=[q0, q1, q2])
        s = s.add_rule(q0, "q0 == 'yes'", q2)
        return s

    s = example_survey()
    survey_dict = s.to_dict()
    s2 = Survey.from_dict(survey_dict)
    results = s2.run()
    print(results)
