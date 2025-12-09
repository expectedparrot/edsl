"""A class for exporting surveys to different formats."""

from typing import Optional, TYPE_CHECKING

import subprocess
import platform
import os
import tempfile

if TYPE_CHECKING:
    from ..scenarios import ScenarioList, FileStore


def open_docx(file_path):
    """
    Open a docx file using the default application in a cross-platform manner.

    :param file_path: str, path to the docx file
    """
    file_path = os.path.abspath(file_path)

    if platform.system() == "Darwin":  # macOS
        subprocess.call(("open", file_path))
    elif platform.system() == "Windows":  # Windows
        os.startfile(file_path)
    else:  # linux variants
        subprocess.call(("xdg-open", file_path))


class SurveyExport:
    """A class for exporting surveys to different formats."""

    def __init__(self, survey):
        """Initialize with a Survey object."""
        self.survey = survey

    def css(self):
        from .survey_css import SurveyCSS

        return SurveyCSS.default_style().generate_css()

    # def get_description(self) -> str:
    #     """Return the description of the survey."""
    #     from edsl import QuestionFreeText

    #     question_texts = "\n".join([q.question_text for q in self.survey._questions])
    #     q = QuestionFreeText(
    #         question_name="description",
    #         question_text=f"""A survey was conducted with the following questions:
    #                          {question_texts}
    #                          Please write a description of the survey.
    #                          """,
    #     )
    #     return q.run().select("description").first()

    def docx(
        self,
        filename: Optional[str] = None,
    ) -> "FileStore":
        """Generate a docx document for the survey."""
        from docx import Document
        from edsl import FileStore  # Added import for FileStore

        if not filename:  # handles None and ""
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            filename = tmp.name
            tmp.close()

        doc = Document()
        doc.add_heading("EDSL Survey")
        doc.add_paragraph("\n")
        for index, question in enumerate(self.survey._questions):
            h = doc.add_paragraph()  # Add question as a paragraph
            h.add_run(f"Question {index + 1} ({question.question_name})").bold = True
            h.add_run(f"; {question.question_type}").italic = True
            p = doc.add_paragraph()
            p.add_run(question.question_text)
            if question.question_type == "linear_scale":
                for key, value in getattr(question, "option_labels", {}).items():
                    doc.add_paragraph(str(key) + ": " + str(value), style="ListBullet")
            else:
                if hasattr(question, "question_options"):
                    for option in getattr(question, "question_options", []):
                        doc.add_paragraph(str(option), style="ListBullet")

        # Determine filename: use provided filename or create a temporary one
        if filename is None:
            current_directory = os.getcwd()
            # Create a temporary file to hold the .docx content
            temp_file = tempfile.NamedTemporaryFile(
                "wb", delete=False, suffix=".docx", dir=current_directory
            )
            filename = temp_file.name
            temp_file.close()

        # Save the document to the determined filename
        doc.save(filename)

        # Return a FileStore object for the generated file
        return FileStore(filename)

    def show(self):
        self.to_scenario_list(questions_only=False, rename=True).print(format="rich")

    def to_scenario_list(
        self,
        questions_only: bool = True,
        rename=False,
        remove_jinja2_syntax: bool = False,
    ) -> "ScenarioList":
        import re
        from ..scenarios import ScenarioList, Scenario

        if questions_only:
            to_iterate_over = self.survey._questions
        else:
            to_iterate_over = self.survey.recombined_questions_and_instructions()

        if rename:
            renaming_dict = {
                "name": "identifier",
                "question_name": "identifier",
                "question_text": "text",
            }
        else:
            renaming_dict = {}

        all_keys = []
        scenarios = ScenarioList()
        for item in to_iterate_over:
            d = item.to_dict()
            if item.__class__.__name__ == "Instruction":
                d["question_type"] = "NA / instruction"

            # Remove Jinja2 syntax from question_text if requested
            if remove_jinja2_syntax and "question_text" in d:
                # Remove {{ }} brackets and their contents, preserving spacing
                d["question_text"] = re.sub(
                    r"\s*\{\{.*?\}\}\s*", " ", d["question_text"]
                )
                # Clean up extra whitespace that may result from removal
                d["question_text"] = " ".join(d["question_text"].split()).strip()

            for key in renaming_dict:
                if key in d:
                    d[renaming_dict[key]] = d.pop(key)
            # Preserve order by using list with manual deduplication
            for key in d.keys():
                if key not in all_keys:
                    all_keys.append(key)
            scenarios.append(Scenario(d))

        for scenario in scenarios:
            for key in all_keys:
                if key not in scenario:
                    scenario[key] = None

        return scenarios

    def code(self, filename: str = None, survey_var_name: str = "survey") -> list[str]:
        """Create the Python code representation of a survey.

        :param filename: The name of the file to save the code to.
        :param survey_var_name: The name of the survey variable.

        >>> from edsl.surveys import Survey
        >>> survey = Survey.example()
        >>> print(survey.code())
        from edsl.surveys.Survey import Survey
        ...
        ...
        survey = Survey(questions=[q0, q1, q2])
        ...
        """
        import black

        header_lines = ["from edsl.surveys.Survey import Survey"]
        header_lines.append("from edsl import Question")
        lines = ["\n".join(header_lines)]
        for question in self.survey._questions:
            question.question_text = question["question_text"].replace("\n", " ")
            # remove dublicate spaces
            question.question_text = " ".join(question.question_text.split())
            lines.append(f"{question.question_name} = " + repr(question))
        lines.append(
            f"{survey_var_name} = Survey(questions = [{', '.join(self.survey.question_names)}])"
        )
        # return lines
        code_string = "\n".join(lines)
        formatted_code = black.format_str(code_string, mode=black.FileMode())

        if filename:
            print("The code has been saved to", filename)
            print("The survey itself is saved to 'survey' object")
            with open(filename, "w") as file:
                file.write(formatted_code)
            return

        return formatted_code

    def html(
        self,
        scenario: Optional[dict] = None,
        filename: Optional[str] = None,
        return_link=False,
        css: Optional[str] = None,
        cta: Optional[str] = "Open HTML file",
        include_question_name=False,
    ):
        import os
        from edsl import FileStore  # Added import for FileStore

        if scenario is None:
            scenario = {}

        if css is None:
            css = self.css()

        if filename is None:
            current_directory = os.getcwd()
            filename = tempfile.NamedTemporaryFile(
                "w", delete=False, suffix=".html", dir=current_directory
            ).name

        html_header = f"""<html>
        <head><title></title>
        <style>
        { css }
        </style>
        </head>
        <body>
        <div class="survey_container">
        """

        html_footer = """
        </div>
        </body>
        </html>"""

        output = html_header

        with open(filename, "w") as f:
            f.write(html_header)
            for question in self.survey._questions:
                f.write(
                    question.html(
                        scenario=scenario, include_question_name=include_question_name
                    )
                )
                output += question.html(
                    scenario=scenario, include_question_name=include_question_name
                )
            f.write(html_footer)
            output += html_footer

        return FileStore(filename)

    def latex(
        self,
        filename: Optional[str] = None,
        include_question_name: bool = False,
        standalone: bool = True,
    ) -> "FileStore":
        """Generate a LaTeX (.tex) document for the survey.

        Parameters
        ----------
        filename : Optional[str]
            The filename to write to. If not provided, a temporary file is created
            in the current working directory with a ``.tex`` suffix.
        include_question_name : bool, default False
            If True, includes the internal ``question_name`` of each question in
            the rendered LaTeX.
        standalone : bool, default True
            If True, a full LaTeX document is produced with ``\\documentclass`` and
            ``\\begin{document}`` / ``\\end{document}``. If False, only the snippet
            corresponding to the survey content is written (suitable for inclusion
            in a larger document).

        Returns
        -------
        FileStore
            A ``FileStore`` object pointing to the generated ``.tex`` file so it
            can easily be downloaded, viewed, or further processed.
        """
        # Local import to avoid heavy dependency at import-time.
        from edsl import FileStore

        # Determine filename
        if filename is None:
            current_directory = os.getcwd()
            temp_file = tempfile.NamedTemporaryFile(
                "w", delete=False, suffix=".tex", dir=current_directory
            )
            filename = temp_file.name
            temp_file.close()

        # Basic LaTeX document structure (optional)
        if standalone:
            header = r"""\documentclass{article}
\usepackage[utf8]{inputenc}
\usepackage{enumitem}
\begin{document}
\section*{EDSL Survey}
"""
            footer = "\n\\end{document}\n"
        else:
            header = ""
            footer = ""

        def _escape_latex(text: str) -> str:
            """Escape characters that have special meaning in LaTeX."""
            replacements = {
                "&": r"\&",
                "%": r"\%",
                "$": r"\$",
                "#": r"\#",
                "_": r"\\_",
                "{": r"\{",
                "}": r"\}",
                "~": r"\textasciitilde{}",
                "^": r"\textasciicircum{}",
                "\\": r"\textbackslash{}",
            }
            for k, v in replacements.items():
                text = text.replace(k, v)
            return text

        with open(filename, "w", encoding="utf-8") as f:
            f.write(header)
            for idx, question in enumerate(self.survey._questions):
                # Heading for each question
                heading_parts = [f"Question {idx + 1}"]
                if include_question_name:
                    heading_parts.append(
                        f"(\\texttt{{{_escape_latex(question.question_name)}}})"
                    )
                heading_parts.append(
                    f"-- \\textit{{{_escape_latex(question.question_type)}}}"
                )
                heading = " ".join(heading_parts)
                f.write(f"\\subsection*{{{heading}}}\n")

                # Question text
                f.write(f"{_escape_latex(question.question_text)}\\\n\n")

                # Handle options or labels depending on question type
                if question.question_type == "linear_scale":
                    option_labels = getattr(question, "option_labels", {}) or {}
                    if option_labels:
                        f.write("\\begin{itemize}\n")
                        for key, value in option_labels.items():
                            f.write(
                                f"  \\item {_escape_latex(str(key))}: {_escape_latex(str(value))}\n"
                            )
                        f.write("\\end{itemize}\n\n")
                else:
                    if hasattr(question, "question_options"):
                        options = getattr(question, "question_options", []) or []
                        if options:
                            f.write("\\begin{enumerate}[label=\\alph*)]\n")
                            for option in options:
                                f.write(f"  \\item {_escape_latex(str(option))}\n")
                            f.write("\\end{enumerate}\n\n")

            f.write(footer)

        # Return a FileStore object for the generated file
        return FileStore(filename)


if __name__ == "__main__":
    import doctest

    doctest.testmod(optionflags=doctest.ELLIPSIS)
