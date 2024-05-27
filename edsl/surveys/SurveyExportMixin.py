"""A mixin class for exporting surveys to different formats."""
from docx import Document
from typing import Union, Optional
import black


class SurveyExportMixin:
    """A mixin class for exporting surveys to different formats."""

    def css(self):
        return """
    /* General styles for the survey container */
    .survey_container {
        width: 80%;
        margin: 0 auto;
        padding: 20px;
        background-color: #f9f9f9;
        border: 1px solid #ddd;
        border-radius: 8px;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }

    /* Styles for individual survey questions */
    .survey_question {
        margin-bottom: 20px;
        padding: 15px;
        background-color: #fff;
        border: 1px solid #ddd;
        border-radius: 8px;
    }

    /* Styles for the question text */
    .question_text {
        font-size: 18px;
        font-weight: bold;
        margin-bottom: 10px;
        color: #333;
    }

    /* Styles for the question options list */
    .question_options {
        list-style-type: none;
        padding: 0;
        margin: 0;
    }

    /* Styles for each option item */
    .question_options li {
        margin-bottom: 10px;
        font-size: 16px;
        color: #555;
    }

    /* Styles for radio buttons and checkboxes */
    .question_options input[type="radio"],
    .question_options input[type="checkbox"] {
        margin-right: 10px;
    }

    /* Styles for text input questions */
    input[type="text"] {
        width: 100%;
        padding: 10px;
        font-size: 16px;
        border: 1px solid #ddd;
        border-radius: 4px;
        box-sizing: border-box;
    }
    """

    def docx(self, filename=None) -> Union["Document", None]:
        """Generate a docx document for the survey."""
        doc = Document()
        doc.add_heading("EDSL Survey")
        doc.add_paragraph(f"\n")
        for index, question in enumerate(self._questions):
            h = doc.add_paragraph()  # Add question as a paragraph
            h.add_run(f"Question {index + 1} ({question.question_name})").bold = True
            h.add_run(f"; {question.question_type}").italic = True
            p = doc.add_paragraph()
            p.add_run(question.question_text)
            if question.question_type == "linear_scale":
                for key, value in getattr(question, "option_labels", {}).items():
                    doc.add_paragraph(str(key) + ": " + str(value), style="ListBullet")
            else:
                if hasattr(question, "question_options"):
                    for option in getattr(question, "question_options", []):
                        doc.add_paragraph(str(option), style="ListBullet")
        if filename:
            doc.save(filename)
            print("The survey has been saved to", filename)
            return
        return doc

    def code(self, filename: str = None, survey_var_name: str = "survey") -> list[str]:
        """Create the Python code representation of a survey.

        :param filename: The name of the file to save the code to.
        :param survey_var_name: The name of the survey variable.
        """
        header_lines = ["from edsl.surveys.Survey import Survey"]
        header_lines.append("from edsl import Question")
        lines = ["\n".join(header_lines)]
        for question in self._questions:
            question.question_text = question["question_text"].replace("\n", " ")
            # remove dublicate spaces
            question.question_text = " ".join(question.question_text.split())
            lines.append(f"{question.question_name} = " + repr(question))
        lines.append(
            f"{survey_var_name} = Survey(questions = [{', '.join(self.question_names)}])"
        )
        # return lines
        code_string = "\n".join(lines)
        formatted_code = black.format_str(code_string, mode=black.FileMode())

        if filename:
            print("The code has been saved to", filename)
            print("The survey itself is saved to 'survey' object")
            with open(filename, "w") as file:
                file.write(formatted_code)
            return

        return formatted_code
    
    def html(self, 
             scenario: Optional[dict] = None, 
             filename:Optional[str] = None, 
             return_link = False,
             css:Optional[str] = None, 
             cta:Optional[str] = "Open HTML file"): 

        from IPython.display import display, HTML
        import tempfile
        import os
        from edsl.utilities.utilities import is_notebook

        if scenario is None:
            scenario = {}

        if css is None:
            css = self.css()

        if filename is None:
            current_directory = os.getcwd()
            filename = tempfile.NamedTemporaryFile("w", delete=False, suffix=".html", dir=current_directory).name

        html_header = f"""<html>
        <head><title></title>
        <style>
        { css }
        </style>
        </head>
        <body>
        <div class="survey_container">
        """

        html_footer = """
        </div>
        </body>
        </html>"""

        with open(filename, 'w') as f:
            with open(filename, 'w') as f:
                f.write(html_header)
                for question in self._questions:
                    f.write(question.html(scenario = scenario))
                f.write(html_footer)

        if is_notebook():
            html_url = f'/files/{filename}'
            html_link = f'<a href="{html_url}" target="_blank">{cta}</a>'
            display(HTML(html_link))
        else:
            print(f"Survey saved to {filename}")

        if return_link:
            return filename

