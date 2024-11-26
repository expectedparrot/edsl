"""A Scenario is a dictionary with a key/value to parameterize a question."""

from __future__ import annotations
import copy
import hashlib
import os
import json
from collections import UserDict
from typing import Union, List, Optional, Generator
from uuid import uuid4

from edsl.Base import Base
from edsl.scenarios.ScenarioHtmlMixin import ScenarioHtmlMixin
from edsl.utilities.decorators import add_edsl_version, remove_edsl_version
from edsl.exceptions.scenarios import ScenarioError


class DisplayJSON:
    def __init__(self, dict):
        self.text = json.dumps(dict, indent=4)

    def __repr__(self):
        return self.text


class DisplayYAML:
    def __init__(self, dict):
        import yaml

        self.text = yaml.dump(dict)

    def __repr__(self):
        return self.text


class Scenario(Base, UserDict, ScenarioHtmlMixin):
    """A Scenario is a dictionary of keys/values.

    They can be used parameterize EDSL questions."""

    __documentation__ = "https://docs.expectedparrot.com/en/latest/scenarios.html"

    def __init__(self, data: Union[dict, None] = None, name: str = None):
        """Initialize a new Scenario.

        # :param data: A dictionary of keys/values for parameterizing questions.
        #"""
        if not isinstance(data, dict) and data is not None:
            raise EDSLScenarioError(
                "You must pass in a dictionary to initialize a Scenario."
            )

        self.data = data if data is not None else {}
        self.name = name

    def replicate(self, n: int) -> "ScenarioList":
        """Replicate a scenario n times to return a ScenarioList.

        :param n: The number of times to replicate the scenario.

        Example:

        >>> s = Scenario({"food": "wood chips"})
        >>> s.replicate(2)
        ScenarioList([Scenario({'food': 'wood chips'}), Scenario({'food': 'wood chips'})])
        """
        from edsl.scenarios.ScenarioList import ScenarioList

        return ScenarioList([copy.deepcopy(self) for _ in range(n)])

    @property
    def has_jinja_braces(self) -> bool:
        """Return whether the scenario has jinja braces. This matters for rendering.

        >>> s = Scenario({"food": "I love {{wood chips}}"})
        >>> s.has_jinja_braces
        True
        """
        for _, value in self.items():
            if isinstance(value, str):
                if "{{" in value and "}}" in value:
                    return True
        return False

    def convert_jinja_braces(
        self, replacement_left="<<", replacement_right=">>"
    ) -> Scenario:
        """Convert Jinja braces to some other character.

        >>> s = Scenario({"food": "I love {{wood chips}}"})
        >>> s.convert_jinja_braces()
        Scenario({'food': 'I love <<wood chips>>'})

        """
        new_scenario = Scenario()
        for key, value in self.items():
            if isinstance(value, str):
                new_scenario[key] = value.replace("{{", replacement_left).replace(
                    "}}", replacement_right
                )
            else:
                new_scenario[key] = value
        return new_scenario

    def __add__(self, other_scenario: "Scenario") -> "Scenario":
        """Combine two scenarios by taking the union of their keys

        If the other scenario is None, then just return self.

        :param other_scenario: The other scenario to combine with.

        Example:

        >>> s1 = Scenario({"price": 100, "quantity": 2})
        >>> s2 = Scenario({"color": "red"})
        >>> s1 + s2
        Scenario({'price': 100, 'quantity': 2, 'color': 'red'})
        >>> (s1 + s2).__class__.__name__
        'Scenario'
        """
        if other_scenario is None:
            return self
        else:
            data1 = copy.deepcopy(self.data)
            data2 = copy.deepcopy(other_scenario.data)
            s = Scenario(data1 | data2)
            return s

    def rename(
        self, old_name_or_replacement_dict: dict, new_name: Optional[str] = None
    ) -> "Scenario":
        """Rename the keys of a scenario.

        :param replacement_dict: A dictionary of old keys to new keys.

        Example:

        >>> s = Scenario({"food": "wood chips"})
        >>> s.rename({"food": "food_preference"})
        Scenario({'food_preference': 'wood chips'})

        >>> s = Scenario({"food": "wood chips"})
        >>> s.rename("food", "snack")
        Scenario({'snack': 'wood chips'})
        """
        if isinstance(old_name_or_replacement_dict, str) and new_name is not None:
            replacement_dict = {old_name_or_replacement_dict: new_name}
        else:
            replacement_dict = old_name_or_replacement_dict

        new_scenario = Scenario()
        for key, value in self.items():
            if key in replacement_dict:
                new_scenario[replacement_dict[key]] = value
            else:
                new_scenario[key] = value
        return new_scenario

    def table(self, tablefmt: str = "grid") -> str:
        from edsl.results.Dataset import Dataset

        keys = [key for key, value in self.items()]
        values = [value for key, value in self.items()]
        d = Dataset([{"key": keys}, {"value": values}])
        return d.table(tablefmt=tablefmt)

    def json(self):
        return DisplayJSON(self.to_dict(add_edsl_version=False))

    def yaml(self):
        import yaml

        return DisplayYAML(self.to_dict(add_edsl_version=False))

    def to_dict(self, add_edsl_version=True) -> dict:
        """Convert a scenario to a dictionary.

        Example:

        >>> s = Scenario({"food": "wood chips"})
        >>> s.to_dict()
        {'food': 'wood chips', 'edsl_version': '...', 'edsl_class_name': 'Scenario'}

        >>> s.to_dict(add_edsl_version = False)
        {'food': 'wood chips'}

        """
        from edsl.scenarios.FileStore import FileStore

        d = self.data.copy()
        for key, value in d.items():
            if isinstance(value, FileStore):
                d[key] = value.to_dict(add_edsl_version=add_edsl_version)
        if add_edsl_version:
            from edsl import __version__

            d["edsl_version"] = __version__
            d["edsl_class_name"] = "Scenario"

        return d

    def __hash__(self) -> int:
        """
        Return a hash of the scenario.

        Example:

        >>> s = Scenario({"food": "wood chips"})
        >>> hash(s)
        1153210385458344214
        """
        from edsl.utilities.utilities import dict_hash

        return dict_hash(self.to_dict(add_edsl_version=False))

    def print(self):
        from rich import print_json
        import json

        print_json(json.dumps(self.to_dict()))

    def __repr__(self):
        return "Scenario(" + repr(self.data) + ")"

    def to_dataset(self) -> "Dataset":
        # d = Dataset([{'a.b':[1,2,3,4]}])
        from edsl.results.Dataset import Dataset

        keys = [key for key, value in self.items()]
        values = [value for key, value in self.items()]
        return Dataset([{"key": keys}, {"value": values}])

    def _repr_html_(self):
        from tabulate import tabulate
        import reprlib

        d = self.to_dict(add_edsl_version=False)
        # return self.to_dataset()
        r = reprlib.Repr()
        r.maxstring = 70

        data = [[k, r.repr(v)] for k, v in d.items()]
        from tabulate import tabulate

        if hasattr(self, "__documentation__"):
            footer = f"<a href='{self.__documentation__}'>(docs)</a></p>"
        else:
            footer = ""

        table = str(tabulate(data, headers=["keys", "values"], tablefmt="html"))
        return f"<pre>{table}</pre>" + footer

    def select(self, list_of_keys: List[str]) -> "Scenario":
        """Select a subset of keys from a scenario.

        :param list_of_keys: The keys to select.

        Example:

        >>> s = Scenario({"food": "wood chips", "drink": "water"})
        >>> s.select(["food"])
        Scenario({'food': 'wood chips'})
        """
        new_scenario = Scenario()
        for key in list_of_keys:
            new_scenario[key] = self[key]
        return new_scenario

    def drop(self, list_of_keys: List[str]) -> "Scenario":
        """Drop a subset of keys from a scenario.

        :param list_of_keys: The keys to drop.

        Example:

        >>> s = Scenario({"food": "wood chips", "drink": "water"})
        >>> s.drop(["food"])
        Scenario({'drink': 'water'})
        """
        new_scenario = Scenario()
        for key in self.keys():
            if key not in list_of_keys:
                new_scenario[key] = self[key]
        return new_scenario

    def keep(self, list_of_keys: List[str]) -> "Scenario":
        """Keep a subset of keys from a scenario.

        :param list_of_keys: The keys to keep.

        Example:

        >>> s = Scenario({"food": "wood chips", "drink": "water"})
        >>> s.keep(["food"])
        Scenario({'food': 'wood chips'})
        """

        return self.select(list_of_keys)

    @classmethod
    def from_url(cls, url: str, field_name: Optional[str] = "text") -> "Scenario":
        """Creates a scenario from a URL.

        :param url: The URL to create the scenario from.
        :param field_name: The field name to use for the text.

        """
        import requests

        text = requests.get(url).text
        return cls({"url": url, field_name: text})

    @classmethod
    def from_file(cls, file_path: str, field_name: str) -> "Scenario":
        """Creates a scenario from a file.

        >>> import tempfile
        >>> with tempfile.NamedTemporaryFile(suffix=".txt", mode="w") as f:
        ...     _ = f.write("This is a test.")
        ...     _ = f.flush()
        ...     s = Scenario.from_file(f.name, "file")
        >>> s
        Scenario({'file': FileStore(path='...')})

        """
        from edsl.scenarios.FileStore import FileStore

        fs = FileStore(file_path)
        return cls({field_name: fs})

    @classmethod
    def from_image(
        cls, image_path: str, image_name: Optional[str] = None
    ) -> "Scenario":
        """
        Creates a scenario with a base64 encoding of an image.

        Args:
            image_path (str): Path to the image file.

        Returns:
            Scenario: A new Scenario instance with image information.

        """
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Image file not found: {image_path}")

        if image_name is None:
            image_name = os.path.basename(image_path).split(".")[0]

        return cls.from_file(image_path, image_name)

    @classmethod
    def from_pdf(cls, pdf_path):
        # Ensure the file exists
        import fitz

        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"The file {pdf_path} does not exist.")

        # Open the PDF file
        document = fitz.open(pdf_path)

        # Get the filename from the path
        filename = os.path.basename(pdf_path)

        # Iterate through each page and extract text
        text = ""
        for page_num in range(len(document)):
            page = document.load_page(page_num)
            blocks = page.get_text("blocks")  # Extract text blocks

            # Sort blocks by their vertical position (y0) to maintain reading order
            blocks.sort(key=lambda b: (b[1], b[0]))  # Sort by y0 first, then x0

            # Combine the text blocks in order
            for block in blocks:
                text += block[4] + "\n"

        # Create a dictionary for the combined text
        page_info = {"filename": filename, "text": text}
        return Scenario(page_info)

    @classmethod
    def from_docx(cls, docx_path: str) -> "Scenario":
        """Creates a scenario from the text of a docx file.

        :param docx_path: The path to the docx file.

        Example:

        >>> from docx import Document
        >>> doc = Document()
        >>> _ = doc.add_heading("EDSL Survey")
        >>> _ = doc.add_paragraph("This is a test.")
        >>> doc.save("test.docx")
        >>> s = Scenario.from_docx("test.docx")
        >>> s
        Scenario({'file_path': 'test.docx', 'text': 'EDSL Survey\\nThis is a test.'})
        >>> import os; os.remove("test.docx")
        """
        from docx import Document

        doc = Document(docx_path)

        # Extract all text
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Join the text from all paragraphs
        text = "\n".join(full_text)
        return Scenario({"file_path": docx_path, "text": text})

    @staticmethod
    def _line_chunks(text, num_lines: int) -> Generator[str, None, None]:
        """Split a text into chunks of a given size.

        :param text: The text to split.
        :param num_lines: The number of lines in each chunk.

        Example:

        >>> list(Scenario._line_chunks("This is a test.\\nThis is a test. This is a test.", 1))
        ['This is a test.', 'This is a test. This is a test.']
        """
        lines = text.split("\n")
        for i in range(0, len(lines), num_lines):
            chunk = "\n".join(lines[i : i + num_lines])
            yield chunk

    @staticmethod
    def _word_chunks(text, num_words: int) -> Generator[str, None, None]:
        """Split a text into chunks of a given size.

        :param text: The text to split.
        :param num_words: The number of words in each chunk.

        Example:

        >>> list(Scenario._word_chunks("This is a test.", 2))
        ['This is', 'a test.']
        """
        words = text.split()
        for i in range(0, len(words), num_words):
            chunk = " ".join(words[i : i + num_words])
            yield chunk

    def chunk(
        self,
        field,
        num_words: Optional[int] = None,
        num_lines: Optional[int] = None,
        include_original=False,
        hash_original=False,
    ) -> "ScenarioList":
        """Split a field into chunks of a given size.

        :param field: The field to split.
        :param num_words: The number of words in each chunk.
        :param num_lines: The number of lines in each chunk.
        :param include_original: Whether to include the original field in the new scenarios.
        :param hash_original: Whether to hash the original field in the new scenarios.

        If you specify `include_original=True`, the original field will be included in the new scenarios with an "_original" suffix.

        Either `num_words` or `num_lines` must be specified, but not both.

        The `hash_original` parameter is useful if you do not want to store the original text, but still want a unique identifier for it.

        Example:

        >>> s = Scenario({"text": "This is a test.\\nThis is a test.\\n\\nThis is a test."})
        >>> s.chunk("text", num_lines = 1)
        ScenarioList([Scenario({'text': 'This is a test.', 'text_chunk': 0}), Scenario({'text': 'This is a test.', 'text_chunk': 1}), Scenario({'text': '', 'text_chunk': 2}), Scenario({'text': 'This is a test.', 'text_chunk': 3})])

        >>> s.chunk("text", num_words = 2)
        ScenarioList([Scenario({'text': 'This is', 'text_chunk': 0}), Scenario({'text': 'a test.', 'text_chunk': 1}), Scenario({'text': 'This is', 'text_chunk': 2}), Scenario({'text': 'a test.', 'text_chunk': 3}), Scenario({'text': 'This is', 'text_chunk': 4}), Scenario({'text': 'a test.', 'text_chunk': 5})])

        >>> s = Scenario({"text": "Hello World"})
        >>> s.chunk("text", num_words = 1, include_original = True)
        ScenarioList([Scenario({'text': 'Hello', 'text_chunk': 0, 'text_original': 'Hello World'}), Scenario({'text': 'World', 'text_chunk': 1, 'text_original': 'Hello World'})])

        >>> s = Scenario({"text": "Hello World"})
        >>> s.chunk("text", num_words = 1, include_original = True, hash_original = True)
        ScenarioList([Scenario({'text': 'Hello', 'text_chunk': 0, 'text_original': 'b10a8db164e0754105b7a99be72e3fe5'}), Scenario({'text': 'World', 'text_chunk': 1, 'text_original': 'b10a8db164e0754105b7a99be72e3fe5'})])

        >>> s.chunk("text")
        Traceback (most recent call last):
        ...
        ValueError: You must specify either num_words or num_lines.

        >>> s.chunk("text", num_words = 1, num_lines = 1)
        Traceback (most recent call last):
        ...
        ValueError: You must specify either num_words or num_lines, but not both.
        """
        from edsl.scenarios.ScenarioList import ScenarioList

        if num_words is not None:
            chunks = list(self._word_chunks(self[field], num_words))

        if num_lines is not None:
            chunks = list(self._line_chunks(self[field], num_lines))

        if num_words is None and num_lines is None:
            raise ValueError("You must specify either num_words or num_lines.")

        if num_words is not None and num_lines is not None:
            raise ValueError(
                "You must specify either num_words or num_lines, but not both."
            )

        scenarios = []
        for i, chunk in enumerate(chunks):
            new_scenario = copy.deepcopy(self)
            new_scenario[field] = chunk
            new_scenario[field + "_chunk"] = i
            if include_original:
                if hash_original:
                    new_scenario[field + "_original"] = hashlib.md5(
                        self[field].encode()
                    ).hexdigest()
                else:
                    new_scenario[field + "_original"] = self[field]
            scenarios.append(new_scenario)
        return ScenarioList(scenarios)

    @classmethod
    @remove_edsl_version
    def from_dict(cls, d: dict) -> "Scenario":
        """Convert a dictionary to a scenario.

        Example:

        >>> Scenario.from_dict({"food": "wood chips"})
        Scenario({'food': 'wood chips'})
        """
        from edsl.scenarios.FileStore import FileStore

        for key, value in d.items():
            # TODO: we should check this better if its a FileStore + add remote security check against path traversal
            if (
                isinstance(value, dict) and "base64_string" in value and "path" in value
            ) or isinstance(value, FileStore):
                d[key] = FileStore.from_dict(value)
        return cls(d)

    def _table(self) -> tuple[dict, list]:
        """Prepare generic table data."""
        table_data = []
        for attr_name, attr_value in self.__dict__.items():
            table_data.append({"Attribute": attr_name, "Value": repr(attr_value)})
        column_names = ["Attribute", "Value"]
        return table_data, column_names

    def rich_print(self) -> "Table":
        """Display an object as a rich table."""
        from rich.table import Table

        table_data, column_names = self._table()
        table = Table(title=f"{self.__class__.__name__} Attributes")
        for column in column_names:
            table.add_column(column, style="bold")

        for row in table_data:
            row_data = [row[column] for column in column_names]
            table.add_row(*row_data)

        return table

    @classmethod
    def example(cls, randomize: bool = False, has_image=False) -> Scenario:
        """
        Returns an example Scenario instance.

        :param randomize: If True, adds a random string to the value of the example key.
        """
        if not has_image:
            addition = "" if not randomize else str(uuid4())
            return cls(
                {
                    "persona": f"A reseacher studying whether LLMs can be used to generate surveys.{addition}",
                }
            )
        else:
            return cls.from_image(cls.example_image())

    def code(self) -> List[str]:
        """Return the code for the scenario."""
        lines = []
        lines.append("from edsl.scenario import Scenario")
        lines.append(f"s = Scenario({self.data})")
        # return f"Scenario({self.data})"
        return lines


if __name__ == "__main__":
    import doctest

    doctest.testmod(optionflags=doctest.ELLIPSIS)
