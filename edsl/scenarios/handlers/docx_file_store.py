import os
import tempfile

from ..file_methods import FileMethods

class DocxMethods(FileMethods):
    suffix = "docx"

    def extract_text(self):
        from docx import Document

        self.doc = Document(self.path)

        # Extract all text
        full_text = []
        for para in self.doc.paragraphs:
            full_text.append(para.text)

        text = "\n".join(full_text)
        return text

    def view_system(self):
        import os
        import subprocess

        if os.path.exists(self.path):
            try:
                if (os_name := os.name) == "posix":
                    subprocess.run(["open", self.path], check=True)  # macOS
                elif os_name == "nt":
                    os.startfile(self.path)  # Windows
                else:
                    subprocess.run(["xdg-open", self.path], check=True)  # Linux
            except Exception as e:
                print(f"Error opening DOCX: {e}")
        else:
            print("DOCX file was not found.")

    def view_notebook(self):
        try:
            import mammoth
        except ImportError:
            print("mammoth is not installed. Please install it using 'pip install mammoth'.")
            return
        from IPython.display import HTML, display

        with open(self.path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = f"""
            <div style="width: 800px; height: 800px; padding: 20px; 
                       border: 1px solid #ccc; overflow-y: auto;">
                {result.value}
            </div>
            """
            display(HTML(html))

    def example(self):
        from docx import Document

        os.makedirs("test_dir", exist_ok=True)
        doc1 = Document()
        _ = doc1.add_heading("First Survey")
        doc1.save("test_dir/test1.docx")
        doc2 = Document()
        _ = doc2.add_heading("Second Survey")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc2.save(tmp.name)
            tmp.close()

        return tmp.name


if __name__ == "__main__":
    import doctest
    doctest.testmod()
