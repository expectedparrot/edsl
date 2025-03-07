from typing import Dict, List, Any, Optional, List


def is_hashable(v):
    try:
        hash(v)
        return True
    except TypeError:
        return False


class TreeNode:
    def __init__(self, key=None, value=None):
        self.key = key
        self.value = value
        self.children = {}


class Tree:
    def __init__(self, data: "Dataset", node_order: Optional[List[str]] = None):
        d = {}
        for entry in data:
            d.update(entry)
        self.data = d
        self.root = None

        self.node_order = node_order

        self.construct_tree(node_order)

    def unique_values_by_keys(self) -> dict:
        unique_values = {}
        for key, raw_values in self.data.items():
            values = [v if is_hashable(v) else str(v) for v in raw_values]
            unique_values[key] = list(set(values))
        return unique_values

    def construct_tree(self, node_order: Optional[List[str]] = None):
        # Validate node_order
        if node_order is None:
            unique_values = self.unique_values_by_keys()
            # Sort keys by number of unique values
            node_order = sorted(
                unique_values, key=lambda k: len(unique_values[k]), reverse=True
            )
        else:
            if not set(node_order).issubset(set(self.data.keys())):
                invalid_keys = set(node_order) - set(self.data.keys())
                raise ValueError(f"Invalid keys in node_order: {invalid_keys}")

        self.root = TreeNode()

        for i in range(len(self.data[node_order[0]])):
            current = self.root
            for level in node_order[:-1]:
                value = self.data[level][i]
                if not is_hashable(value):
                    value = str(value)
                if value not in current.children:
                    current.children[value] = TreeNode(key=level, value=value)
                current = current.children[value]

            leaf_key = node_order[-1]
            leaf_value = self.data[leaf_key][i]
            if not is_hashable(leaf_value):
                leaf_value = str(leaf_value)
            if leaf_value not in current.children:
                current.children[leaf_value] = TreeNode(key=leaf_key, value=leaf_value)

    def __repr__(self):
        if self.node_order is not None:
            return f"Tree(Dataset({self.data}), node_order={self.node_order})"
        else:
            return f"Tree(Dataset({self.data}))"

    def print_tree(
        self, node: Optional[TreeNode] = None, level: int = 0, print_keys: bool = False
    ):
        if node is None:
            node = self.root
            if node is None:
                print("Tree has not been constructed yet.")
                return

        if node.value is not None:
            if print_keys and node.key is not None:
                print("  " * level + f"{node.key}: {node.value}")
            else:
                print("  " * level + str(node.value))
        for child in node.children.values():
            self.print_tree(child, level + 1, print_keys)

    def to_docx(self, filename: Optional[str] = None):
        if filename is None:
            filename = "tree_structure.docx"

        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()

        # Create styles for headings
        for i in range(1, 10):  # Up to 9 levels of headings
            style_name = f"Heading {i}"
            if style_name not in doc.styles:
                doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        # Get or create the 'Body Text' style
        if "Body Text" not in doc.styles:
            body_style = doc.styles.add_style("Body Text", WD_STYLE_TYPE.PARAGRAPH)
        else:
            body_style = doc.styles["Body Text"]

        body_style.font.size = Pt(11)

        self._add_to_docx(doc, self.root, 0)
        import base64
        from io import BytesIO
        import base64

        # Save document to bytes buffer
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        base64_string = base64.b64encode(doc_buffer.getvalue()).decode("utf-8")
        from edsl.scenarios.FileStore import FileStore

        # Create and return FileStore instance
        return FileStore(
            path="tree_structure.docx",  # Default name
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            binary=True,
            suffix="docx",
            base64_string=base64_string,
        )

    def _repr_html_(self):
        """Returns an interactive HTML representation of the tree with collapsible sections."""

        # Generate a unique ID for this tree instance
        import uuid

        tree_id = f"tree_{uuid.uuid4().hex[:8]}"

        styles = f"""
        <div class="{tree_id}">
            <style>
                .{tree_id} details {{
                    margin-left: 20px;
                }}
                .{tree_id} summary {{
                    cursor: pointer;
                    margin: 2px 0;
                }}
                .{tree_id} .value {{
                    font-family: monospace;
                    background: #f5f5f5;
                    padding: 2px 6px;
                    border-radius: 3px;
                    margin: 1px 0;
                }}
                .{tree_id} .key {{
                    color: #666;
                    font-style: italic;
                }}
            </style>
        """

        def node_to_html(node, level=0, print_keys=True):
            if node is None:
                return "Tree has not been constructed yet."

            html = []

            if node.value is not None:
                # Create the node content
                content = []
                if print_keys and node.key is not None:
                    content.append(f'<span class="key">{node.key}: </span>')
                content.append(f'<span class="value">{node.value}</span>')
                content_html = "".join(content)

                if node.children:
                    # Node with children - removed the "open" condition to keep all nodes closed by default
                    html.append("<details>")
                    html.append(f"<summary>{content_html}</summary>")
                    for child in node.children.values():
                        html.append(node_to_html(child, level + 1, print_keys))
                    html.append("</details>")
                else:
                    # Leaf node
                    html.append(f"<div>{content_html}</div>")
            else:
                # Root node with no value
                if node.children:
                    for child in node.children.values():
                        html.append(node_to_html(child, level, print_keys))

            return "\n".join(html)

        tree_html = node_to_html(self.root)
        return f"{styles}{tree_html}</div>"

    def _add_to_docx(self, doc, node: TreeNode, level: int):
        if node.value is not None:
            if level == 0:
                doc.add_heading(str(node.value), level=level + 1)
            elif node.children:  # If the node has children, it's not the last level
                para = doc.add_paragraph(str(node.value))
                para.style = f"Heading {level+1}"
            else:  # If the node has no children, it's the last level (body text)
                para = doc.add_paragraph(str(node.value))
                para.style = "Body Text"

        # Process child nodes (moved outside the if block)
        for child in node.children.values():
            self._add_to_docx(doc, child, level + 1)

    def to_dict(self, node: Optional[TreeNode] = None) -> dict:
        """Converts the tree structure into a nested dictionary.
        
        Args:
            node: The current node being processed. Defaults to the root node.
            
        Returns:
            A nested dictionary representing the tree structure.
        """
        if node is None:
            node = self.root
            if node is None:
                return {}

        result = {}
        for value, child in node.children.items():
            if child.children:
                result[value] = self.to_dict(child)
            else:
                result[value] = None  # or {} if you prefer empty dict for leaf nodes
        return result

    @classmethod
    def example(cls) -> "Tree":
        """Creates an example Tree instance with geographic data.
        
        Returns:
            Tree: A sample tree with continent/country/city/population data
        
        Examples:
            >>> tree = Tree.example()     
            >>> tree.to_dict()
            {'North America': {'US': {'New York': {8419000: None}}, 'Canada': {'Toronto': {2930000: None}}}, 'Asia': {'China': {'Beijing': {21540000: None}}, 'Japan': {'Tokyo': {13960000: None}}}, 'Europe': {'France': {'Paris': {2161000: None}}}}
        """
        from edsl.results.Dataset import Dataset
        
        data = Dataset([
            {"continent": ["North America", "Asia", "Europe", "North America", "Asia"]},
            {"country": ["US", "China", "France", "Canada", "Japan"]},
            {"city": ["New York", "Beijing", "Paris", "Toronto", "Tokyo"]},
            {"population": [8419000, 21540000, 2161000, 2930000, 13960000]},
        ])
        
        tree = cls(data)
        tree.construct_tree(["continent", "country", "city", "population"])
        return tree


if __name__ == "__main__":
    #tree = Tree.example()
    #print(tree.to_dict())
    import doctest 
    doctest.testmod()
