import os
import tempfile
from docx import Document
from docx.shared import Inches
from .warning_utils import print_info, print_success, print_error


class ReportDOCX:
    """Helper class for generating DOCX reports from Report objects."""
    
    def __init__(self, report):
        """Initialize with a Report instance.
        
        Args:
            report: The Report instance to generate DOCX for.
        """
        self.report = report
    
    def generate(self, filename: str):
        """Generate a Word (.docx) report.

        Charts will be embedded as PNGs, tables converted to Word tables when possible.

        Args:
            filename (str): Destination docx file path (e.g. "report.docx").
        """
        print_info("Generating DOCX report…")

        doc = Document()

        # Add title and overview sections
        self._add_title_and_sections(doc)
        
        # Add individual analysis sections
        self._add_analysis_sections(doc)

        # Save the document
        doc.save(filename)
        print_success(f"DOCX report saved to {filename}")
    
    def _add_title_and_sections(self, doc):
        """Add title and main overview sections."""
        # Title & high-level sections
        doc.add_heading('Survey Report', level=1)

        # Add question summary table
        if self.report.include_questions_table:
            doc.add_heading('Question Summary', level=2)
            self._create_question_summary_table(doc)

        if self.report.include_overview:
            doc.add_heading('Survey Overview', level=2)
            doc.add_paragraph(self.report.survey_overview)

        if self.report.include_respondents_section:
            doc.add_heading('Respondent Overview', level=2)
            doc.add_paragraph(self.report.respondent_overview)

        if self.report.include_scenario_section:
            doc.add_heading('Scenario Overview', level=2)
            doc.add_paragraph(self.report.scenario_overview)
    
    def _add_analysis_sections(self, doc):
        """Add individual analysis sections."""
        for question_names, output_dict in self.report.items():
            section_title = self.report._format_question_header(question_names)
            doc.add_heading(section_title, level=2)
            
            # Add question metadata table
            self._create_question_metadata_table(doc, question_names)

            # Add output subsections
            self._add_output_subsections(doc, question_names, output_dict)
    
    def _add_output_subsections(self, doc, question_names, output_dict):
        """Add output subsections for a question analysis."""
        for output_name, output_obj in output_dict.items():
            # Skip tables that can't be analyzed (like AllResponsesTable)
            if hasattr(output_obj, 'can_be_analyzed') and not output_obj.can_be_analyzed:
                continue
                
            display_name = getattr(output_obj, 'pretty_short_name', output_name)
            doc.add_heading(display_name, level=3)

            # Add writeup
            self._add_writeup(doc, question_names, output_name)
            
            # Add visualization/table
            self._embed_output(doc, output_obj)
    
    def _add_writeup(self, doc, question_names, output_name):
        """Add writeup text for an output if available and enabled."""
        if question_names in self.report.writeups and output_name in self.report.writeups[question_names]:
            # Check if writeup is enabled for this analysis
            analysis_key = tuple(question_names) if isinstance(question_names, (list, tuple)) else (question_names,)
            writeup_enabled = self.report._analysis_writeup_filters.get(analysis_key, True)
            if writeup_enabled:
                writeup_text = self.report.writeups[question_names][output_name]
                doc.add_paragraph(writeup_text)
    
    def _embed_output(self, doc, output_obj):
        """Embed visualization or table output in Word document."""
        # Special handling for ThemeFinderOutput - embed all charts like in HTML
        if output_obj.__class__.__name__ == 'ThemeFinderOutput':
            return self._embed_theme_finder_output(doc, output_obj)
        
        # Standard handling for other output types
        # TRY 1: Attempt to embed as image via .svg property first, then fall back to .png
        svg_location = None
        png_location = None
        
        # Try SVG first
        try:
            svg_location = getattr(output_obj, "svg", None)
        except Exception:
            # SVG access failed, continue to PNG
            pass
        
        # Try PNG if no SVG
        if svg_location is None:
            try:
                png_location = getattr(output_obj, "png", None)
            except Exception:
                # PNG access failed, continue to dataframe/chart embedding
                pass

        # Embed SVG (convert to PNG for Word)
        if svg_location is not None:
            if self._try_embed_svg(doc, svg_location):
                return  # Successfully embedded

        # Embed PNG directly
        if png_location is not None:
            if self._try_embed_png(doc, png_location):
                return  # Successfully embedded

        # TRY 2: If image embedding failed or wasn't available, attempt DataFrame to table or Altair chart
        if self._try_embed_dataframe_or_chart(doc, output_obj):
            return  # Successfully embedded
        
        # Fallback: Add error message and throw exception
        error_msg = f"Failed to embed output: {str(output_obj)}"
        print_error(error_msg)
        raise RuntimeError(error_msg)
    
    def _embed_theme_finder_output(self, doc, theme_output):
        """Embed all ThemeFinder charts like in HTML report."""
        try:
            # Get all charts that should be included
            charts = theme_output.get_all_charts_for_docx()
            
            if not charts:
                print_error("No charts available from ThemeFinder")
                doc.add_paragraph("No theme analysis charts could be generated.")
                return True
            
            print_info(f"Embedding {len(charts)} ThemeFinder charts...")
            
            for chart_title, chart, description in charts:
                # Add chart title as heading
                chart_heading = doc.add_heading(chart_title, level=4)
                
                # Add description if provided
                if description:
                    doc.add_paragraph(description)
                
                # Try to embed the chart
                try:
                    if chart is not None and self._embed_altair_chart(doc, chart):
                        print_success(f"Successfully embedded: {chart_title}")
                    else:
                        print_error(f"Failed to embed: {chart_title}")
                        doc.add_paragraph(f"[Could not embed chart: {chart_title}]")
                except Exception as e:
                    print_error(f"Error embedding {chart_title}: {e}")
                    doc.add_paragraph(f"[Error embedding chart: {chart_title}]")
                
                # Add some spacing between charts
                doc.add_paragraph("")
            
            return True
            
        except Exception as e:
            print_error(f"Error in _embed_theme_finder_output: {e}")
            doc.add_paragraph(f"Error embedding theme analysis: {e}")
            return True  # Return True to continue processing
    
    def _try_embed_svg(self, doc, svg_location) -> bool:
        """Try to embed SVG by converting to PNG."""
        try:
            svg_path_tmp = str(svg_location.path)
            # Convert SVG to PNG for Word document since Word doesn't support SVG embedding
            from PIL import Image
            import cairosvg
            png_path_tmp = svg_path_tmp.replace('.svg', '.png')
            cairosvg.svg2png(url=svg_path_tmp, write_to=png_path_tmp)
            doc.add_picture(png_path_tmp, width=Inches(6))
            
            # Clean up temporary PNG
            if os.path.exists(png_path_tmp):
                try:
                    os.remove(png_path_tmp)
                except OSError:
                    pass
            return True
        except Exception as e:
            print_error(f"Error embedding SVG image: {e}")
            raise
    
    def _try_embed_png(self, doc, png_location) -> bool:
        """Try to embed PNG image."""
        try:
            png_path_tmp = str(png_location.path)
            doc.add_picture(png_path_tmp, width=Inches(6))
            return True
        except Exception as e:
            print_error(f"Error embedding PNG image: {e}")
            raise
    
    def _try_embed_dataframe_or_chart(self, doc, output_obj) -> bool:
        """Try to embed DataFrame as table or Altair chart as image."""
        try:
            if not hasattr(output_obj, "output"):
                return False
            
            # Try to get the output - this is where theme-finder might fail
            try:
                generic_output = output_obj.output()
            except Exception as e:
                print_error(f"Failed to get output from {type(output_obj)}: {e}")
                return False
                
            import pandas as pd
            import altair as alt
            
            # Check if it's a valid Altair chart
            if isinstance(generic_output, (alt.Chart, alt.LayerChart, alt.FacetChart)) or isinstance(generic_output, alt.TopLevelMixin):
                return self._embed_altair_chart(doc, generic_output)
            elif isinstance(generic_output, pd.DataFrame):
                return self._embed_dataframe_as_table(doc, generic_output)
            else:
                print_info(f"Unrecognized output type from {type(output_obj)}: {type(generic_output)}")
                print_info(f"Attempting PNG conversion...")
                
                # Try to convert to PNG if the object has save, to_png, or other image methods
                if self._try_convert_to_png_and_embed(doc, generic_output, output_obj):
                    return True
                    
                print_error(f"Cannot embed output type from {type(output_obj)}: {type(generic_output)}")
                return False
                
        except Exception as e:
            print_error(f"Error in _try_embed_dataframe_or_chart for {type(output_obj)}: {e}")
            return False
    
    def _try_convert_to_png_and_embed(self, doc, output_object, source_obj) -> bool:
        """Try to convert an unknown chart object to PNG and embed it."""
        import tempfile
        import os
        
        # Create temporary PNG file
        tmp_png = None
        try:
            tmp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            tmp_png = tmp_file.name
            tmp_file.close()
            
            # Try different methods to save the object as PNG
            saved = False
            
            # Method 1: Try .save() method (common for Altair-like objects)
            if hasattr(output_object, 'save') and callable(getattr(output_object, 'save')):
                try:
                    print_info("Trying .save() method...")
                    output_object.save(tmp_png)
                    saved = True
                except Exception as e:
                    print_info(f"save() method failed: {e}")
            
            # Method 2: Try .to_png() method
            if not saved and hasattr(output_object, 'to_png') and callable(getattr(output_object, 'to_png')):
                try:
                    print_info("Trying .to_png() method...")
                    output_object.to_png(tmp_png)
                    saved = True
                except Exception as e:
                    print_info(f"to_png() method failed: {e}")
            
            # Method 3: Try matplotlib-style savefig
            if not saved and hasattr(output_object, 'savefig') and callable(getattr(output_object, 'savefig')):
                try:
                    print_info("Trying .savefig() method...")
                    output_object.savefig(tmp_png, format='png')
                    saved = True
                except Exception as e:
                    print_info(f"savefig() method failed: {e}")
            
            # Method 4: Try using matplotlib if the object can be converted
            if not saved:
                try:
                    print_info("Trying matplotlib conversion...")
                    import matplotlib.pyplot as plt
                    
                    # Clear any existing plots
                    plt.clf()
                    
                    # Try to plot the object directly
                    if hasattr(output_object, 'plot'):
                        output_object.plot()
                    else:
                        # Try to convert to string and create a text plot
                        plt.text(0.5, 0.5, str(output_object)[:200], 
                                horizontalalignment='center',
                                verticalalignment='center',
                                transform=plt.gca().transAxes,
                                fontsize=10)
                        plt.title(f"Output from {type(source_obj).__name__}")
                    
                    plt.savefig(tmp_png, format='png', bbox_inches='tight', dpi=150)
                    plt.close()
                    saved = True
                except Exception as e:
                    print_info(f"matplotlib conversion failed: {e}")
            
            # If we managed to save something, try to embed it
            if saved and os.path.exists(tmp_png) and os.path.getsize(tmp_png) > 0:
                try:
                    doc.add_picture(tmp_png, width=Inches(6))
                    print_success(f"Successfully embedded PNG for {type(source_obj).__name__}")
                    return True
                except Exception as e:
                    print_error(f"Failed to add PNG to document: {e}")
            else:
                print_error("No valid PNG was generated")
                
        except Exception as e:
            print_error(f"Error in PNG conversion: {e}")
        finally:
            # Clean up temporary file
            if tmp_png and os.path.exists(tmp_png):
                try:
                    os.remove(tmp_png)
                except OSError:
                    pass
        
        return False
    
    def _embed_dataframe_as_table(self, doc, df: 'pd.DataFrame') -> bool:
        """Embed a pandas DataFrame as a Word table."""
        try:
            table = doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(df.columns):
                hdr_cells[i].text = str(col_name)
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for j, val in enumerate(row):
                    row_cells[j].text = str(val)
            return True
        except Exception as e:
            print_error(f"Error embedding DataFrame as table: {e}")
            raise
    
    def _embed_altair_chart(self, doc, chart) -> bool:
        """Embed an Altair chart by saving as PNG."""
        tmp_png = None
        try:
            tmp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
            tmp_png = tmp_file.name
            tmp_file.close()
            # alt.Chart.save uses altair_saver under the hood
            chart.save(tmp_png)
            doc.add_picture(tmp_png, width=Inches(6))
            return True
        except Exception as e:
            print_error(f"Error embedding Altair chart: {e}")
            raise
        finally:
            if tmp_png and os.path.exists(tmp_png):
                try:
                    os.remove(tmp_png)
                except OSError:
                    pass
    
    def _create_question_summary_table(self, doc):
        """Create Word table summarizing all questions and their inclusion in the report."""
        # Get all questions that are included in the report
        included_questions = set()
        for analysis in self.report.analyses:
            for question_name in analysis:
                included_questions.add(question_name)
        
        # Create table with headers
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Set header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Question Name'
        hdr_cells[1].text = 'Question Type'
        hdr_cells[2].text = 'Included in Report'
        
        # Add rows for each question
        for question in self.report.results.survey.questions:
            question_name = question.question_name
            question_type = question.question_type
            is_included = question_name in included_questions
            
            row_cells = table.add_row().cells
            row_cells[0].text = question_name
            row_cells[1].text = question_type
            row_cells[2].text = '✓' if is_included else '-'
    
    def _create_question_metadata_table(self, doc, question_names: list[str]):
        """Create Word table with question metadata."""
        total_respondents = len(self.report.results)
        
        for question_name in question_names:
            question = next((q for q in self.report.results.survey.questions if q.question_name == question_name), None)
            if question:
                question_type = question.question_type
                question_options = ""
                if hasattr(question, 'question_options') and question.question_options:
                    question_options = ", ".join(str(opt) for opt in question.question_options)
                elif hasattr(question, 'option_labels') and question.option_labels:
                    question_options = ", ".join(str(opt) for opt in question.option_labels)
                
                table = doc.add_table(rows=4, cols=2)
                table.style = 'Table Grid'
                
                # Question Name row
                table.rows[0].cells[0].text = 'Question Name'
                table.rows[0].cells[1].text = question_name
                
                # Question Type row
                table.rows[1].cells[0].text = 'Question Type'
                table.rows[1].cells[1].text = question_type
                
                # Question Options row
                table.rows[2].cells[0].text = 'Question Options'
                table.rows[2].cells[1].text = question_options
                
                # Total Respondents row
                table.rows[3].cells[0].text = 'Total Respondents'
                table.rows[3].cells[1].text = str(total_respondents)
                
                # Add separator between questions if there are multiple
                if len(question_names) > 1 and question_name != question_names[-1]:
                    doc.add_paragraph("") 