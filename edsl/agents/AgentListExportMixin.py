"""Mixin for exporting agent list to docx."""
from docx import Document


class AgentListExportMixin:
    """Mixin for exporting agent list to docx."""

    def docx(self) -> "Document":
        """Generate a docx document for the survey."""
        doc = Document()

        doc.add_heading("EDSL Auto-Generated Agent Description")

        doc.add_paragraph(f"\n")

        for index, agent in enumerate(self.data):
            # Add question as a paragraph
            h = doc.add_paragraph()
            h.add_run(f"Agent {index + 1}").bold = True

            p = doc.add_paragraph()
            # p.add_run(agent.persona)
            for key, value in agent.traits.items():
                p.add_run(f"{key}: ").bold = True
                p.add_run(f"{value}\n")

        return doc
