from __future__ import annotations
from collections import UserList
from docx import Document
from typing import Optional, Union
from edsl.agents import Agent


class AgentList(UserList):
    def __init__(self, data: Optional[list] = None):
        if data is not None:
            super().__init__(data)
        else:
            super().__init__()

    def to(self, question_or_survey: Union["Question", "Survey"]):
        return question_or_survey.by(*self)

    def update_traits(self, new_attributes: list):
        for agent in self.data:
            agent.update_traits(new_attributes)

    def print(self, html=False):
        html = ""
        for agent in self.data:
            html += agent.dict_to_html()
        from edsl.utilities.interface import gen_html_sandwich
        from edsl.utilities.interface import view_html

        view_html(gen_html_sandwich(html))

    def to_dict(self):
        return {"agent_list": [agent.to_dict() for agent in self.data]}

    @classmethod
    def from_dict(cls, data: dict) -> "AgentList":
        """Deserializes the dictionary back to an Agent List object."""
        agents = [Agent.from_dict(agent_dict) for agent_dict in data["agent_list"]]
        return cls(agents)

    def docx(self) -> Document:
        "Generates a docx document for the survey"
        doc = Document()

        doc.add_heading("EDSL Auto-Generated Agent Description")

        doc.add_paragraph(f"\n")

        for index, agent in enumerate(self.data):
            # Add question as a paragraph
            h = doc.add_paragraph()
            h.add_run(f"Agent {index + 1}").bold = True

            p = doc.add_paragraph()
            # p.add_run(agent.persona)
            for key, value in agent.traits.items():
                p.add_run(f"{key}: ").bold = True
                p.add_run(f"{value}\n")

        return doc


if __name__ == "__main__":
    a = AgentList([1, 2, 3])
    print(a)

    from edsl.questions.QuestionMultipleChoice import QuestionMultipleChoice
    from edsl.agents.Agent import Agent

    q = QuestionMultipleChoice(
        question_text="How are you feeling?",
        question_options=["Very sad.", "Sad.", "Neutral.", "Happy.", "Very happy."],
        question_name="feelings",
    )

    a = Agent(traits={"feeling": "Very sad."})
    agent_list = AgentList([a])
    results = agent_list.to(q).run()

    print(results)

    results2 = q.by(agent_list).run()

    print(results2)
