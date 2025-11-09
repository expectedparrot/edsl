from __future__ import annotations
import sys
import json
import random
from collections import UserList
from typing import Any, Union, Optional, TYPE_CHECKING, Callable

from ..base import PersistenceMixin, HashingMixin

from .dataset_tree import Tree
from .exceptions import DatasetKeyError, DatasetValueError, DatasetTypeError


from .display.table_display import TableDisplay

# from .smart_objects import FirstObject
from .dataset_operations_mixin import DatasetOperationsMixin

if TYPE_CHECKING:
    from ..surveys import Survey
    from ..questions import QuestionBase
    from ..jobs import Job  # noqa: F401


class Dataset(UserList, DatasetOperationsMixin, PersistenceMixin, HashingMixin):
    """
    A versatile data container for tabular data with powerful manipulation capabilities.

    The Dataset class is a fundamental data structure in EDSL that represents tabular data
    in a column-oriented format. It provides a rich set of methods for data manipulation,
    transformation, analysis, visualization, and export through the DatasetOperationsMixin.

    Key features:

    1. Column-oriented data structure optimized for LLM experiment results
    2. Rich data manipulation API similar to dplyr/pandas (filter, select, mutate, etc.)
    3. Visualization capabilities including tables, plots, and reports
    4. Export to various formats (CSV, Excel, SQLite, pandas, etc.)
    5. Serialization for storage and transport
    6. Tree-based data exploration

    A Dataset typically contains multiple columns, each represented as a dictionary
    with a single key-value pair. The key is the column name and the value is a list
    of values for that column. All columns must have the same length.

    The Dataset class inherits from:
    - UserList: Provides list-like behavior for storing column data
    - DatasetOperationsMixin: Provides data manipulation methods
    - PersistenceMixin: Provides serialization capabilities
    - HashingMixin: Provides hashing functionality for comparison and storage

    Datasets are typically created by transforming other EDSL container types like
    Results, AgentList, or ScenarioList, but can also be created directly from data.
    """

    def __init__(
        self, data: list[dict[str, Any]] = None, print_parameters: Optional[dict] = None
    ):
        """
        Initialize a new Dataset instance.

        Parameters:
            data: A list of dictionaries, where each dictionary represents a column
                 in the dataset. Each dictionary should have a single key-value pair,
                 where the key is the column name and the value is a list of values.
                 All value lists must have the same length.
            print_parameters: Optional dictionary of parameters controlling how the
                             dataset is displayed when printed.

        Examples:
            >>> # Create a dataset with two columns
            >>> d = Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])
            >>> len(d)
            3

            >>> # Dataset with a single column
            >>> Dataset([{'answer.how_feeling': ['OK', 'Great', 'Terrible']}])
            Dataset([{'answer.how_feeling': ['OK', 'Great', 'Terrible']}])
        """
        super().__init__(data)
        # self.data = data
        self.print_parameters = print_parameters

    def __len__(self) -> int:
        """Return the number of observations in the dataset.

        Need to override the __len__ method to return the number of observations in the dataset because
        otherwise, the UserList class would return the number of dictionaries in the dataset.

        >>> d = Dataset([{'a.b':[1,2,3,4]}])
        >>> len(d)
        4
        """
        _, values = list(self.data[0].items())[0]
        return len(values)

    def drop(self, field_name):
        """
        Returns a new Dataset with the specified field removed.

        Args:
            field_name (str): The name of the field to remove.

        Returns:
            Dataset: A new Dataset instance without the specified field.

        Raises:
            KeyError: If the field_name doesn't exist in the dataset.

        Examples:
            >>> from .dataset import Dataset
            >>> d = Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])
            >>> d.drop('a')
            Dataset([{'b': [4, 5, 6]}])

            >>> # Testing drop with nonexistent field raises DatasetKeyError - tested in unit tests
        """
        from .dataset import Dataset

        # Check if field exists in the dataset
        if field_name not in self.relevant_columns():
            raise DatasetKeyError(f"Field '{field_name}' not found in dataset")

        # Create a new dataset without the specified field
        new_data = [entry for entry in self.data if field_name not in entry]
        return Dataset(new_data)

    def tail(self, n: int = 5) -> Dataset:
        """Return the last n observations in the dataset.

        >>> d = Dataset([{'a.b':[1,2,3,4]}])
        >>> d.tail(2)
        Dataset([{'a.b': [3, 4]}])
        """
        new_data = []
        for observation in self.data:
            key, values = list(observation.items())[0]
            new_data.append({key: values[-n:]})
        return Dataset(new_data)

    def head(self, n: int = 5) -> Dataset:
        """Return the first n observations in the dataset.

        >>> d = Dataset([{'a.b':[1,2,3,4]}])
        >>> d.head(2)
        Dataset([{'a.b': [1, 2]}])
        """
        new_data = []
        for observation in self.data:
            key, values = list(observation.items())[0]
            new_data.append({key: values[:n]})
        return Dataset(new_data)

    # def expand(self, field):
    #     return self.to_scenario_list().expand(field)

    def keys(self) -> list[str]:
        """Return the keys of the dataset.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3, 4]}, {'b': [5, 6, 7, 8]}])
            >>> d.keys()
            ['a', 'b']

            >>> d = Dataset([{'x.y': [1, 2]}, {'z.w': [3, 4]}])
            >>> d.keys()
            ['x.y', 'z.w']
        """
        return [list(o.keys())[0] for o in self]

    def filter(self, expression) -> "Dataset":
        """Filter the dataset based on a boolean expression.

        Args:
            expression: A string expression that evaluates to a boolean value.
                       Can reference column names in the dataset.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3, 4]}, {'b': [5, 6, 7, 8]}])
            >>> d.filter('a > 2').data
            [{'a': [3, 4]}, {'b': [7, 8]}]

            >>> d = Dataset([{'x': ['a', 'b', 'c']}, {'y': [1, 2, 3]}])
            >>> d.filter('y < 3').data
            [{'x': ['a', 'b']}, {'y': [1, 2]}]
        """
        return self.to_scenario_list().filter(expression).to_dataset()

    def vibe_filter(
        self,
        criteria: str,
        *,
        model: str = "gpt-4o",
        temperature: float = 0.1,
        show_expression: bool = False,
    ) -> "Dataset":
        """
        Filter the dataset using natural language criteria.

        This method uses an LLM to generate a filter expression based on
        natural language criteria, then applies it using the dataset's filter method.

        Parameters:
            criteria: Natural language description of the filtering criteria.
                Examples:
                - "Keep only people over 30"
                - "Remove outliers in the satisfaction scores"
                - "Only include responses from the last month"
                - "Filter out any rows with missing data"
            model: OpenAI model to use for generating the filter (default: "gpt-4o")
            temperature: Temperature for generation (default: 0.1 for consistent logic)
            show_expression: If True, prints the generated filter expression

        Returns:
            Dataset: A new Dataset containing only the rows that match the criteria

        Examples:
            >>> from edsl.dataset import Dataset
            >>> d = Dataset([{'age': [25, 35, 42]}, {'occupation': ['student', 'engineer', 'teacher']}])
            >>> # filtered = d.vibe_filter("Keep only people over 30")

        Notes:
            - Requires OPENAI_API_KEY environment variable to be set
            - The LLM generates a filter expression using column names directly
            - Uses the dataset's built-in filter() method for safe evaluation
            - Use show_expression=True to see the generated filter logic
        """
        from .vibes.vibe_filter import VibeFilter

        # Get column names and sample data
        columns = self.relevant_columns()

        # Get a few sample rows to help the LLM understand the data structure
        sample_dicts = self.to_dicts(remove_prefix=False)[:5]

        # Create the filter generator
        filter_gen = VibeFilter(model=model, temperature=temperature)

        # Generate the filter expression
        filter_expr = filter_gen.create_filter(columns, sample_dicts, criteria)

        if show_expression:
            print(f"Generated filter expression: {filter_expr}")

        # Use the dataset's built-in filter method which returns Dataset
        return self.filter(filter_expr)

    def mutate(
        self, new_var_string: str, functions_dict: Optional[dict[str, Callable]] = None
    ) -> "Dataset":
        """Create new columns by applying functions to existing columns.

        Args:
            new_var_string: A string expression defining the new variable.
                           Can reference existing column names.
            functions_dict: Optional dictionary of custom functions to use in the expression.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])
            >>> d.mutate('c = a + b').data
            [{'a': [1, 2, 3]}, {'b': [4, 5, 6]}, {'c': [5, 7, 9]}]

            >>> d = Dataset([{'x': [1, 2, 3]}])
            >>> d.mutate('y = x * 2').data
            [{'x': [1, 2, 3]}, {'y': [2, 4, 6]}]
        """
        return (
            self.to_scenario_list().mutate(new_var_string, functions_dict).to_dataset()
        )

    def collapse(self, field: str, separator: Optional[str] = None) -> "Dataset":
        """Collapse multiple values in a field into a single value using a separator.

        Args:
            field: The name of the field to collapse.
            separator: Optional string to use as a separator between values.
                      Defaults to a space if not specified.

        Examples:
            >>> d = Dataset([{'words': [['hello', 'world'], ['good', 'morning']]}])
            >>> d.collapse('words').data
            [{'words': [[['hello', 'world'], ['good', 'morning']]]}]

            >>> d = Dataset([{'numbers': [1, 2, 3]}])
            >>> d.collapse('numbers', separator=',').data
            [{'numbers': ['1,2,3']}]
        """
        return self.to_scenario_list().collapse(field, separator).to_dataset()

    def long(self, *args, exclude_fields: Union[list[str], str] = None) -> Dataset:
        """Convert the dataset from wide to long format.

        Args:
            *args: Field names to exclude, passed as separate positional arguments.
                Example: .long('field1', 'field2')
            exclude_fields: Alternative way to specify fields to exclude. Can be:
                - A list of field names: ['field1', 'field2']
                - A comma-separated string: 'field1, field2'
                - Shorthand names without prefixes (e.g., 'city' instead of 'scenario.city')
                  are allowed if unambiguous

        Examples:
            >>> d = Dataset([{'a': [1, 2], 'b': [3, 4]}])
            >>> d.long().data
            [{'row': [0, 0, 1, 1]}, {'key': ['a', 'b', 'a', 'b']}, {'value': [1, 3, 2, 4]}]

            >>> d = Dataset([{'x': [1, 2], 'y': [3, 4], 'z': [5, 6]}])
            >>> d.long(exclude_fields=['z']).data
            [{'row': [0, 0, 1, 1]}, {'key': ['x', 'y', 'x', 'y']}, {'value': [1, 3, 2, 4]}, {'z': [5, 5, 6, 6]}]

            >>> # Can use comma-separated string
            >>> d.long(exclude_fields='z').data
            [{'row': [0, 0, 1, 1]}, {'key': ['x', 'y', 'x', 'y']}, {'value': [1, 3, 2, 4]}, {'z': [5, 5, 6, 6]}]

            >>> # Can use multiple positional arguments
            >>> d.long('y', 'z').data
            [{'row': [0, 1]}, {'key': ['x', 'x']}, {'value': [1, 2]}, {'y': [3, 4]}, {'z': [5, 6]}]
        """
        headers, data = self._tabular()

        # Handle different input methods
        if args:
            # Positional arguments provided: .long('field1', 'field2')
            exclude_fields = list(args)
        elif exclude_fields is None:
            exclude_fields = []
        elif isinstance(exclude_fields, str):
            # Comma-separated string: .long(exclude_fields='field1, field2')
            exclude_fields = [f.strip() for f in exclude_fields.split(",")]

        # Resolve shorthand field names to full names
        resolved_exclude_fields = []
        for field in exclude_fields:
            if field in headers:
                # Exact match found
                resolved_exclude_fields.append(field)
            else:
                # Look for fields that end with this name (after a dot)
                matches = [h for h in headers if h.endswith("." + field)]
                if len(matches) == 0:
                    raise DatasetValueError(
                        f"Field '{field}' not found in headers: {headers}"
                    )
                elif len(matches) > 1:
                    raise DatasetValueError(
                        f"Field '{field}' is ambiguous. Could refer to: {matches}"
                    )
                else:
                    resolved_exclude_fields.append(matches[0])

        exclude_fields = resolved_exclude_fields

        # Initialize result dictionaries for each column
        result_dict = {}

        for index, row in enumerate(data):
            row_values = dict(zip(headers, row))
            excluded_values = {field: row_values[field] for field in exclude_fields}

            # Transform non-excluded fields to long format
            for header, value in row_values.items():
                if header not in exclude_fields:
                    # Initialize lists in result_dict if needed
                    if not result_dict:
                        result_dict = {
                            "row": [],
                            "key": [],
                            "value": [],
                            **{field: [] for field in exclude_fields},
                        }

                    # Add values to each column
                    result_dict["row"].append(index)
                    result_dict["key"].append(header)
                    result_dict["value"].append(value)
                    for field in exclude_fields:
                        result_dict[field].append(excluded_values[field])

        return Dataset([{k: v} for k, v in result_dict.items()])

    def wide(self) -> "Dataset":
        """
        Convert a long-format dataset (with row, key, value columns) to wide format.

        Examples:
            >>> d = Dataset([{'row': [0, 0, 1, 1]}, {'key': ['a', 'b', 'a', 'b']}, {'value': [1, 3, 2, 4]}])
            >>> d.wide().data
            [{'a': [1, 2]}, {'b': [3, 4]}]

            >>> d = Dataset([{'row': [0, 0, 1, 1]}, {'key': ['x', 'y', 'x', 'y']}, {'value': [1, 3, 2, 4]}, {'z': [5, 5, 6, 6]}])
            >>> d.wide().data
            [{'x': [1, 2]}, {'y': [3, 4]}, {'z': [5, 6]}]
        """
        # Extract the component arrays
        row_dict = next(col for col in self if "row" in col)
        key_dict = next(col for col in self if "key" in col)
        value_dict = next(col for col in self if "value" in col)

        rows = row_dict["row"]
        keys = key_dict["key"]
        values = value_dict["value"]

        if not (len(rows) == len(keys) == len(values)):
            raise DatasetValueError("All input arrays must have the same length")

        # Get unique keys and row indices
        unique_keys = sorted(set(keys))
        unique_rows = sorted(set(rows))

        # Create a dictionary to store the result
        result = {key: [None] * len(unique_rows) for key in unique_keys}

        # Populate the result dictionary
        for row_idx, key, value in zip(rows, keys, values):
            # Find the position in the output array for this row
            output_row_idx = unique_rows.index(row_idx)
            result[key][output_row_idx] = value

        # Add any additional columns that weren't part of the key-value transformation
        additional_columns = []
        for col in self:
            col_key = list(col.keys())[0]
            if col_key not in ["row", "key", "value"]:
                # Get unique values for this column
                unique_values = []
                for row_idx in unique_rows:
                    # Find the first occurrence of this row index
                    for i, r in enumerate(rows):
                        if r == row_idx:
                            unique_values.append(col[col_key][i])
                            break
                additional_columns.append({col_key: unique_values})

        # Convert to list of column dictionaries format
        result_columns = [{key: values} for key, values in result.items()]
        return Dataset(result_columns + additional_columns)

    def __repr__(self) -> str:
        """Return a string representation of the dataset.

        Uses traditional repr format when running doctests, otherwise uses
        rich-based display for better readability. In Jupyter notebooks,
        returns a minimal string since _repr_html_ handles the display.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])
            >>> repr(d)
            "Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])"

            >>> d = Dataset([{'x': ['a', 'b']}])
            >>> repr(d)
            "Dataset([{'x': ['a', 'b']}])"
        """
        import os

        if os.environ.get("EDSL_RUNNING_DOCTESTS") == "True":
            return self._eval_repr_()

        # Check if we're in a Jupyter notebook environment
        # If so, return minimal representation since _repr_html_ will handle display
        try:
            from IPython import get_ipython

            ipy = get_ipython()
            if ipy is not None and "IPKernelApp" in ipy.config:
                # We're in a Jupyter notebook/kernel, not IPython terminal
                return "Dataset(...)"
        except (NameError, ImportError):
            pass

        return self._summary_repr()

    def _eval_repr_(self) -> str:
        """Return an eval-able string representation of the dataset.

        This representation can be used with eval() to recreate the Dataset object.
        Used primarily for doctests and debugging.
        """
        return f"Dataset({self.data})"

    def _summary_repr(
        self, max_rows: int = 5, max_cols: int = 10, max_value_length: int = 30
    ) -> str:
        """Generate a summary representation of the Dataset with Rich formatting.

        Args:
            max_rows: Maximum number of rows to show before truncating
            max_cols: Maximum number of columns to show before truncating
            max_value_length: Maximum length of a value before truncating
        """
        from rich.console import Console
        from rich.text import Text
        import io
        from edsl.config import RICH_STYLES

        # Build the Rich text with consistent styling
        output = Text()
        output.append("Dataset(\n", style=RICH_STYLES["primary"])

        # Handle empty dataset
        if not self.data:
            output.append("    data=[]\n", style=RICH_STYLES["dim"])
            output.append(")", style=RICH_STYLES["primary"])
            console = Console(file=io.StringIO(), force_terminal=True, width=120)
            console.print(output, end="")
            return console.file.getvalue()

        num_obs = len(self)
        num_cols = len(self.keys())

        output.append(
            f"    num_observations={num_obs},\n", style=RICH_STYLES["default"]
        )
        output.append(f"    num_columns={num_cols},\n", style=RICH_STYLES["default"])

        # Show column names
        if num_cols > 0:
            cols = self.keys()
            output.append("    columns=[\n", style=RICH_STYLES["default"])

            for i, col in enumerate(cols[:max_cols]):
                col_str = str(col)
                if len(col_str) > 40:
                    col_str = col_str[:37] + "..."
                output.append("        ", style=RICH_STYLES["default"])
                output.append(f"'{col_str}'", style=RICH_STYLES["key"])
                output.append(",\n", style=RICH_STYLES["default"])

            if num_cols > max_cols:
                output.append(
                    f"        ... ({num_cols - max_cols} more)\n",
                    style=RICH_STYLES["dim"],
                )

            output.append("    ],\n", style=RICH_STYLES["default"])

            # Show preview of data
            if num_obs > 0:
                output.append("    preview={\n", style=RICH_STYLES["default"])
                headers, rows = self._tabular()

                # Show column headers and first few values
                for i, col in enumerate(headers[:max_cols]):
                    col_values = [row[i] for row in rows[:max_rows]]

                    # Format values with truncation
                    formatted_values = []
                    for val in col_values:
                        val_str = repr(val)
                        if len(val_str) > max_value_length:
                            val_str = val_str[: max_value_length - 3] + "..."
                        formatted_values.append(val_str)

                    output.append("        ", style=RICH_STYLES["default"])
                    output.append(f"'{col}'", style=RICH_STYLES["secondary"])
                    output.append(
                        f": [{', '.join(formatted_values)}", style=RICH_STYLES["value"]
                    )

                    if num_obs > max_rows:
                        output.append(", ...", style=RICH_STYLES["dim"])

                    output.append("],\n", style=RICH_STYLES["default"])

                if num_cols > max_cols:
                    output.append(
                        f"        ... ({num_cols - max_cols} more columns)\n",
                        style=RICH_STYLES["dim"],
                    )

                output.append("    }\n", style=RICH_STYLES["default"])

        output.append(")", style=RICH_STYLES["primary"])

        # Render to string
        console = Console(file=io.StringIO(), force_terminal=True, width=120)
        console.print(output, end="")
        return console.file.getvalue()

    def write(self, filename: str, tablefmt: Optional[str] = None) -> None:
        """Write the dataset to a file in the specified format.

        Args:
            filename: The name of the file to write to.
            tablefmt: Optional format for the table (e.g., 'csv', 'html', 'latex').
        """
        return self.table(tablefmt=tablefmt).write(filename)

    def _repr_html_(self):
        """Return an HTML representation of the dataset for Jupyter notebooks.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])
            >>> html = d._repr_html_()
            >>> isinstance(html, str)
            True
        """
        return self.table(print_parameters=self.print_parameters)._repr_html_()

    def _tabular(self) -> tuple[list[str], list[list[Any]]]:
        """Convert the dataset to a tabular format (headers and rows).

        Returns:
            A tuple containing:
            - List of column headers
            - List of rows, where each row is a list of values

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])
            >>> headers, rows = d._tabular()
            >>> headers
            ['a', 'b']
            >>> rows
            [[1, 4], [2, 5], [3, 6]]

            >>> d = Dataset([{'x': ['a', 'b']}, {'y': [1, 2]}])
            >>> headers, rows = d._tabular()
            >>> headers
            ['x', 'y']
            >>> rows
            [['a', 1], ['b', 2]]
        """
        # Extract headers
        headers = []
        for entry in self.data:
            headers.extend(entry.keys())
        headers = list(dict.fromkeys(headers))  # Ensure unique headers

        # Extract data
        if not self.data:
            return headers, []

        # Get max_len, but handle empty datasets
        lengths = [len(values) for entry in self.data for values in entry.values()]
        max_len = max(lengths) if lengths else 0
        rows = []
        for i in range(max_len):
            row = []
            for header in headers:
                for entry in self.data:
                    if header in entry:
                        values = entry[header]
                        row.append(values[i] if i < len(values) else None)
                        break
                else:
                    row.append(None)  # Default to None if header is missing
            rows.append(row)

        return headers, rows

    def _key_to_value(self, key: str) -> Any:
        """Retrieve the value associated with the given key from the dataset.

        Args:
            key: The key to look up in the dataset.

        Returns:
            The list of values associated with the key.

        Examples:
            >>> d = Dataset([{'a.b': [1, 2, 3, 4]}])
            >>> d._key_to_value('a.b')
            [1, 2, 3, 4]

            >>> d = Dataset([{'x.y': [1, 2]}, {'z.w': [3, 4]}])
            >>> d._key_to_value('w')
            [3, 4]
        """
        potential_matches = []
        for data_dict in self.data:
            data_key, data_values = list(data_dict.items())[0]
            if key == data_key:
                return data_values
            if key == data_key.split(".")[-1]:
                potential_matches.append((data_key, data_values))

        if len(potential_matches) == 1:
            return potential_matches[0][1]
        elif len(potential_matches) > 1:
            from .exceptions import DatasetKeyError

            raise DatasetKeyError(
                f"Key '{key}' found in more than one location: {[m[0] for m in potential_matches]}"
            )

        from .exceptions import DatasetKeyError

        raise DatasetKeyError(f"Key '{key}' not found in any of the dictionaries.")

    def first(self) -> dict[str, Any]:
        """Get the first value of the first key in the first dictionary.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3, 4]}, {'b': [5, 6, 7, 8]}])
            >>> d.first()
            1

            >>> d = Dataset([{'x': ['first', 'second']}])
            >>> d.first()
            'first'
        """

        def get_values(d):
            """Get the values of the first key in the dictionary."""
            return list(d.values())[0]

        return get_values(self.data[0])[0]

    def latex(self, **kwargs):
        """Return a LaTeX representation of the dataset.

        Args:
            **kwargs: Additional arguments to pass to the table formatter.


        """
        return self.table().latex()

    def remove_prefix(self) -> Dataset:
        """Remove the prefix from column names that contain dots.

        Examples:
            >>> d = Dataset([{'a.b': [1, 2, 3]}, {'c.d': [4, 5, 6]}])
            >>> d.remove_prefix().data
            [{'b': [1, 2, 3]}, {'d': [4, 5, 6]}]

            >>> d = Dataset([{'x.y.z': [1, 2]}, {'a.b.c': [3, 4]}])
            >>> d.remove_prefix().data
            [{'y': [1, 2]}, {'b': [3, 4]}]
        """
        new_data = []
        for observation in self.data:
            key, values = list(observation.items())[0]
            if "." in key:
                new_key = key.split(".")[1]
                new_data.append({new_key: values})
            else:
                new_data.append({key: values})
        return Dataset(new_data)

    def print(self, pretty_labels=None, **kwargs):
        """
        Print the dataset in a formatted way.

        Args:
            pretty_labels: A dictionary mapping column names to their display names
            **kwargs: Additional arguments
                format: The output format ("html", "markdown", "rich", "latex")

        Returns:
            TableDisplay object

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])
            >>> display = d.print(format='rich')
            >>> display is not None
            True

            >>> d = Dataset([{'long_column_name': [1, 2]}])
            >>> display = d.print(pretty_labels={'long_column_name': 'Short'})
            >>> display is not None
            True
        """
        if "format" in kwargs:
            if kwargs["format"] not in ["html", "markdown", "rich", "latex"]:
                raise DatasetValueError(f"Format '{kwargs['format']}' not supported.")

            # If rich format is requested, set tablefmt accordingly
            if kwargs["format"] == "rich":
                kwargs["tablefmt"] = "rich"

        if pretty_labels is None:
            pretty_labels = {}
        else:
            return self.rename(pretty_labels).print(**kwargs)

        # Pass through any tablefmt parameter
        tablefmt = kwargs.get("tablefmt", "rich")
        return self.table(tablefmt=tablefmt)

    def rename(self, rename_dic) -> Dataset:
        """Rename columns in the dataset according to the provided dictionary.

        Raises:
            DatasetKeyError: If any key in rename_dic does not exist in the dataset.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])
            >>> d.rename({'a': 'x', 'b': 'y'}).data
            [{'x': [1, 2, 3]}, {'y': [4, 5, 6]}]

            >>> d = Dataset([{'old_name': [1, 2]}])
            >>> d.rename({'old_name': 'new_name'}).data
            [{'new_name': [1, 2]}]
        """
        from .exceptions import DatasetKeyError

        # Collect all existing keys in the dataset
        existing_keys = set()
        for observation in self.data:
            if observation:  # Ensure observation is not empty
                key = list(observation.keys())[0]
                existing_keys.add(key)

        # Validate that all keys in rename_dic exist in the dataset
        missing_keys = set(rename_dic.keys()) - existing_keys
        if missing_keys:
            raise DatasetKeyError(
                f"The following keys in rename_dic are not present in the dataset: {missing_keys}"
            )

        new_data = []
        for observation in self.data:
            key, values = list(observation.items())[0]
            new_key = rename_dic.get(key, key)
            new_data.append({new_key: values})
        return Dataset(new_data)

    def merge(self, other: Dataset, by_x, by_y) -> Dataset:
        """Merge the dataset with another dataset on the given keys.

        Examples:
            >>> d1 = Dataset([{'key': [1, 2, 3]}, {'value1': ['a', 'b', 'c']}])
            >>> d2 = Dataset([{'key': [2, 3, 4]}, {'value2': ['x', 'y', 'z']}])
            >>> merged = d1.merge(d2, 'key', 'key')
            >>> len(merged.data[0]['key'])
            3

            >>> d1 = Dataset([{'id': [1, 2]}, {'name': ['Alice', 'Bob']}])
            >>> d2 = Dataset([{'id': [2, 3]}, {'age': [25, 30]}])
            >>> merged = d1.merge(d2, 'id', 'id')
            >>> len(merged.data[0]['id'])
            2
        """
        df1 = self.to_pandas()
        df2 = other.to_pandas()
        merged_df = df1.merge(df2, how="left", left_on=by_x, right_on=by_y)
        return Dataset.from_pandas_dataframe(merged_df)

    def to(self, survey_or_question: Union["Survey", "QuestionBase"]) -> "Job":
        """Transform the dataset using a survey or question.

        Args:
            survey_or_question: Either a Survey or QuestionBase object to apply to the dataset.

        Examples:
            >>> from edsl import QuestionFreeText
            >>> from edsl.jobs import Jobs
            >>> d = Dataset([{'name': ['Alice', 'Bob']}])
            >>> q = QuestionFreeText(question_text="How are you, {{ name }}?", question_name="how_feeling")
            >>> job = d.to(q)
            >>> isinstance(job, Jobs)
            True
        """
        from ..surveys import Survey
        from ..questions import QuestionBase

        if isinstance(survey_or_question, Survey):
            return survey_or_question.by(self.to_scenario_list())
        elif isinstance(survey_or_question, QuestionBase):
            return Survey([survey_or_question]).by(self.to_scenario_list())

    def select(self, *keys) -> Dataset:
        """Return a new dataset with only the selected keys.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3, 4]}, {'b': [5, 6, 7, 8]}, {'c': [9, 10, 11, 12]}])
            >>> d.select('a', 'c').data
            [{'a': [1, 2, 3, 4]}, {'c': [9, 10, 11, 12]}]

            >>> d = Dataset([{'x': [1, 2]}, {'y': [3, 4]}])
            >>> d.select('x').data
            [{'x': [1, 2]}]
        """
        for key in keys:
            if key not in self.keys():
                from .exceptions import DatasetValueError

                raise DatasetValueError(
                    f"Key '{key}' not found in the dataset. "
                    f"Available keys: {self.keys()}"
                )

        if isinstance(keys, str):
            keys = [keys]

        new_data = []
        for observation in self.data:
            observation_key = list(observation.keys())[0]
            if observation_key in keys:
                new_data.append(observation)
        return Dataset(new_data)

    def to_json(self):
        """Return a JSON representation of the dataset.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])
            >>> d.to_json()
            [{'a': [1, 2, 3]}, {'b': [4, 5, 6]}]

            >>> d = Dataset([{'x': ['a', 'b']}])
            >>> d.to_json()
            [{'x': ['a', 'b']}]
        """
        return json.loads(
            json.dumps(self.data)
        )  # janky but I want to make sure it's serializable & deserializable

    def shuffle(self, seed=None) -> Dataset:
        """Return a new dataset with the observations shuffled.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3, 4]}, {'b': [5, 6, 7, 8]}])
            >>> shuffled = d.shuffle(seed=42)
            >>> len(shuffled.data[0]['a']) == len(d.data[0]['a'])
            True

            >>> d = Dataset([{'x': ['a', 'b', 'c']}])
            >>> shuffled = d.shuffle(seed=123)
            >>> set(shuffled.data[0]['x']) == set(d.data[0]['x'])
            True
        """
        if seed is not None:
            random.seed(seed)

        indices = None

        for entry in self:
            key, values = list(entry.items())[0]
            if indices is None:
                indices = list(range(len(values)))
                random.shuffle(indices)
            entry[key] = [values[i] for i in indices]

        return self

    def expand_field(self, field):
        """Expand a field in the dataset.

        Renamed to avoid conflict with the expand method defined earlier.
        """
        return self.to_scenario_list().expand(field).to_dataset()

    def sample(
        self,
        n: int = None,
        frac: float = None,
        with_replacement: bool = True,
        seed: Union[str, int, float] = None,
    ) -> Dataset:
        """Return a new dataset with a sample of the observations.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3, 4, 5]}, {'b': [6, 7, 8, 9, 10]}])
            >>> sampled = d.sample(n=3, seed=42)
            >>> len(sampled.data[0]['a'])
            3

            >>> d = Dataset([{'x': ['a', 'b', 'c', 'd']}])
            >>> sampled = d.sample(frac=0.5, seed=123)
            >>> len(sampled.data[0]['x'])
            2
        """
        if seed is not None:
            random.seed(seed)

        # Validate the input for sampling parameters
        if n is None and frac is None:
            from .exceptions import DatasetValueError

            raise DatasetValueError(
                "Either 'n' or 'frac' must be provided for sampling."
            )

        if n is not None and frac is not None:
            from .exceptions import DatasetValueError

            raise DatasetValueError("Only one of 'n' or 'frac' should be specified.")

        # Get the length of the lists from the first entry
        first_key, first_values = list(self[0].items())[0]
        total_length = len(first_values)

        # Determine the number of samples based on 'n' or 'frac'
        if n is None:
            n = int(total_length * frac)

        if not with_replacement and n > total_length:
            from .exceptions import DatasetValueError

            raise DatasetValueError(
                "Sample size cannot be greater than the number of available elements when sampling without replacement."
            )

        # Sample indices based on the method chosen
        if with_replacement:
            indices = [random.randint(0, total_length - 1) for _ in range(n)]
        else:
            indices = random.sample(range(total_length), k=n)

        # Apply the same indices to all entries
        for entry in self:
            key, values = list(entry.items())[0]
            entry[key] = [values[i] for i in indices]

        return self

    def get_sort_indices(self, lst: list[Any], reverse: bool = False) -> list[int]:
        """
        Return the indices that would sort the list, using either numpy or pure Python.
        None values are placed at the end of the sorted list.

        Args:
            lst: The list to be sorted
            reverse: Whether to sort in descending order
            use_numpy: Whether to use numpy implementation (falls back to pure Python if numpy is unavailable)

        Returns:
            A list of indices that would sort the list
        """
        enumerated = list(enumerate(lst))
        sorted_pairs = sorted(
            enumerated, key=lambda x: (x[1] is None, x[1]), reverse=reverse
        )
        return [index for index, _ in sorted_pairs]

    def order_by(self, sort_key: str, reverse: bool = False) -> Dataset:
        """Return a new dataset with the observations sorted by the given key.

        Args:
            sort_key: The key to sort the observations by
            reverse: Whether to sort in reverse order

        Examples:
            >>> d = Dataset([{'a': [3, 1, 4, 1, 5]}, {'b': ['x', 'y', 'z', 'w', 'v']}])
            >>> sorted_d = d.order_by('a')
            >>> sorted_d.data
            [{'a': [1, 1, 3, 4, 5]}, {'b': ['y', 'w', 'x', 'z', 'v']}]

            >>> d = Dataset([{'a': [3, 1, 4, 1, 5]}, {'b': ['x', 'y', 'z', 'w', 'v']}])
            >>> sorted_d = d.order_by('a', reverse=True)
            >>> sorted_d.data
            [{'a': [5, 4, 3, 1, 1]}, {'b': ['v', 'z', 'x', 'y', 'w']}]

            >>> d = Dataset([{'a': [3, None, 1, 4, None]}, {'b': ['x', 'y', 'z', 'w', 'v']}])
            >>> sorted_d = d.order_by('a')
            >>> sorted_d.data
            [{'a': [1, 3, 4, None, None]}, {'b': ['z', 'x', 'w', 'y', 'v']}]
        """
        number_found = 0
        for obs in self.data:
            key, values = list(obs.items())[0]
            if sort_key == key or sort_key == key.split(".")[-1]:
                relevant_values = values
                number_found += 1

        if number_found == 0:
            raise DatasetKeyError(
                f"Key '{sort_key}' not found in any of the dictionaries."
            )
        elif number_found > 1:
            raise DatasetKeyError(
                f"Key '{sort_key}' found in more than one dictionary."
            )

        sort_indices_list = self.get_sort_indices(relevant_values, reverse=reverse)
        new_data = []
        for observation in self.data:
            key, values = list(observation.items())[0]
            new_values = [values[i] for i in sort_indices_list]
            new_data.append({key: new_values})

        return Dataset(new_data)

    def tree(self, node_order: Optional[list[str]] = None) -> Tree:
        """Return a tree representation of the dataset.

        >>> d = Dataset([{'a':[1,2,3,4]}, {'b':[4,3,2,1]}])
        >>> d.tree()
        Tree(Dataset({'a': [1, 2, 3, 4], 'b': [4, 3, 2, 1]}), node_order=['a', 'b'])
        """
        if node_order is None:
            node_order = self.keys()
        return Tree(self, node_order=node_order)

    def table(
        self,
        *fields,
        tablefmt: Optional[str] = "rich",
        max_rows: Optional[int] = None,
        pretty_labels=None,
        print_parameters: Optional[dict] = None,
    ):
        if pretty_labels is not None:
            new_fields = []
            for field in fields:
                new_fields.append(pretty_labels.get(field, field))
            return self.rename(pretty_labels).table(
                *new_fields, tablefmt=tablefmt, max_rows=max_rows
            )

        self.print_parameters = print_parameters

        headers, data = self._tabular()

        if tablefmt is not None and tablefmt != "rich":
            # Rich format is handled separately, so we don't validate it against tabulate_formats
            from tabulate import tabulate_formats

            if tablefmt not in tabulate_formats:
                print(
                    f"Error: The following table format is not supported: {tablefmt}",
                    file=sys.stderr,
                )
                print(
                    f"\nAvailable formats are: {tabulate_formats} and 'rich'",
                    file=sys.stderr,
                )
                return None

        if max_rows:
            if len(data) < max_rows:
                max_rows = None

        if fields:
            full_data = data
            data = []
            indices = []
            for field in fields:
                if field not in headers:
                    print(
                        f"Error: The following field was not found: {field}",
                        file=sys.stderr,
                    )
                    print(f"\nAvailable fields are: {headers}", file=sys.stderr)

                    # Optional: Suggest similar fields using difflib
                    import difflib

                    matches = difflib.get_close_matches(field, headers)
                    if matches:
                        print(f"\nDid you mean: {matches[0]} ?", file=sys.stderr)
                    return None
                indices.append(headers.index(field))
            headers = fields
            for row in full_data:
                data.append([row[i] for i in indices])

        if max_rows is not None:
            if max_rows > len(data):
                from .exceptions import DatasetValueError

                raise DatasetValueError(
                    "max_rows cannot be greater than the number of rows in the dataset."
                )
            last_line = data[-1]
            spaces = len(data[max_rows])
            filler_line = ["." for i in range(spaces)]
            data = data[:max_rows]
            data.append(filler_line)
            data.append(last_line)

        return TableDisplay(
            data=data, headers=headers, tablefmt=tablefmt, raw_data_set=self
        )

    def summary(self) -> "Dataset":
        """Return a summary of the dataset.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])
            >>> d.summary().data
            [{'num_observations': [3]}, {'keys': [['a', 'b']]}]
        """
        return Dataset([{"num_observations": [len(self)]}, {"keys": [self.keys()]}])

    @classmethod
    def example(self, n: int = None) -> "Dataset":
        """Return an example dataset.

        Examples:
            >>> Dataset.example()
            Dataset([{'a': [1, 2, 3, 4]}, {'b': [4, 3, 2, 1]}])

            >>> Dataset.example(n=2)
            Dataset([{'a': [1, 1]}, {'b': [2, 2]}])
        """
        if n is None:
            return Dataset([{"a": [1, 2, 3, 4]}, {"b": [4, 3, 2, 1]}])
        else:
            return Dataset([{"a": [1] * n}, {"b": [2] * n}])

    @classmethod
    def from_edsl_object(cls, object):
        d = object.to_dict(add_edsl_version=False)
        return cls([{"key": list(d.keys())}, {"value": list(d.values())}])

    @classmethod
    def from_pandas_dataframe(cls, df):
        result = cls([{col: df[col].tolist()} for col in df.columns])
        return result

    def to_dict(self) -> dict:
        """
        Convert the dataset to a dictionary.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])
            >>> d.to_dict()
            {'data': [{'a': [1, 2, 3]}, {'b': [4, 5, 6]}]}

            >>> d = Dataset([{'x': ['a', 'b']}])
            >>> d.to_dict()
            {'data': [{'x': ['a', 'b']}]}
        """
        return {"data": self.data}

    @classmethod
    def from_dict(cls, data: dict) -> "Dataset":
        """
        Convert a dictionary to a dataset.

        Examples:
            >>> d = Dataset.from_dict({'data': [{'a': [1, 2, 3]}, {'b': [4, 5, 6]}]})
            >>> isinstance(d, Dataset)
            True
            >>> d.data
            [{'a': [1, 2, 3]}, {'b': [4, 5, 6]}]

            >>> d = Dataset.from_dict({'data': [{'x': ['a', 'b']}]})
            >>> d.data
            [{'x': ['a', 'b']}]
        """
        return cls(data["data"])

    def to_docx(self, output_file: str, title: str = None) -> None:
        """
        Convert the dataset to a Word document.

        Args:
            output_file (str): Path to save the Word document
            title (str, optional): Title for the document

        Examples:
            >>> import tempfile
            >>> d = Dataset([{'a': [1, 2, 3]}, {'b': [4, 5, 6]}])
            >>> with tempfile.NamedTemporaryFile(suffix='.docx') as tmp:
            ...     d.to_docx(tmp.name, title='Test Document')
            ...     import os
            ...     os.path.exists(tmp.name)
            True
        """
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create document
        doc = Document()

        # Add title if provided
        if title:
            title_heading = doc.add_heading(title, level=1)
            title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Get headers and data
        headers, data = self._tabular()

        # Create table
        table = doc.add_table(rows=len(data) + 1, cols=len(headers))
        table.style = "Table Grid"

        # Add headers
        for j, header in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = str(header)

        # Add data
        for i, row in enumerate(data):
            for j, cell_content in enumerate(row):
                cell = table.cell(i + 1, j)
                cell.text = str(cell_content) if cell_content is not None else ""

        # Adjust column widths
        for column in table.columns:
            max_width = 0
            for cell in column.cells:
                text_width = len(str(cell.text))
                max_width = max(max_width, text_width)
            for cell in column.cells:
                cell.width = Inches(min(max_width * 0.1 + 0.5, 6))

        # Save the document
        doc.save(output_file)

    def to_markdown(
        self,
        filename: Optional[str] = None,
        col_spacer_lines: int = 2,
        row_spacer_lines: int = 2,
        include_row_headers: bool = True,
        include_column_headers: bool = True,
        row_header_text: str = "Row",
        row_header_level: int = 2,
        column_header_level: int = 3,
    ):
        """
        Concatenate markdown cells into a single markdown file and return a FileStore.

        The concatenation proceeds left-to-right within each row, then top-to-bottom
        across rows. If a cell's content is wrapped in a ```markdown fenced block,
        the fence is removed before concatenation.

        Args:
            filename: Optional explicit path to write the markdown file. If not
                      provided, a temporary ``.md`` file is created.
            col_spacer_lines: Number of newline characters to insert between cells
                              within the same row (left-to-right). Defaults to 0.
            row_spacer_lines: Number of newline characters to insert between rows
                              (top-to-bottom). Defaults to 2.
            include_row_headers: If True, prepend each row with a header like
                                 "## Row 1". Defaults to True.
            include_column_headers: If True, prepend each cell with a header like
                                   "### column.name". Defaults to True.
            row_header_text: Base text used for the row header (e.g., "Row").
            row_header_level: Markdown header level for row headers (e.g., 2 for H2).
            column_header_level: Markdown header level for column headers (e.g., 3 for H3).

        Returns:
            FileStore: A FileStore pointing to the written markdown file.

        Examples:
            >>> from edsl.dataset import Dataset
            >>> content = "```markdown\\n# Title\\n\\nSome text.\\n```"
            >>> d = Dataset([{ 'md': [content] }])
            >>> fs = d.to_markdown()  # doctest: +ELLIPSIS
            >>> fs.suffix
            'md'
        """
        from ..scenarios import FileStore
        import tempfile

        def strip_markdown_fence(text: Any) -> str:
            if text is None:
                return ""
            s = str(text).strip()
            # Only remove fences when explicitly marked as ```markdown
            if s.startswith("```markdown"):
                # Remove opening line
                s = s[len("```markdown") :].lstrip("\n")
                # Remove the last closing fence if present
                if s.endswith("```"):
                    s = s[:-3]
                else:
                    fence_idx = s.rfind("```")
                    if fence_idx != -1:
                        s = s[:fence_idx]
                return s.strip("\n")
            return s

        # Build concatenated markdown content
        headers, rows = self._tabular()
        row_blocks: list[str] = []
        col_sep = "\n" * (col_spacer_lines if col_spacer_lines is not None else 0)
        row_sep = "\n" * (row_spacer_lines if row_spacer_lines is not None else 0)

        # Build row blocks with optional row/column headers
        for row_index, row in enumerate(rows, start=1):
            parts: list[str] = []
            if include_row_headers:
                parts.append(f"{'#' * row_header_level} {row_header_text} {row_index}")
            for col_index, cell in enumerate(row):
                cleaned = strip_markdown_fence(cell)
                if cleaned == "":
                    continue
                if include_column_headers:
                    parts.append(f"{'#' * column_header_level} {headers[col_index]}")
                parts.append(cleaned)

            row_block = col_sep.join(parts).strip()
            if row_block:
                row_blocks.append(row_block)

        content = row_sep.join(row_blocks)

        # Write to file
        if filename is None:
            with tempfile.NamedTemporaryFile(
                delete=False, suffix=".md", mode="w", encoding="utf-8"
            ) as f:
                f.write(content)
                path = f.name
        else:
            with open(filename, "w", encoding="utf-8") as f:
                f.write(content)
            path = filename

        return FileStore(path)

    def unique(self) -> "Dataset":
        """
        Remove duplicate rows from the dataset.

        Returns:
            A new Dataset with duplicate rows removed.

        Examples:
            >>> d = Dataset([{'a': [1, 2, 3, 1]}, {'b': [4, 5, 6, 4]}])
            >>> d.unique().data
            [{'a': [1, 2, 3]}, {'b': [4, 5, 6]}]

            >>> d = Dataset([{'x': ['a', 'b', 'a']}, {'y': [1, 2, 1]}])
            >>> d.unique().data
            [{'x': ['a', 'b']}, {'y': [1, 2]}]

            >>> # Dataset with a single column
            >>> Dataset([{'value': [1, 2, 3, 2, 1, 3]}]).unique().data
            [{'value': [1, 2, 3]}]
        """
        # Convert data to tuples for each row to make them hashable
        rows = []
        for i in range(len(self)):
            row = tuple(entry[list(entry.keys())[0]][i] for entry in self.data)
            rows.append(row)

        # Keep track of unique rows and their indices
        unique_rows = []
        indices = []

        # Use a set to track seen rows
        seen = set()
        for i, row in enumerate(rows):
            if row not in seen:
                seen.add(row)
                unique_rows.append(row)
                indices.append(i)

        # Create a new dataset with only the unique rows
        new_data = []
        for entry in self.data:
            key, values = list(entry.items())[0]
            new_values = [values[i] for i in indices]
            new_data.append({key: new_values})

        return Dataset(new_data)

    def expand(self, field: str, number_field: bool = False) -> "Dataset":
        """
        Expand a field containing lists into multiple rows.

        Args:
            field: The field containing lists to expand
            number_field: If True, adds a number field indicating the position in the original list

        Returns:
            A new Dataset with the expanded rows

        Examples:
            >>> d = Dataset([{'a': [[1, 2, 3], [4, 5, 6]]}, {'b': ['x', 'y']}])
            >>> d.expand('a').data
            [{'a': [1, 2, 3, 4, 5, 6]}, {'b': ['x', 'x', 'x', 'y', 'y', 'y']}]

            >>> d = Dataset([{'items': [['apple', 'banana'], ['orange']]}, {'id': [1, 2]}])
            >>> d.expand('items', number_field=True).data
            [{'items': ['apple', 'banana', 'orange']}, {'id': [1, 1, 2]}, {'items_number': [1, 2, 1]}]
        """
        from collections.abc import Iterable

        # Find the field in the dataset
        field_data = None
        for entry in self.data:
            key = list(entry.keys())[0]
            if key == field:
                field_data = entry[key]
                break

        if field_data is None:
            raise DatasetKeyError(
                f"Field '{field}' not found in dataset. Available fields are: {self.keys()}"
            )

        # Validate that the field contains lists
        if not all(isinstance(v, list) for v in field_data):
            # find the first entry that does not contain a list
            for i, entry in enumerate(self.data):
                key, values = list(entry.items())[0]
                if not isinstance(values, list):
                    raise DatasetTypeError(
                        f"Field '{field}' must contain lists in all entries. Entry {i} contains {type(values)}"
                    )

        # Create new expanded data structure
        new_data = []

        # Process each field
        for entry in self.data:
            key, values = list(entry.items())[0]
            new_values = []

            if key == field:
                # This is the field to expand - flatten all sublists
                for row_values in values:
                    if not isinstance(row_values, Iterable) or isinstance(
                        row_values, str
                    ):
                        row_values = [row_values]
                    new_values.extend(row_values)
            else:
                # For other fields, repeat each value the appropriate number of times
                for i, row_value in enumerate(values):
                    expand_length = len(field_data[i]) if i < len(field_data) else 0
                    new_values.extend([row_value] * expand_length)

            new_data.append({key: new_values})

        # Add number field if requested
        if number_field:
            number_values = []
            for i, lst in enumerate(field_data):
                number_values.extend(range(1, len(lst) + 1))
            new_data.append({f"{field}_number": number_values})

        return Dataset(new_data)


if __name__ == "__main__":
    import doctest

    doctest.testmod(optionflags=doctest.ELLIPSIS)
