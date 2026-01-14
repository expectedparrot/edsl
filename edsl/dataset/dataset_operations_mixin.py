"""
This module provides mixin classes that enable powerful data manipulation operations
across various EDSL list-like objects.

The DataOperationsBase class defines common operations for working with structured data,
including data transformation, visualization, export, querying, and analysis. These
operations are inherited by different specialized mixins (DatasetOperationsMixin,
ResultsOperationsMixin, etc.) which implement class-specific behaviors.

The design pattern used here allows different container types (Results, Dataset,
ScenarioList, AgentList) to share the same data manipulation interface, enabling
fluid operations across different parts of the EDSL ecosystem.
"""

import io
import warnings
import textwrap
import sqlite3
from typing import Optional, Tuple, Union, List, TYPE_CHECKING  # Callable not used

if TYPE_CHECKING:
    from ..scenarios import FileStore
from functools import wraps
from .exceptions import (
    DatasetKeyError,
    DatasetValueError,
    DatasetTypeError,
    DatasetExportError,
)

if TYPE_CHECKING:
    from docx import Document
    from .dataset import Dataset
    from ..scenarios import ScenarioList
    from ..jobs import Job  # noqa: F401


class DataOperationsBase:
    """
    Base class providing common data operations for EDSL container objects.

    This class serves as the foundation for various data manipulation mixins,
    providing a consistent interface for operations like filtering, aggregation,
    transformation, visualization, and export across different types of EDSL
    containers (Results, Dataset, ScenarioList, AgentList).

    Key functionality categories:

    1. Data Transformation:
       - Filtering with `filter()`
       - Creating new columns with `mutate()`
       - Reshaping with `long()`, `wide()`, `flatten()`, etc.
       - Selecting specific columns with `select()`

    2. Visualization and Display:
       - Tabular display with `table()`
       - Plotting with `ggplot2()`
       - Generating reports with `report()`

    3. Data Export:
       - To various formats with `to_csv()`, `to_excel()`, etc.
       - To other data structures with `to_pandas()`, `to_dicts()`, etc.

    4. Analysis:
       - SQL-based querying with `sql()`
       - Aggregation with `tally()`
       - Tree-based exploration

    These operations are designed to be applied fluently in sequence, enabling
    expressive data manipulation pipelines.
    """

    def relevant_columns(
        self, data_type: Optional[str] = None, remove_prefix: bool = False
    ) -> list:
        """Return the set of keys that are present in the dataset.

        :param data_type: The data type to filter by.
        :param remove_prefix: Whether to remove the prefix from the column names.

        >>> from ..dataset import Dataset
        >>> d = Dataset([{'a.b':[1,2,3,4]}])
        >>> d.relevant_columns()
        ['a.b']

        >>> d.relevant_columns(remove_prefix=True)
        ['b']

        >>> d = Dataset([{'a':[1,2,3,4]}, {'b':[5,6,7,8]}])
        >>> d.relevant_columns()
        ['a', 'b']

        >>> from edsl.results import Results; Results.example().select('how_feeling', 'how_feeling_yesterday').relevant_columns()
        ['answer.how_feeling', 'answer.how_feeling_yesterday']

        >>> from edsl.results import Results
        >>> sorted(Results.example().select().relevant_columns(data_type = "model"))
        ['model.canned_response', 'model.inference_service', 'model.model', 'model.model_index', 'model.temperature']

        >>> # Testing relevant_columns with invalid data_type raises DatasetValueError - tested in unit tests
        """
        columns = [list(x.keys())[0] for x in self]
        if remove_prefix:
            columns = [column.split(".")[-1] for column in columns]

        def get_data_type(column):
            if "." in column:
                return column.split(".")[0]
            else:
                return None

        if data_type:
            all_columns = columns[:]
            columns = [
                column for column in columns if get_data_type(column) == data_type
            ]
            if len(columns) == 0:
                all_data_types = sorted(
                    list(set(get_data_type(column) for column in all_columns))
                )
                raise DatasetValueError(
                    f"No columns found for data type: {data_type}. Available data types are: {all_data_types}."
                )

        return columns

    def num_observations(self):
        """Return the number of observations in the dataset.

        >>> from edsl.results import Results
        >>> Results.example().num_observations()
        4
        """
        _num_observations = None
        for entry in self:
            key, values = list(entry.items())[0]
            if _num_observations is None:
                _num_observations = len(values)
            else:
                if len(values) != _num_observations:
                    raise DatasetValueError(
                        f"The number of observations is not consistent across columns. "
                        f"Column '{key}' has {len(values)} observations, but previous columns had {_num_observations} observations."
                    )

        # Return 0 for empty datasets instead of None
        return _num_observations if _num_observations is not None else 0

    def vibe_plot(
        self,
        description: str,
        show_code: bool = True,
        show_expression: bool = False,
        height: float = 4,
        width: float = 6,
    ):
        """
        Generate and display a ggplot2 visualization using natural language description.

        Parameters:
            description: Natural language description of the desired plot
            show_code: If True, displays the generated R code alongside the plot
            show_expression: If True, prints the R code used (alias for show_code)
            height: Plot height in inches (default: 4)
            width: Plot width in inches (default: 6)

        Returns:
            A plot object that renders in Jupyter notebooks

        Examples:
            >>> from edsl.results import Results
            >>> r = Results.example()
            >>> # Generate a plot from a description (requires R/ggplot2):
            >>> # plot = r.vibe_plot("bar chart of how_feeling")
            >>> # Display with R code shown:
            >>> # plot = r.vibe_plot("bar chart of how_feeling", show_expression=True)
            >>> # Custom dimensions:
            >>> # plot = r.vibe_plot("scatter plot of age vs income", height=8, width=10)
        """
        from edsl_services.dataset_vibes.viz import GGPlotGenerator

        gen = GGPlotGenerator(model="gpt-4o", temperature=0.1)

        # Either show_code or show_expression will trigger displaying the code
        should_show = show_code or show_expression

        if should_show:
            # Get the code display object
            code_display = gen.make_plot_code(
                self.to_pandas(remove_prefix=True),
                description,
                return_display=True,
                show_code=True,
            )
            # Extract the actual code string for ggplot2
            r_code = code_display.code

            # Display the code (in Jupyter it will show with copy button, in terminal just the code)
            try:
                from IPython.display import display
                from ..utilities.utilities import is_notebook

                if is_notebook():
                    display(code_display)
                else:
                    print(r_code)
            except ImportError:
                # Not in a notebook environment
                print(r_code)
        else:
            # Get just the code string
            r_code = gen.make_plot_code(self.to_pandas(remove_prefix=True), description)

        return self.ggplot2(r_code, height=height, width=width)

    def vibe_sql(
        self,
        description: str,
        show_code: bool = True,
        show_expression: bool = False,
        transpose: bool = None,
        transpose_by: str = None,
        remove_prefix: bool = True,
        shape: str = "wide",
    ):
        """
        Generate and execute a SQL query using natural language description.

        Parameters:
            description: Natural language description of the desired query
            show_code: If True, displays the generated SQL query with copy button
            show_expression: If True, displays the generated SQL query (alias for show_code)
            transpose: Whether to transpose the resulting table (rows become columns)
            transpose_by: Column to use as the new index when transposing
            remove_prefix: Whether to remove type prefixes from column names
            shape: Data shape to use ("wide" or "long")

        Returns:
            A Dataset object containing the query results

        Examples:
            >>> from edsl.results import Results
            >>> r = Results.example()
            >>> # Generate and execute a query from a description:
            >>> # result = r.vibe_sql("Show all people over 30")
            >>> # With query shown:
            >>> # result = r.vibe_sql("Count by occupation", show_expression=True)
            >>> # Aggregation query:
            >>> # result = r.vibe_sql("Average age by city")
        """
        from edsl_services.dataset_vibes.sql import VibeSQLGenerator

        gen = VibeSQLGenerator(model="gpt-4o", temperature=0.1)

        # Either show_code or show_expression will trigger displaying the code
        should_show = show_code or show_expression

        if should_show:
            # Get the SQL query with display object
            query_display = gen.make_sql_query(
                self.to_pandas(remove_prefix=remove_prefix),
                description,
                return_display=True,
                show_code=True,
            )
            # Extract the actual SQL query string
            sql_query = query_display.code

            # Display the code (in Jupyter it will show with copy button, in terminal just the query)
            try:
                from IPython.display import display
                from ..utilities.utilities import is_notebook

                if is_notebook():
                    display(query_display)
                else:
                    print(sql_query)
            except ImportError:
                # Not in a notebook environment
                print("Generated SQL query:")
                print(sql_query)
        else:
            # Get just the SQL query string without display
            sql_query = gen.make_sql_query(
                self.to_pandas(remove_prefix=remove_prefix), description
            )

        # Execute the query and return the result
        return self.sql(
            sql_query,
            transpose=transpose,
            transpose_by=transpose_by,
            remove_prefix=remove_prefix,
            shape=shape,
        )

    def chart(self):
        """
        Create a chart from the results.
        """
        import altair as alt

        return alt.Chart(self.to_pandas(remove_prefix=True))

    def make_tabular(
        self, remove_prefix: bool, pretty_labels: Optional[dict] = None
    ) -> tuple[list, List[list]]:
        """Turn the results into a tabular format.

        :param remove_prefix: Whether to remove the prefix from the column names.

        >>> from edsl.results import Results
        >>> r = Results.example()
        >>> r.select('how_feeling').make_tabular(remove_prefix = True)
        (['how_feeling'], [['OK'], ['Great'], ['Terrible'], ['OK']])

        >>> r.select('how_feeling').make_tabular(remove_prefix = True, pretty_labels = {'how_feeling': "How are you feeling"})
        (['How are you feeling'], [['OK'], ['Great'], ['Terrible'], ['OK']])
        """

        def create_dict_from_list_of_dicts(list_of_dicts):
            for entry in list_of_dicts:
                key, list_of_values = list(entry.items())[0]
                yield key, list_of_values

        tabular_repr = dict(create_dict_from_list_of_dicts(self.data))

        full_header = [list(x.keys())[0] for x in self]

        rows = []
        for i in range(self.num_observations()):
            row = [tabular_repr[h][i] for h in full_header]
            rows.append(row)

        if remove_prefix:
            header = [h.split(".")[-1] for h in full_header]
        else:
            header = full_header

        if pretty_labels is not None:
            header = [pretty_labels.get(h, h) for h in header]

        return header, rows

    def print_long(self):
        """Print the results in a long format.
        >>> from edsl.results import Results
        >>> r = Results.example()
        >>> r.select('how_feeling').print_long()
        answer.how_feeling: OK
        answer.how_feeling: Great
        answer.how_feeling: Terrible
        answer.how_feeling: OK
        """
        for entry in self:
            key, list_of_values = list(entry.items())[0]
            for value in list_of_values:
                print(f"{key}: {value}")

    def get_tabular_data(
        self,
        remove_prefix: bool = False,
        pretty_labels: Optional[dict] = None,
    ) -> Tuple[List[str], List[List]]:
        """Internal method to get tabular data in a standard format.

        Args:
            remove_prefix: Whether to remove the prefix from column names
            pretty_labels: Dictionary mapping original column names to pretty labels

        Returns:
            Tuple containing (header_row, data_rows)
        """
        if pretty_labels is None:
            pretty_labels = {}

        return self.make_tabular(
            remove_prefix=remove_prefix, pretty_labels=pretty_labels
        )

    def to_jsonl(self, filename: Optional[str] = None):
        """Export the results to a FileStore instance containing JSONL data."""
        from .file_exports import JSONLExport

        exporter = JSONLExport(data=self, filename=filename)
        return exporter.export()

    def to_sqlite(
        self,
        filename: Optional[str] = None,
        remove_prefix: bool = False,
        pretty_labels: Optional[dict] = None,
        table_name: str = "results",
        if_exists: str = "replace",
    ):
        """Export the results to a SQLite database file."""
        from .file_exports import SQLiteExport

        exporter = SQLiteExport(
            data=self,
            filename=filename,
            remove_prefix=remove_prefix,
            pretty_labels=pretty_labels,
            table_name=table_name,
            if_exists=if_exists,
        )
        return exporter.export()

    def to_csv(
        self,
        filename: Optional[str] = None,
        remove_prefix: bool = False,
        pretty_labels: Optional[dict] = None,
    ) -> "FileStore":
        """Export the results to a FileStore instance containing CSV data."""
        from .file_exports import CSVExport

        exporter = CSVExport(
            data=self,
            filename=filename,
            remove_prefix=remove_prefix,
            pretty_labels=pretty_labels,
        )
        return exporter.export()

    def to_excel(
        self,
        filename: Optional[str] = None,
        remove_prefix: bool = False,
        pretty_labels: Optional[dict] = None,
        sheet_name: Optional[str] = None,
    ):
        """Export the results to a FileStore instance containing Excel data."""
        from .file_exports import ExcelExport

        exporter = ExcelExport(
            data=self,
            filename=filename,
            remove_prefix=remove_prefix,
            pretty_labels=pretty_labels,
            sheet_name=sheet_name,
        )
        return exporter.export()

    def to_docx(
        self,
        filename: Optional[str] = None,
        remove_prefix: bool = False,
        pretty_labels: Optional[dict] = None,
    ) -> "FileStore":
        """Export the results to a FileStore instance containing DOCX data.

        Each row of the dataset will be rendered on its own page, with a 2-column
        table that lists the keys and associated values for that observation.
        """
        # Import here to avoid heavy dependency unless the method is called
        from .file_exports import DocxExport

        exporter = DocxExport(
            data=self,
            filename=filename,
            remove_prefix=remove_prefix,
            pretty_labels=pretty_labels,
        )
        return exporter.export()

    def clipboard_data(self) -> str:
        """Return TSV representation of this object for clipboard operations.

        This method is called by the clipboard() method in the base class to provide
        a custom format for copying objects to the system clipboard.

        Returns:
            str: Tab-separated values representation of the object
        """
        # Use the to_csv method to get CSV data
        csv_filestore = self.to_csv()

        # Get the CSV content and convert it to TSV
        csv_content = csv_filestore.text

        # Convert CSV to TSV by replacing commas with tabs
        # This is a simple approach, but we should handle quoted fields properly
        import csv
        import io

        # Parse the CSV content
        csv_reader = csv.reader(io.StringIO(csv_content))
        rows = list(csv_reader)

        # Convert to TSV format
        tsv_lines = []
        for row in rows:
            tsv_lines.append("\t".join(row))

        return "\n".join(tsv_lines)

    def _db(self, remove_prefix: bool = True, shape: str = "wide"):
        """Create a SQLite database in memory and return the connection.

        Args:
            remove_prefix: Whether to remove the prefix from the column names
            shape: The shape of the data in the database ("wide" or "long")

        Returns:
            A database connection

        Examples:
            >>> from sqlalchemy import text
            >>> from edsl import Results
            >>> conn = Results.example()._db()
            >>> len(conn.execute("SELECT * FROM self").fetchall())
            4
            >>> conn = Results.example()._db(shape = "long")
            >>> len(conn.execute("SELECT * FROM self").fetchall())
            200
        """
        import sqlite3
        import csv

        conn = sqlite3.connect(":memory:")

        # Get CSV data from the dataset
        if remove_prefix and shape == "wide":
            csv_string = self.remove_prefix().to_csv().text
        else:
            csv_string = self.to_csv().text

        # Parse CSV
        reader = csv.DictReader(io.StringIO(csv_string))
        rows = list(reader)

        if not rows:
            # Create empty table
            conn.execute("CREATE TABLE self (empty TEXT)")
            return conn

        columns = list(rows[0].keys())

        if shape == "long":
            # Melt the data to long format
            long_rows = []
            row_number = 1
            for row in rows:
                for key, value in row.items():
                    data_type = key.split(".")[0] if "." in key else None
                    key_name = ".".join(key.split(".")[1:]) if "." in key else key
                    long_rows.append(
                        {
                            "row_number": row_number,
                            "key": key_name,
                            "value": value,
                            "data_type": data_type,
                        }
                    )
                    row_number += 1

            # Create long format table
            conn.execute(
                """
                CREATE TABLE self (
                    row_number INTEGER,
                    key TEXT,
                    value TEXT,
                    data_type TEXT
                )
            """
            )
            conn.executemany(
                "INSERT INTO self VALUES (?, ?, ?, ?)",
                [
                    (r["row_number"], r["key"], r["value"], r["data_type"])
                    for r in long_rows
                ],
            )
        else:
            # Create wide format table
            # Sanitize column names for SQL
            safe_columns = [f'"{col}"' for col in columns]
            create_sql = f"CREATE TABLE self ({', '.join(f'{col} TEXT' for col in safe_columns)})"
            conn.execute(create_sql)

            placeholders = ", ".join(["?" for _ in columns])
            insert_sql = f"INSERT INTO self VALUES ({placeholders})"
            conn.executemany(insert_sql, [tuple(row.values()) for row in rows])

        conn.commit()
        return conn

    def sql(
        self,
        query: str,
        transpose: bool = None,
        transpose_by: str = None,
        remove_prefix: bool = True,
        shape: str = "wide",
    ) -> "Dataset":
        """
        Execute SQL queries on the dataset.

        This powerful method allows you to use SQL to query and transform your data,
        combining the expressiveness of SQL with EDSL's data structures. It works by
        creating an in-memory SQLite database from your data and executing the query
        against it.

        Parameters:
            query: SQL query string to execute
            transpose: Whether to transpose the resulting table (rows become columns)
            transpose_by: Column to use as the new index when transposing
            remove_prefix: Whether to remove type prefixes (e.g., "answer.") from column names
            shape: Data shape to use ("wide" or "long")
                  - "wide": Default tabular format with columns for each field
                  - "long": Melted format with key-value pairs, useful for certain queries

        Returns:
            A Dataset object containing the query results

        Notes:
            - The data is stored in a table named "self" in the SQLite database
            - In wide format, column names include their type prefix unless remove_prefix=True
            - In long format, the data is melted into columns: row_number, key, value, data_type
            - Complex objects like lists and dictionaries are converted to strings

        Examples:
            >>> from edsl import Results
            >>> r = Results.example()

            # Basic selection
            >>> len(r.sql("SELECT * FROM self", shape="wide"))
            4

            # Filtering with WHERE clause
            >>> r.sql("SELECT * FROM self WHERE how_feeling = 'Great'").num_observations()
            1

            # Aggregation
            >>> r.sql("SELECT how_feeling, COUNT(*) as count FROM self GROUP BY how_feeling").keys()
            ['how_feeling', 'count']

            # Using long format
            >>> len(r.sql("SELECT * FROM self", shape="long"))
            200
        """
        from .dataset import Dataset

        conn = self._db(remove_prefix=remove_prefix, shape=shape)
        try:
            cursor = conn.execute(query)
        except sqlite3.OperationalError as e:
            raise DatasetValueError(f"Error executing SQL query: {e}; query: {query}")
            from traceback import format_exc
            raise DatasetValueError(f"Error executing SQL query: {e}; query: {query}; traceback: {format_exc()}")

        columns = [description[0] for description in cursor.description]
        rows = cursor.fetchall()

        if not rows:
            return Dataset([])

        # Convert row-oriented data to column-oriented format for Dataset
        # Dataset format: [{'col1': [val1, val2, ...]}, {'col2': [val1, val2, ...]}]
        column_data = {col: [] for col in columns}
        for row in rows:
            for col, val in zip(columns, row):
                column_data[col].append(val)

        # Handle transpose
        if transpose or transpose_by:
            # Determine index column
            index_col = transpose_by if transpose_by else columns[0]

            # Create transposed data - rows become columns
            new_columns = column_data[index_col]  # These become the new column names
            other_cols = [c for c in columns if c != index_col]

            # Build transposed column data
            transposed_data = []
            # First add the index column (original column names become values)
            transposed_data.append({"index": other_cols})
            # Then add each row as a column
            for i, new_col_name in enumerate(new_columns):
                col_values = [column_data[c][i] for c in other_cols]
                transposed_data.append({str(new_col_name): col_values})

            return Dataset(transposed_data)

        # Create Dataset in column-oriented format
        dataset_data = [{col: values} for col, values in column_data.items()]
        return Dataset(dataset_data)

    def to_pandas(self, remove_prefix: bool = False, lists_as_strings=False):
        """Convert the results to a pandas DataFrame, ensuring that lists remain as lists.

        Args:
            remove_prefix: Whether to remove the prefix from the column names.
            lists_as_strings: Whether to convert lists to strings.

        Returns:
            A pandas DataFrame.
        """
        # pandas is imported in _to_pandas_strings
        return self._to_pandas_strings(remove_prefix)

    def _to_pandas_strings(self, remove_prefix: bool = False):
        """Convert the results to a pandas DataFrame.

        Args:
            remove_prefix: Whether to remove the prefix from the column names.

        Examples:
            >>> from edsl.results import Results
            >>> r = Results.example()
            >>> r.select('how_feeling').to_pandas()  # doctest: +SKIP
              answer.how_feeling
            0                 OK
            1              Great
            2           Terrible
            3                 OK
        """
        try:
            import pandas as pd
        except ImportError:
            raise ImportError(
                "pandas is required for to_pandas(). "
                "Install with: pip install edsl[pandas] or pip install pandas"
            )

        # Handle empty dataset case
        if not self.data:
            return pd.DataFrame()

        csv_string = self.to_csv(remove_prefix=remove_prefix).text
        csv_buffer = io.StringIO(csv_string)
        df = pd.read_csv(csv_buffer)
        # df_sorted = df.sort_index(axis=1)  # Sort columns alphabetically
        return df

    def to_polars(self, remove_prefix: bool = False, lists_as_strings=False):
        """Convert the results to a Polars DataFrame.

        Args:
            remove_prefix: Whether to remove the prefix from the column names.
            lists_as_strings: Whether to convert lists to strings.

        Returns:
            A Polars DataFrame.
        """
        # polars is imported in _to_polars_strings
        return self._to_polars_strings(remove_prefix)

    def _to_polars_strings(self, remove_prefix: bool = False):
        """Convert the results to a Polars DataFrame.

        Args:
            remove_prefix: Whether to remove the prefix from the column names.

        Returns:
            A Polars DataFrame.
        """
        try:
            import polars as pl
        except ImportError:
            raise ImportError(
                "polars is required for to_polars(). "
                "Install with: pip install polars"
            )

        csv_string = self.to_csv(remove_prefix=remove_prefix).text
        df = pl.read_csv(io.StringIO(csv_string))
        return df

    def tree(self, node_order: Optional[List[str]] = None):
        """Convert the results to a Tree.

        Args:
            node_order: The order of the nodes.

        Returns:
            A Tree object.
        """
        from .dataset_tree import Tree

        return Tree(self, node_order=node_order)

    def to_scenario_list(self, remove_prefix: bool = True) -> "ScenarioList":
        """Convert the results to a list of dictionaries, one per scenario.

        :param remove_prefix: Whether to remove the prefix from the column names.

        >>> from edsl.results import Results
        >>> r = Results.example()
        >>> r.select('how_feeling').to_scenario_list()
        ScenarioList([Scenario({'how_feeling': 'OK'}), Scenario({'how_feeling': 'Great'}), Scenario({'how_feeling': 'Terrible'}), Scenario({'how_feeling': 'OK'})])
        """
        from ..scenarios import ScenarioList, Scenario

        list_of_dicts = self.to_dicts(remove_prefix=remove_prefix)
        scenarios = []
        for d in list_of_dicts:
            scenarios.append(Scenario(d))
        return ScenarioList(scenarios)

    def to_agent_list(self, remove_prefix: bool = True):
        """Convert the results to a list of dictionaries, one per agent.

        :param remove_prefix: Whether to remove the prefix from the column names.

        >>> from edsl.results import Results
        >>> r = Results.example()
        >>> r.select('how_feeling').to_agent_list()
        AgentList([Agent(traits = {'how_feeling': 'OK'}), Agent(traits = {'how_feeling': 'Great'}), Agent(traits = {'how_feeling': 'Terrible'}), Agent(traits = {'how_feeling': 'OK'})])
        """
        from ..agents import Agent, AgentList

        list_of_dicts = self.to_dicts(remove_prefix=remove_prefix)
        agents = []
        for d in list_of_dicts:
            if "name" in d:
                d["agent_name"] = d.pop("name")
                agents.append(Agent(d, name=d["agent_name"]))
            elif "agent_parameters" in d:
                agent_parameters = d.pop("agent_parameters")
                agent_name = agent_parameters.get("name", None)
                instruction = agent_parameters.get("instruction", None)
                agents.append(Agent(d, name=agent_name, instruction=instruction))
            else:
                agents.append(Agent(d))
        return AgentList(agents)

    def to_dicts(self, remove_prefix: bool = True) -> list[dict]:
        """Convert the results to a list of dictionaries.

        :param remove_prefix: Whether to remove the prefix from the column names.

        >>> from edsl.results import Results
        >>> r = Results.example()
        >>> r.select('how_feeling').to_dicts()
        [{'how_feeling': 'OK'}, {'how_feeling': 'Great'}, {'how_feeling': 'Terrible'}, {'how_feeling': 'OK'}]

        """
        list_of_keys = []
        list_of_values = []
        for entry in self:
            key, values = list(entry.items())[0]
            list_of_keys.append(key)
            list_of_values.append(values)

        if remove_prefix:
            list_of_keys = [key.split(".")[-1] for key in list_of_keys]

        list_of_dicts = []
        for entries in zip(*list_of_values):
            list_of_dicts.append(dict(zip(list_of_keys, entries)))

        return list_of_dicts

    def to_list(self, flatten=False, remove_none=False, unzipped=False) -> list[list]:
        """Convert the results to a list of lists.

        :param flatten: Whether to flatten the list of lists.
        :param remove_none: Whether to remove None values from the list.

        >>> from edsl.results import Results
        >>> Results.example().select('how_feeling', 'how_feeling_yesterday')
        Dataset([{'answer.how_feeling': ['OK', 'Great', 'Terrible', 'OK']}, {'answer.how_feeling_yesterday': ['Great', 'Good', 'OK', 'Terrible']}])

        >>> Results.example().select('how_feeling', 'how_feeling_yesterday').to_list()
        [('OK', 'Great'), ('Great', 'Good'), ('Terrible', 'OK'), ('OK', 'Terrible')]

        >>> r = Results.example()
        >>> r.select('how_feeling').to_list()
        ['OK', 'Great', 'Terrible', 'OK']

        >>> from edsl.dataset import Dataset
        >>> Dataset([{'a.b': [[1, 9], 2, 3, 4]}]).select('a.b').to_list(flatten = True)
        [1, 9, 2, 3, 4]

        >>> from edsl.dataset import Dataset
        >>> # Testing to_list flatten with multiple columns raises DatasetValueError - tested in unit tests


        """
        if len(self.relevant_columns()) > 1 and flatten:
            raise DatasetValueError(
                "Cannot flatten a list of lists when there are multiple columns selected."
            )

        if len(self.relevant_columns()) == 1:
            # if only one 'column' is selected (which is typical for this method
            list_to_return = list(self[0].values())[0]
        else:
            keys = self.relevant_columns()
            data = self.to_dicts(remove_prefix=False)
            list_to_return = []
            for d in data:
                list_to_return.append(tuple([d[key] for key in keys]))

        if remove_none:
            list_to_return = [item for item in list_to_return if item is not None]

        if flatten:
            new_list = []
            for item in list_to_return:
                if isinstance(item, list):
                    new_list.extend(item)
                else:
                    new_list.append(item)
            list_to_return = new_list

        # return PrettyList(list_to_return)
        return list_to_return

    def html(
        self,
        filename: Optional[str] = None,
        cta: str = "Open in browser",
        return_link: bool = False,
    ):
        import os
        import tempfile
        from ..utilities.utilities import is_notebook

        df = self.to_pandas()

        if filename is None:
            current_directory = os.getcwd()
            filename = tempfile.NamedTemporaryFile(
                "w", delete=False, suffix=".html", dir=current_directory
            ).name

        with open(filename, "w") as f:
            f.write(df.to_html())

        if is_notebook():
            try:
                from IPython.display import HTML, display

                html_url = f"/files/{filename}"
                html_link = f'<a href="{html_url}" target="_blank">{cta}</a>'
                display(HTML(html_link))
            except ImportError:
                print(f"Saved to {filename}")
        else:
            print(f"Saved to {filename}")
            import webbrowser
            import os

            webbrowser.open(f"file://{os.path.abspath(filename)}")

        if return_link:
            return filename

    def _prepare_report_data(
        self,
        *fields: Optional[str],
        top_n: Optional[int] = None,
        header_fields: Optional[List[str]] = None,
    ) -> tuple:
        """Prepares data for report generation in various formats.

        Args:
            *fields: The fields to include in the report. If none provided, all fields are used.
            top_n: Optional limit on the number of observations to include.
            header_fields: Optional list of fields to include in the main header instead of as sections.

        Returns:
            A tuple containing (field_data, num_obs, fields, header_fields)
        """
        # If no fields specified, use all columns
        if not fields:
            fields = self.relevant_columns()

        # Initialize header_fields if not provided
        if header_fields is None:
            header_fields = []

        # Validate all fields
        all_fields = list(fields) + [f for f in header_fields if f not in fields]
        for field in all_fields:
            if field not in self.relevant_columns():
                raise DatasetKeyError(f"Field '{field}' not found in dataset")

        # Get data for each field
        field_data = {}
        for field in all_fields:
            for entry in self:
                if field in entry:
                    field_data[field] = entry[field]
                    break

        # Number of observations to process
        num_obs = self.num_observations()
        if top_n is not None:
            num_obs = min(num_obs, top_n)

        return field_data, num_obs, fields, header_fields

    def _report_markdown(
        self, field_data, num_obs, fields, header_fields, divider: bool = True
    ) -> str:
        """Generates a markdown report from the prepared data.

        Args:
            field_data: Dictionary mapping field names to their values
            num_obs: Number of observations to include
            fields: Fields to include as sections
            header_fields: Fields to include in the observation header
            divider: If True, adds a horizontal rule between observations

        Returns:
            A string containing the markdown report
        """
        report_lines = []
        for i in range(num_obs):
            # Create header with observation number and any header fields
            header = f"# Observation: {i+1}"
            if header_fields:
                header_parts = []
                for field in header_fields:
                    value = field_data[field][i]
                    # Get the field name without prefix for cleaner display
                    display_name = field.split(".")[-1] if "." in field else field
                    # Format with backticks for monospace
                    header_parts.append(f"`{display_name}`: {value}")
                if header_parts:
                    header += f" ({', '.join(header_parts)})"
            report_lines.append(header)

            # Add the remaining fields
            for field in fields:
                if field not in header_fields:
                    report_lines.append(f"## {field}")
                    value = field_data[field][i]
                    if isinstance(value, list) or isinstance(value, dict):
                        import json

                        report_lines.append(f"```\n{json.dumps(value, indent=2)}\n```")
                    else:
                        report_lines.append(str(value))

            # Add divider between observations if requested
            if divider and i < num_obs - 1:
                report_lines.append("\n---\n")
            else:
                report_lines.append("")  # Empty line between observations

        return "\n".join(report_lines)

    def _report_docx(self, field_data, num_obs, fields, header_fields) -> "Document":
        """Generates a Word document report from the prepared data.

        Args:
            field_data: Dictionary mapping field names to their values
            num_obs: Number of observations to include
            fields: Fields to include as sections
            header_fields: Fields to include in the observation header

        Returns:
            A docx.Document object containing the report
        """
        try:
            from docx import Document
            from docx.shared import Pt
            import json
        except ImportError:
            from .exceptions import DatasetImportError

            raise DatasetImportError(
                "The python-docx package is required for DOCX export. Install it with 'pip install python-docx'."
            )

        doc = Document()

        for i in range(num_obs):
            # Create header with observation number and any header fields
            header_text = f"Observation: {i+1}"
            if header_fields:
                header_parts = []
                for field in header_fields:
                    value = field_data[field][i]
                    # Get the field name without prefix for cleaner display
                    display_name = field.split(".")[-1] if "." in field else field
                    header_parts.append(f"{display_name}: {value}")
                if header_parts:
                    header_text += f" ({', '.join(header_parts)})"

            doc.add_heading(header_text, level=1)

            # Add the remaining fields
            for field in fields:
                if field not in header_fields:
                    doc.add_heading(field, level=2)
                    value = field_data[field][i]

                    if isinstance(value, (list, dict)):
                        # Format structured data with indentation
                        formatted_value = json.dumps(value, indent=2)
                        p = doc.add_paragraph()
                        p.add_run(formatted_value).font.name = "Courier New"
                        p.add_run().font.size = Pt(10)
                    else:
                        doc.add_paragraph(str(value))

            # Add page break between observations except for the last one
            if i < num_obs - 1:
                doc.add_page_break()

        return doc

    def report(
        self,
        *fields: Optional[str],
        top_n: Optional[int] = None,
        header_fields: Optional[List[str]] = None,
        divider: bool = True,
        return_string: bool = False,
        format: str = "markdown",
        filename: Optional[str] = None,
    ) -> Optional[Union[str, "Document"]]:
        """Generates a report of the results by iterating through rows.

        Args:
            *fields: The fields to include in the report. If none provided, all fields are used.
            top_n: Optional limit on the number of observations to include.
            header_fields: Optional list of fields to include in the main header instead of as sections.
            divider: If True, adds a horizontal rule between observations (markdown only).
            return_string: If True, returns the markdown string. If False (default in notebooks),
                          only displays the markdown without returning.
            format: Output format - either "markdown" or "docx".
            filename: If provided and format is "docx", saves the document to this file.

        Returns:
            Depending on format and return_string:
            - For markdown: A string if return_string is True, otherwise None (displays in notebook)
            - For docx: A docx.Document object, or None if filename is provided (saves to file)

        Examples:
            >>> from edsl.results import Results
            >>> r = Results.example()
            >>> report = r.select('how_feeling').report(return_string=True)
            >>> "# Observation: 1" in report
            True
            >>> doc = r.select('how_feeling').report(format="docx")
            >>> isinstance(doc, object)
            True
        """
        from ..utilities.utilities import is_notebook

        # Prepare the data for the report
        field_data, num_obs, fields, header_fields = self._prepare_report_data(
            *fields, top_n=top_n, header_fields=header_fields
        )

        # Generate the report in the requested format
        if format.lower() == "markdown":
            report_text = self._report_markdown(
                field_data, num_obs, fields, header_fields, divider
            )

            # In notebooks, display as markdown
            is_nb = is_notebook()
            if is_nb and not return_string:
                from IPython.display import Markdown, display

                display(Markdown(report_text))
                return None

            # Return the string if requested or if not in a notebook
            return report_text

        elif format.lower() == "docx":
            doc = self._report_docx(field_data, num_obs, fields, header_fields)

            # Save to file if filename is provided
            if filename:
                doc.save(filename)
                print(f"Report saved to {filename}")
                return None

            return doc

        else:
            raise DatasetExportError(
                f"Unsupported format: {format}. Use 'markdown' or 'docx'."
            )

    def confusion_matrix(
        self, actual_field: str, predicted_field: str, normalize: Optional[str] = None
    ) -> "Dataset":
        """
        Create a confusion matrix comparing two categorical fields.

        A confusion matrix is a table that shows the frequency of actual vs predicted values,
        useful for evaluating classification results. Rows typically represent actual values
        and columns represent predicted values.

        Parameters:
            actual_field: The field containing actual/true values
            predicted_field: The field containing predicted values
            normalize: Optional normalization mode:
                - None (default): Show raw counts
                - "true": Normalize over true (row) conditions
                - "pred": Normalize over predicted (column) conditions
                - "all": Normalize over all observations

        Returns:
            A Dataset object with the confusion matrix where:
            - First column contains the unique values from actual_field
            - Additional columns contain counts for each unique value in predicted_field
            - Column names are the unique values from predicted_field
            - Optionally normalized based on the normalize parameter

        Notes:
            - Field names can be specified with or without prefixes (e.g., 'field' or 'answer.field')
            - Missing values (None) are included as a separate category
            - The resulting Dataset can be displayed as a table using .table()
            - Results are sorted by the actual_field values

        Examples:
            >>> from edsl.dataset import Dataset
            >>> d = Dataset([
            ...     {'actual': ['cat', 'cat', 'dog', 'dog', 'dog']},
            ...     {'predicted': ['cat', 'dog', 'dog', 'dog', 'cat']}
            ... ])
            >>> cm = d.confusion_matrix('actual', 'predicted')
            >>> sorted(cm.keys())
            ['actual', 'cat', 'dog']

            >>> # With normalization over rows (true conditions)
            >>> cm_norm = d.confusion_matrix('actual', 'predicted', normalize='true')
            >>> sorted(cm_norm.keys())
            ['actual', 'cat', 'dog']
        """
        # Validate normalize parameter
        if normalize not in [None, "true", "pred", "all"]:
            raise DatasetValueError(
                f"normalize must be None, 'true', 'pred', or 'all', got '{normalize}'"
            )

        # Get the values for both fields
        actual_values = self._key_to_value(actual_field)
        predicted_values = self._key_to_value(predicted_field)

        if len(actual_values) != len(predicted_values):
            raise DatasetValueError(
                f"Fields must have the same length. "
                f"{actual_field} has {len(actual_values)} values, "
                f"{predicted_field} has {len(predicted_values)} values."
            )

        # Get unique values for both fields (sorted for consistency)
        unique_actual = sorted(set(actual_values), key=lambda x: (x is None, x))
        unique_predicted = sorted(set(predicted_values), key=lambda x: (x is None, x))

        # Build the confusion matrix as a dictionary of counts
        matrix = {}
        for actual_val in unique_actual:
            matrix[actual_val] = {}
            for pred_val in unique_predicted:
                matrix[actual_val][pred_val] = 0

        # Count occurrences
        for actual_val, pred_val in zip(actual_values, predicted_values):
            matrix[actual_val][pred_val] += 1

        # Apply normalization if requested
        if normalize == "true":
            # Normalize by row (actual values)
            for actual_val in unique_actual:
                row_sum = sum(matrix[actual_val].values())
                if row_sum > 0:
                    for pred_val in unique_predicted:
                        matrix[actual_val][pred_val] /= row_sum
        elif normalize == "pred":
            # Normalize by column (predicted values)
            col_sums = {pred_val: 0 for pred_val in unique_predicted}
            for actual_val in unique_actual:
                for pred_val in unique_predicted:
                    col_sums[pred_val] += matrix[actual_val][pred_val]
            for actual_val in unique_actual:
                for pred_val in unique_predicted:
                    if col_sums[pred_val] > 0:
                        matrix[actual_val][pred_val] /= col_sums[pred_val]
        elif normalize == "all":
            # Normalize by total count
            total = sum(
                matrix[actual_val][pred_val]
                for actual_val in unique_actual
                for pred_val in unique_predicted
            )
            if total > 0:
                for actual_val in unique_actual:
                    for pred_val in unique_predicted:
                        matrix[actual_val][pred_val] /= total

        # Convert to Dataset format
        # First column is the actual values
        dataset_data = [{actual_field: unique_actual}]

        # Add a column for each predicted value
        for pred_val in unique_predicted:
            col_name = str(pred_val) if pred_val is not None else "None"
            col_values = [matrix[actual_val][pred_val] for actual_val in unique_actual]
            dataset_data.append({col_name: col_values})

        from ..dataset import Dataset

        result = Dataset(dataset_data)
        # Mark this as a confusion matrix for validation in perplexity()
        result._is_confusion_matrix = True
        result._confusion_matrix_metadata = {
            "actual_field": actual_field,
            "predicted_values": unique_predicted,
            "actual_values": unique_actual,
            "normalize": normalize,
        }
        return result

    def perplexity(self, per_class: bool = False) -> Union[float, "Dataset"]:
        """
        Calculate perplexity from a confusion matrix.

        Perplexity measures the uncertainty or "confusion" of a classifier based on the
        cross-entropy of predictions. It's calculated as exp(average negative log-likelihood)
        where the log-likelihood is computed from the probability assigned to the correct
        class. Lower perplexity indicates better model performance.

        This method can only be called on a Dataset that was created by the confusion_matrix()
        method. For each actual class, we look at the probability assigned to predicting that
        correct class (the diagonal of the confusion matrix). The overall perplexity is the
        weighted average across all classes.

        Args:
            per_class: If True, returns a Dataset with per-class perplexity values.
                      If False (default), returns the overall weighted perplexity as a float.

        Returns:
            If per_class=False: A float representing the weighted average perplexity
            If per_class=True: A Dataset with perplexity for each actual class

        Raises:
            DatasetNotConfusionMatrixError: If called on a Dataset that is not a confusion matrix

        Notes:
            - Perplexity of 1.0 means perfect prediction (100% accuracy)
            - Perplexity of N means random guessing (uniform distribution over N classes)
            - Uses natural logarithm (ln) for calculation: perplexity = exp(-ln(p_correct))
            - For raw count matrices, rows are normalized to probabilities
            - Weighted average accounts for class imbalance in the data

        Examples:
            >>> from edsl.dataset import Dataset
            >>> d = Dataset([
            ...     {'actual': ['cat', 'cat', 'dog', 'dog', 'dog']},
            ...     {'predicted': ['cat', 'dog', 'dog', 'dog', 'cat']}
            ... ])
            >>> cm = d.confusion_matrix('actual', 'predicted')
            >>> perp = cm.perplexity()
            >>> isinstance(perp, float)
            True
            >>> perp >= 1.0
            True

            >>> # Get per-class perplexity
            >>> per_class_perp = cm.perplexity(per_class=True)
            >>> 'actual' in per_class_perp.keys()
            True
            >>> 'perplexity' in per_class_perp.keys()
            True
        """
        import math
        from .exceptions import DatasetNotConfusionMatrixError

        # Check if this is a confusion matrix
        if not hasattr(self, "_is_confusion_matrix") or not self._is_confusion_matrix:
            raise DatasetNotConfusionMatrixError(
                "perplexity() can only be called on a Dataset created by confusion_matrix(). "
                "Create a confusion matrix first using: dataset.confusion_matrix(actual_field, predicted_field)"
            )

        # Get metadata
        metadata = self._confusion_matrix_metadata
        actual_field = metadata["actual_field"]
        actual_values = metadata["actual_values"]

        # Get the data as a pandas DataFrame for easier manipulation
        df = self.to_pandas()

        # Calculate perplexity for each actual class (row)
        # We need both the per-class perplexity and the sample counts for weighting
        perplexities = []
        sample_counts = []

        for i, actual_val in enumerate(actual_values):
            # Get the row for this actual value
            row = df[df[actual_field] == actual_val]

            if len(row) == 0:
                continue

            # Extract probability values (all columns except the actual_field column)
            prob_cols = [col for col in df.columns if col != actual_field]
            probs = row[prob_cols].values[0]

            # Normalize if needed (convert counts to probabilities)
            row_sum = sum(probs)
            if row_sum > 0:
                probs = [p / row_sum for p in probs]
            else:
                # Handle edge case of all zeros
                probs = [1.0 / len(probs) for _ in probs]

            # Get the probability of the correct prediction (diagonal element)
            # The correct prediction is at the same index as the actual value
            p_correct = probs[i] if i < len(probs) else 0.0

            # Avoid log(0) by using a small epsilon
            if p_correct <= 0:
                p_correct = 1e-10

            # Calculate perplexity for this class: exp(-ln(p_correct))
            # This is equivalent to: perplexity = 1 / p_correct, but we use exp(-log(p))
            # for numerical stability and to match the standard definition
            perplexity_val = math.exp(-math.log(p_correct))

            perplexities.append(perplexity_val)
            sample_counts.append(row_sum)

        if per_class:
            # Return a Dataset with per-class perplexity
            from ..dataset import Dataset

            return Dataset(
                [{actual_field: list(actual_values)}, {"perplexity": perplexities}]
            )
        else:
            # Return weighted average perplexity
            # Weight by the number of samples in each class
            total_samples = sum(sample_counts)
            if total_samples == 0:
                return 0.0

            # Calculate weighted average of log probabilities, then exp()
            weighted_log_sum = sum(
                count * math.log(1.0 / perp)
                for count, perp in zip(sample_counts, perplexities)
            )
            average_log_prob = weighted_log_sum / total_samples

            return math.exp(-average_log_prob)

    def true_positive_count(self, per_class: bool = False) -> Union[int, "Dataset"]:
        """
        Get the count of true positives from a confusion matrix.

        True positives are the correctly predicted instances for each class,
        represented by the diagonal elements of the confusion matrix.

        This method can only be called on a Dataset that was created by the confusion_matrix()
        method.

        Args:
            per_class: If True, returns a Dataset with true positive counts for each class.
                      If False (default), returns the total count of true positives as an int.

        Returns:
            If per_class=False: An integer representing the total count of correct predictions
            If per_class=True: A Dataset with true positive counts for each actual class

        Raises:
            DatasetNotConfusionMatrixError: If called on a Dataset that is not a confusion matrix

        Notes:
            - True positives are the diagonal elements of the confusion matrix
            - Total true positives equals the number of correct predictions
            - For multi-class problems, this sums all diagonal elements

        Examples:
            >>> from edsl.dataset import Dataset
            >>> d = Dataset([
            ...     {'actual': ['cat', 'cat', 'dog', 'dog', 'dog']},
            ...     {'predicted': ['cat', 'dog', 'dog', 'dog', 'cat']}
            ... ])
            >>> cm = d.confusion_matrix('actual', 'predicted')
            >>> tp = cm.true_positive_count()
            >>> isinstance(tp, int)
            True
            >>> tp
            3

            >>> # Get per-class counts
            >>> per_class_tp = cm.true_positive_count(per_class=True)
            >>> 'actual' in per_class_tp.keys()
            True
            >>> 'true_positive_count' in per_class_tp.keys()
            True
        """
        from .exceptions import DatasetNotConfusionMatrixError

        # Check if this is a confusion matrix
        if not hasattr(self, "_is_confusion_matrix") or not self._is_confusion_matrix:
            raise DatasetNotConfusionMatrixError(
                "true_positive_count() can only be called on a Dataset created by confusion_matrix(). "
                "Create a confusion matrix first using: dataset.confusion_matrix(actual_field, predicted_field)"
            )

        # Get metadata
        metadata = self._confusion_matrix_metadata
        actual_field = metadata["actual_field"]
        actual_values = metadata["actual_values"]

        # Get the data as a pandas DataFrame
        df = self.to_pandas()

        # Extract true positive counts (diagonal elements)
        tp_counts = []
        for i, actual_val in enumerate(actual_values):
            row = df[df[actual_field] == actual_val]
            if len(row) == 0:
                tp_counts.append(0)
                continue

            # Get the column name for this class
            # Columns are ordered as they appear in predicted_values
            prob_cols = [col for col in df.columns if col != actual_field]
            if i < len(prob_cols):
                col_name = prob_cols[i]
                tp_count = int(row[col_name].values[0])
                tp_counts.append(tp_count)
            else:
                tp_counts.append(0)

        if per_class:
            # Return a Dataset with per-class counts
            from ..dataset import Dataset

            return Dataset(
                [
                    {actual_field: list(actual_values)},
                    {"true_positive_count": tp_counts},
                ]
            )
        else:
            # Return total count
            return sum(tp_counts)

    def false_negative_count(self, per_class: bool = False) -> Union[int, "Dataset"]:
        """
        Get the count of false negatives from a confusion matrix.

        False negatives are instances of a class that were incorrectly predicted as
        a different class. For each actual class, this is the sum of all off-diagonal
        elements in that row.

        This method can only be called on a Dataset that was created by the confusion_matrix()
        method.

        Args:
            per_class: If True, returns a Dataset with false negative counts for each class.
                      If False (default), returns the total count of false negatives as an int.

        Returns:
            If per_class=False: An integer representing the total count of false negatives
            If per_class=True: A Dataset with false negative counts for each actual class

        Raises:
            DatasetNotConfusionMatrixError: If called on a Dataset that is not a confusion matrix

        Notes:
            - False negatives = instances of a class predicted as something else
            - For each class: FN = row_sum - true_positives
            - High false negatives indicate the model often misses that class

        Examples:
            >>> from edsl.dataset import Dataset
            >>> d = Dataset([
            ...     {'actual': ['cat', 'cat', 'dog', 'dog', 'dog']},
            ...     {'predicted': ['cat', 'dog', 'dog', 'dog', 'cat']}
            ... ])
            >>> cm = d.confusion_matrix('actual', 'predicted')
            >>> fn = cm.false_negative_count()
            >>> isinstance(fn, int)
            True
            >>> fn
            2

            >>> # Get per-class counts
            >>> per_class_fn = cm.false_negative_count(per_class=True)
            >>> 'actual' in per_class_fn.keys()
            True
            >>> 'false_negative_count' in per_class_fn.keys()
            True
        """
        from .exceptions import DatasetNotConfusionMatrixError

        # Check if this is a confusion matrix
        if not hasattr(self, "_is_confusion_matrix") or not self._is_confusion_matrix:
            raise DatasetNotConfusionMatrixError(
                "false_negative_count() can only be called on a Dataset created by confusion_matrix(). "
                "Create a confusion matrix first using: dataset.confusion_matrix(actual_field, predicted_field)"
            )

        # Get metadata
        metadata = self._confusion_matrix_metadata
        actual_field = metadata["actual_field"]
        actual_values = metadata["actual_values"]

        # Get the data as a pandas DataFrame
        df = self.to_pandas()

        # Extract false negative counts (row sum - diagonal)
        fn_counts = []
        for i, actual_val in enumerate(actual_values):
            row = df[df[actual_field] == actual_val]
            if len(row) == 0:
                fn_counts.append(0)
                continue

            # Get all prediction columns
            prob_cols = [col for col in df.columns if col != actual_field]
            row_values = [int(row[col].values[0]) for col in prob_cols]

            # Row sum
            row_sum = sum(row_values)

            # True positive (diagonal element)
            tp = row_values[i] if i < len(row_values) else 0

            # False negatives = row_sum - tp
            fn_count = row_sum - tp
            fn_counts.append(fn_count)

        if per_class:
            # Return a Dataset with per-class counts
            from ..dataset import Dataset

            return Dataset(
                [
                    {actual_field: list(actual_values)},
                    {"false_negative_count": fn_counts},
                ]
            )
        else:
            # Return total count
            return sum(fn_counts)

    def percent_correctly_predicted(
        self, per_class: bool = False
    ) -> Union[float, "Dataset"]:
        """
        Calculate the percentage of correctly predicted instances from a confusion matrix.

        This is the overall accuracy of the classifier (for per_class=False) or the
        per-class recall/sensitivity (for per_class=True).

        This method can only be called on a Dataset that was created by the confusion_matrix()
        method.

        Args:
            per_class: If True, returns a Dataset with accuracy for each class (recall).
                      If False (default), returns the overall accuracy as a percentage.

        Returns:
            If per_class=False: A float representing the overall accuracy percentage (0-100)
            If per_class=True: A Dataset with per-class accuracy (recall) percentages

        Raises:
            DatasetNotConfusionMatrixError: If called on a Dataset that is not a confusion matrix

        Notes:
            - Overall accuracy = (sum of diagonal) / (total predictions) * 100
            - Per-class accuracy (recall) = true_positives / (true_positives + false_negatives) * 100
            - This is also known as accuracy (overall) or recall/sensitivity (per-class)

        Examples:
            >>> from edsl.dataset import Dataset
            >>> d = Dataset([
            ...     {'actual': ['cat', 'cat', 'dog', 'dog', 'dog']},
            ...     {'predicted': ['cat', 'dog', 'dog', 'dog', 'cat']}
            ... ])
            >>> cm = d.confusion_matrix('actual', 'predicted')
            >>> acc = cm.percent_correctly_predicted()
            >>> isinstance(acc, float)
            True
            >>> 0 <= acc <= 100
            True

            >>> # Get per-class accuracy (recall)
            >>> per_class_acc = cm.percent_correctly_predicted(per_class=True)
            >>> 'actual' in per_class_acc.keys()
            True
            >>> 'percent_correct' in per_class_acc.keys()
            True
        """
        from .exceptions import DatasetNotConfusionMatrixError

        # Check if this is a confusion matrix
        if not hasattr(self, "_is_confusion_matrix") or not self._is_confusion_matrix:
            raise DatasetNotConfusionMatrixError(
                "percent_correctly_predicted() can only be called on a Dataset created by confusion_matrix(). "
                "Create a confusion matrix first using: dataset.confusion_matrix(actual_field, predicted_field)"
            )

        # Get metadata
        metadata = self._confusion_matrix_metadata
        actual_field = metadata["actual_field"]
        actual_values = metadata["actual_values"]

        # Get the data as a pandas DataFrame
        df = self.to_pandas()

        # Calculate per-class percentages
        percentages = []
        total_correct = 0
        total_samples = 0

        for i, actual_val in enumerate(actual_values):
            row = df[df[actual_field] == actual_val]
            if len(row) == 0:
                percentages.append(0.0)
                continue

            # Get all prediction columns
            prob_cols = [col for col in df.columns if col != actual_field]
            row_values = [int(row[col].values[0]) for col in prob_cols]

            # Row sum (total instances of this class)
            row_sum = sum(row_values)

            # True positive (diagonal element)
            tp = row_values[i] if i < len(row_values) else 0

            # Calculate percentage for this class
            if row_sum > 0:
                percentage = (tp / row_sum) * 100
            else:
                percentage = 0.0

            percentages.append(percentage)
            total_correct += tp
            total_samples += row_sum

        if per_class:
            # Return a Dataset with per-class percentages
            from ..dataset import Dataset

            return Dataset(
                [{actual_field: list(actual_values)}, {"percent_correct": percentages}]
            )
        else:
            # Return overall accuracy
            if total_samples > 0:
                return (total_correct / total_samples) * 100
            else:
                return 0.0

    def tally(
        self, *fields: Optional[str], top_n: Optional[int] = None, output="Dataset"
    ) -> Union[dict, "Dataset"]:
        """
        Count frequency distributions of values in specified fields.

        This method tallies the occurrence of unique values within one or more fields,
        similar to a GROUP BY and COUNT in SQL. When multiple fields are provided, it
        performs cross-tabulation across those fields.

        Parameters:
            *fields: Field names to tally. If none provided, uses all available fields.
            top_n: Optional limit to return only the top N most frequent values.
            output: Format for results, either "Dataset" (recommended) or "dict".

        Returns:
            By default, returns a Dataset with columns for the field(s) and a 'count' column.
            If output="dict", returns a dictionary mapping values to counts.

        Notes:
            - For single fields, returns counts of each unique value
            - For multiple fields, returns counts of each unique combination of values
            - Results are sorted in descending order by count
            - Fields can be specified with or without their type prefix

        Examples:
            >>> from edsl import Results
            >>> r = Results.example()

            # Single field frequency count
            >>> r.select('how_feeling').tally('answer.how_feeling', output="dict")
            {'OK': 2, 'Great': 1, 'Terrible': 1}

            # Return as Dataset (default)
            >>> from edsl.dataset import Dataset
            >>> expected = Dataset([{'answer.how_feeling': ['OK', 'Great', 'Terrible']}, {'count': [2, 1, 1]}])
            >>> r.select('how_feeling').tally('answer.how_feeling', output="Dataset") == expected
            True

            # Multi-field cross-tabulation - exact output varies based on data
            >>> result = r.tally('how_feeling', 'how_feeling_yesterday')
            >>> 'how_feeling' in result.keys() and 'how_feeling_yesterday' in result.keys() and 'count' in result.keys()
            True
        """
        from collections import Counter

        if len(fields) == 0:
            fields = self.relevant_columns()

        relevant_columns_without_prefix = [
            column.split(".")[-1] for column in self.relevant_columns()
        ]

        if not all(
            f in self.relevant_columns() or f in relevant_columns_without_prefix
            for f in fields
        ):
            raise DatasetKeyError(
                "One or more specified fields are not in the dataset."
                f"The available fields are: {self.relevant_columns()}"
            )

        if len(fields) == 1:
            field = fields[0]
            values = self._key_to_value(field)
        else:
            values = list(zip(*(self._key_to_value(field) for field in fields)))

        for value in values:
            if isinstance(value, list):
                value = tuple(value)
        try:
            tally = dict(Counter(values))
        except TypeError:
            tally = dict(Counter([str(v) for v in values]))
        except Exception as e:
            raise DatasetValueError(f"Error tallying values: {e}")

        sorted_tally = dict(sorted(tally.items(), key=lambda item: -item[1]))
        if top_n is not None:
            sorted_tally = dict(list(sorted_tally.items())[:top_n])

        from ..dataset import Dataset

        if output == "dict":
            # why did I do this?
            warnings.warn(
                textwrap.dedent(
                    """\
                        The default output from tally will change to Dataset in the future.
                        Use output='Dataset' to get the Dataset object for now.
                        """
                )
            )
            return sorted_tally
        elif output == "Dataset":
            dataset = Dataset(
                [
                    {"value": list(sorted_tally.keys())},
                    {"count": list(sorted_tally.values())},
                ]
            )
            # return dataset
            sl = dataset.to_scenario_list().unpack(
                "value",
                new_names=[fields] if isinstance(fields, str) else fields,
                keep_original=False,
            )
            keys = list(sl[0].keys())
            keys.remove("count")
            keys.append("count")
            return sl.reorder_keys(keys).to_dataset()

    def flatten(self, field: str, keep_original: bool = False) -> "Dataset":
        """
        Expand a field containing dictionaries into separate fields.

        This method takes a field that contains a list of dictionaries and expands
        it into multiple fields, one for each key in the dictionaries. This is useful
        when working with nested data structures or results from extraction operations.

        Parameters:
            field: The field containing dictionaries to flatten
            keep_original: Whether to retain the original field in the result

        Returns:
            A new Dataset with the dictionary keys expanded into separate fields

        Notes:
            - Each key in the dictionaries becomes a new field with name pattern "{field}.{key}"
            - All dictionaries in the field must have compatible structures
            - If a dictionary is missing a key, the corresponding value will be None
            - Non-dictionary values in the field will cause a warning
        Examples:
            >>> from edsl.dataset import Dataset

            # Basic flattening of nested dictionaries
            >>> Dataset([{'a': [{'a': 1, 'b': 2}]}, {'c': [5]}]).flatten('a')
            Dataset([{'c': [5]}, {'a.a': [1]}, {'a.b': [2]}])

            # Works with prefixed fields too
            >>> Dataset([{'answer.example': [{'a': 1, 'b': 2}]}, {'c': [5]}]).flatten('answer.example')
            Dataset([{'c': [5]}, {'answer.example.a': [1]}, {'answer.example.b': [2]}])

            # Keep the original field if needed
            >>> d = Dataset([{'a': [{'a': 1, 'b': 2}]}, {'c': [5]}])
            >>> d.flatten('a', keep_original=True)
            Dataset([{'a': [{'a': 1, 'b': 2}]}, {'c': [5]}, {'a.a': [1]}, {'a.b': [2]}])

            # Can also use unambiguous unprefixed field name
            >>> result = Dataset([{'answer.pros_cons': [{'pros': ['Safety'], 'cons': ['Cost']}]}]).flatten('pros_cons')
            >>> sorted(result.keys()) == ['answer.pros_cons.cons', 'answer.pros_cons.pros']
            True
            >>> sorted(result.to_dicts()[0].items()) == sorted({'cons': ['Cost'], 'pros': ['Safety']}.items())
            True
        """
        from ..dataset import Dataset

        # Ensure the dataset isn't empty
        if not self.data:
            return self.copy()

        # First try direct match with the exact field name
        field_entry = None
        for entry in self.data:
            col_name = next(iter(entry.keys()))
            if field == col_name:
                field_entry = entry
                break

        # If not found, try to match by unprefixed name
        if field_entry is None:
            # Find any columns that have field as their unprefixed name
            candidates = []
            for entry in self.data:
                col_name = next(iter(entry.keys()))
                if "." in col_name:
                    prefix, col_field = col_name.split(".", 1)
                    if col_field == field:
                        candidates.append(entry)

            # If we found exactly one match by unprefixed name, use it
            if len(candidates) == 1:
                field_entry = candidates[0]
            # If we found multiple matches, it's ambiguous
            elif len(candidates) > 1:
                matching_cols = [next(iter(entry.keys())) for entry in candidates]
                from .exceptions import DatasetValueError

                raise DatasetValueError(
                    f"Ambiguous field name '{field}'. It matches multiple columns: {matching_cols}. "
                    f"Please specify the full column name to flatten."
                )
            # If no candidates by unprefixed name, check partial matches
            else:
                partial_matches = []
                for entry in self.data:
                    col_name = next(iter(entry.keys()))
                    if "." in col_name and (
                        col_name.endswith("." + field)
                        or col_name.startswith(field + ".")
                    ):
                        partial_matches.append(entry)

                # If we found exactly one partial match, use it
                if len(partial_matches) == 1:
                    field_entry = partial_matches[0]
                # If we found multiple partial matches, it's ambiguous
                elif len(partial_matches) > 1:
                    matching_cols = [
                        next(iter(entry.keys())) for entry in partial_matches
                    ]
                    from .exceptions import DatasetValueError

                    raise DatasetValueError(
                        f"Ambiguous field name '{field}'. It matches multiple columns: {matching_cols}. "
                        f"Please specify the full column name to flatten."
                    )

        # Get the number of observations
        num_observations = self.num_observations()

        # If we still haven't found the field, it's not in the dataset
        if field_entry is None:
            warnings.warn(
                f"Field '{field}' not found in dataset, returning original dataset"
            )
            return self.copy()

        # Get the actual field name as it appears in the data
        actual_field = next(iter(field_entry.keys()))

        # Create new dictionary for flattened data
        flattened_data = []

        # Copy all existing columns except the one we're flattening (if keep_original is False)
        for entry in self.data:
            col_name = next(iter(entry.keys()))
            if col_name != actual_field or keep_original:
                flattened_data.append(entry.copy())

        # Get field data and make sure it's valid
        field_values = field_entry[actual_field]
        if not all(isinstance(item, dict) for item in field_values if item is not None):
            warnings.warn(
                f"Field '{actual_field}' contains non-dictionary values that cannot be flattened"
            )
            return self.copy()

        # Collect all unique keys across all dictionaries
        all_keys = set()
        for item in field_values:
            if isinstance(item, dict):
                all_keys.update(item.keys())

        # Create new columns for each key
        for key in sorted(all_keys):  # Sort for consistent output
            new_values = []
            for i in range(num_observations):
                value = None
                if i < len(field_values) and isinstance(field_values[i], dict):
                    value = field_values[i].get(key, None)
                new_values.append(value)

            # Add this as a new column
            flattened_data.append({f"{actual_field}.{key}": new_values})

        # Return a new Dataset with the flattened data
        return Dataset(flattened_data)

    def unpack_list(
        self,
        field: str,
        new_names: Optional[List[str]] = None,
        keep_original: bool = True,
    ) -> "Dataset":
        """Unpack list columns into separate columns with provided names or numeric suffixes.

        For example, if a dataset contains:
        [{'data': [[1, 2, 3], [4, 5, 6]], 'other': ['x', 'y']}]

        After d.unpack_list('data'), it should become:
        [{'other': ['x', 'y'], 'data_1': [1, 4], 'data_2': [2, 5], 'data_3': [3, 6]}]

        Args:
            field: The field containing lists to unpack
            new_names: Optional list of names for the unpacked fields. If None, uses numeric suffixes.
            keep_original: If True, keeps the original field in the dataset

        Returns:
            A new Dataset with unpacked columns

        Examples:
            >>> from edsl.dataset import Dataset
            >>> d = Dataset([{'data': [[1, 2, 3], [4, 5, 6]]}])
            >>> d.unpack_list('data')
            Dataset([{'data': [[1, 2, 3], [4, 5, 6]]}, {'data_1': [1, 4]}, {'data_2': [2, 5]}, {'data_3': [3, 6]}])

            >>> d.unpack_list('data', new_names=['first', 'second', 'third'])
            Dataset([{'data': [[1, 2, 3], [4, 5, 6]]}, {'first': [1, 4]}, {'second': [2, 5]}, {'third': [3, 6]}])
        """
        from .dataset import Dataset

        # Create a copy of the dataset
        result = Dataset(self.data.copy())

        # Find the field in the dataset
        field_index = None
        for i, entry in enumerate(result.data):
            if field in entry:
                field_index = i
                break

        if field_index is None:
            raise DatasetKeyError(f"Field '{field}' not found in dataset")

        field_data = result.data[field_index][field]

        # Check if values are lists
        if not all(isinstance(v, list) for v in field_data):
            raise DatasetTypeError(
                f"Field '{field}' does not contain lists in all entries"
            )

        # Get the maximum length of lists
        max_len = max(len(v) for v in field_data)

        # Create new fields for each index
        for i in range(max_len):
            if new_names and i < len(new_names):
                new_field = new_names[i]
            else:
                new_field = f"{field}_{i+1}"

            # Extract the i-th element from each list
            new_values = []
            for item in field_data:
                new_values.append(item[i] if i < len(item) else None)

            result.data.append({new_field: new_values})

        # Remove the original field if keep_original is False
        if not keep_original:
            result.data.pop(field_index)

        return result

    def remove_prefix(self):
        """Returns a new Dataset with the prefix removed from all column names.

        The prefix is defined as everything before the first dot (.) in the column name.
        If removing prefixes would result in duplicate column names, an exception is raised.

        Returns:
            Dataset: A new Dataset with prefixes removed from column names

        Raises:
            ValueError: If removing prefixes would result in duplicate column names

        Examples:
            >>> from edsl.results import Results
            >>> r = Results.example()
            >>> r.select('how_feeling', 'how_feeling_yesterday').relevant_columns()
            ['answer.how_feeling', 'answer.how_feeling_yesterday']
            >>> r.select('how_feeling', 'how_feeling_yesterday').remove_prefix().relevant_columns()
            ['how_feeling', 'how_feeling_yesterday']

            >>> from edsl.dataset import Dataset
            >>> d = Dataset([{'a.x': [1, 2, 3]}, {'b.x': [4, 5, 6]}])
            >>> # d.remove_prefix()

        # Testing remove_prefix with duplicate column names raises DatasetValueError - tested in unit tests
        """
        from .dataset import Dataset

        # Get all column names
        columns = self.relevant_columns()

        # Extract the unprefixed names
        unprefixed = {}
        duplicates = set()

        for col in columns:
            if "." in col:
                unprefixed_name = col.split(".", 1)[1]
                if unprefixed_name in unprefixed:
                    duplicates.add(unprefixed_name)
                unprefixed[unprefixed_name] = col
            else:
                # For columns without a prefix, keep them as is
                unprefixed[col] = col

        # Check for duplicates
        if duplicates:
            raise DatasetValueError(
                f"Removing prefixes would result in duplicate column names: {sorted(list(duplicates))}"
            )

        # Create a new dataset with unprefixed column names
        new_data = []
        for entry in self.data:
            key, values = list(entry.items())[0]
            if "." in key:
                new_key = key.split(".", 1)[1]
            else:
                new_key = key
            new_data.append({new_key: values})

        return Dataset(new_data)

    def report_from_template(
        self,
        template: str,
        *fields: Optional[str],
        top_n: Optional[int] = None,
        remove_prefix: bool = True,
        return_string: bool = False,
        format: str = "text",
        filename: Optional[str] = None,
        separator: str = "\n\n",
        observation_title_template: Optional[str] = None,
        explode: bool = False,
        filestore: bool = False,
    ) -> Optional[Union[str, "Document", List, "FileStore"]]:
        """Generates a report using a Jinja2 template for each row in the dataset.

        This method renders a user-provided Jinja2 template for each observation in the dataset,
        with template variables populated from the row data. This allows for completely customized
        report formatting using pandoc for advanced output formats.

        Args:
            template: Jinja2 template string to render for each row
            *fields: The fields to include in template context. If none provided, all fields are used.
            top_n: Optional limit on the number of observations to include.
            remove_prefix: Whether to remove type prefixes (e.g., "answer.") from field names in template context.
            return_string: If True, returns the rendered content. If False (default in notebooks),
                          only displays the content without returning.
            format: Output format - one of "text", "html", "pdf", or "docx". Formats other than "text" require pandoc.
            filename: If provided, saves the rendered content to this file. For exploded output,
                     this becomes a template (e.g., "report_{index}.html").
            separator: String to use between rendered templates for each row (ignored when explode=True).
            observation_title_template: Optional Jinja2 template for observation titles.
                                       Defaults to "Observation {index}" where index is 1-based.
                                       Template has access to all row data plus 'index' and 'index0' variables.
            explode: If True, creates separate files for each observation instead of one combined file.
            filestore: If True, wraps the generated file(s) in FileStore object(s). If no filename is provided,
                      creates temporary files. For exploded output, returns a list of FileStore objects.

        Returns:
            Depending on explode, format, return_string, and filestore:
            - For text format: String content or None (if displayed in notebook)
            - For html format: HTML string content or None (if displayed in notebook)
            - For docx format: Document object or None (if saved to file)
            - For pdf format: PDF bytes or None (if saved to file)
            - If explode=True: List of created filenames (when filename provided) or list of documents/content
            - If filestore=True: FileStore object(s) containing the generated file(s)

        Notes:
            - Pandoc is required for HTML, PDF, and DOCX output formats
            - Templates are treated as Markdown for all non-text formats
            - PDF output uses XeLaTeX engine through pandoc
            - HTML output includes standalone document structure

        Examples:
            >>> from edsl.results import Results
            >>> r = Results.example()
            >>> template = "Person feels: {{ how_feeling }}"
            >>> report = r.select('how_feeling').report_from_template(template, return_string=True)
            >>> "Person feels: OK" in report
            True
            >>> "Person feels: Great" in report
            True

            # Custom observation titles
            >>> custom_title = "Response {{ index }}: {{ how_feeling }}"
            >>> report = r.select('how_feeling').report_from_template(
            ...     template, observation_title_template=custom_title, return_string=True)
            >>> "Response 1: OK" in report
            True

            # HTML output (requires pandoc)
            >>> html_report = r.select('how_feeling').report_from_template(
            ...     template, format="html", return_string=True)  # doctest: +SKIP
            >>> # Creates HTML with proper document structure

            # PDF output (requires pandoc with XeLaTeX)
            >>> pdf_report = r.select('how_feeling').report_from_template(
            ...     template, format="pdf")  # doctest: +SKIP
            >>> # Returns PDF bytes

            # Basic template functionality
            >>> template2 = "Feeling: {{ how_feeling }}, Index: {{ index }}"
            >>> report2 = r.select('how_feeling').report_from_template(
            ...     template2, return_string=True, top_n=2)
            >>> "Feeling: OK, Index: 1" in report2
            True
        """
        from .report_from_template import TemplateReportGenerator

        generator = TemplateReportGenerator(self)
        return generator.generate_report(
            template,
            *fields,
            top_n=top_n,
            remove_prefix=remove_prefix,
            return_string=return_string,
            format=format,
            filename=filename,
            separator=separator,
            observation_title_template=observation_title_template,
            explode=explode,
            filestore=filestore,
        )


def to_dataset(func):
    """
    Decorator that ensures functions receive a Dataset object as their first argument.

    This decorator automatically converts various EDSL container objects (Results,
    AgentList, ScenarioList) to Dataset objects before passing them to the decorated
    function. This allows methods defined in DataOperationsBase to work seamlessly
    across different container types without duplicating conversion logic.

    Parameters:
        func: The function to decorate

    Returns:
        A wrapped function that ensures its first argument is a Dataset

    Notes:
        - For Results objects, calls select() to convert to a Dataset
        - For AgentList and ScenarioList objects, calls their to_dataset() method
        - For Dataset objects, passes them through unchanged
        - This decorator is used internally by the mixin system to enable method sharing
    """

    @wraps(func)
    def wrapper(self, *args, **kwargs):
        """Execute the function with self converted to a Dataset if needed."""
        # Convert to Dataset based on the class type
        if self.__class__.__name__ == "Results":
            dataset_self = self.select()
        elif self.__class__.__name__ == "AgentList":
            dataset_self = self.to_dataset()
        elif self.__class__.__name__ == "ScenarioList":
            dataset_self = self.to_dataset()
        else:
            dataset_self = self

        # Call the function with the converted self
        return func(dataset_self, *args, **kwargs)

    # Mark the wrapper as being wrapped by to_dataset
    wrapper._is_wrapped = True
    return wrapper


def decorate_methods_from_mixin(cls, mixin_cls):
    """
    Apply the to_dataset decorator to methods inherited from a mixin class.

    This function is part of EDSL's method inheritance system. It takes methods
    from a source mixin class, applies the to_dataset decorator to them, and adds
    them to a target class. This enables the sharing of data manipulation methods
    across different container types while ensuring they receive the right data type.

    The function is careful not to override methods that are already defined in
    more specific parent classes, preserving the method resolution order (MRO).

    Parameters:
        cls: The target class to add decorated methods to
        mixin_cls: The source mixin class providing the methods

    Returns:
        The modified target class with decorated methods added

    Notes:
        - Only public methods (not starting with "_") are decorated and added
        - Methods already defined in more specific parent classes are not overridden
        - Methods from DataOperationsBase are not skipped to ensure all base methods are available
    """
    # Get all attributes, including inherited ones
    for attr_name in dir(mixin_cls):
        # Skip magic methods and private methods
        if not attr_name.startswith("_"):
            attr_value = getattr(mixin_cls, attr_name)
            if callable(attr_value):
                # Check if the method is already defined in the class's MRO
                # but skip DataOperationsBase methods
                for base in cls.__mro__[1:]:  # Skip the class itself
                    if attr_name in base.__dict__ and base is not DataOperationsBase:
                        # Method is overridden in a more specific class, skip decorating
                        break
                else:
                    # Method not overridden, safe to decorate
                    setattr(cls, attr_name, to_dataset(attr_value))
    return cls


# def decorate_methods_from_mixin(cls, mixin_cls):
#     """Decorates all methods from mixin_cls with to_dataset decorator."""

#     # Get all attributes, including inherited ones
#     for attr_name in dir(mixin_cls):
#         # Skip magic methods and private methods
#         if not attr_name.startswith('_'):
#             attr_value = getattr(mixin_cls, attr_name)
#             if callable(attr_value):
#                 setattr(cls, attr_name, to_dataset(attr_value))
#     return cls


class DatasetOperationsMixin(DataOperationsBase):
    """
    Mixin providing data manipulation operations for Dataset objects.

    This mixin class is the cornerstone of EDSL's data manipulation system. It directly
    inherits methods from DataOperationsBase without requiring conversion, as it's
    designed specifically for the Dataset class. It serves as the primary implementation
    of all data operations methods that other container types will inherit and adapt
    through the to_dataset decorator.

    The design follows a standard mixin pattern where common functionality is defined
    in a standalone class that can be "mixed in" to other classes. In EDSL's case,
    this allows different container types (Results, AgentList, ScenarioList) to share
    the same powerful data manipulation interface.

    Key features:

    1. Data Transformation:
       - Filtering with `filter()`
       - Creating new columns with `mutate()`
       - Reshaping with `long()`, `wide()`, `flatten()`, etc.
       - Selecting specific data with `select()`

    2. Visualization:
       - Table display with `table()`
       - R integration with `ggplot2()`
       - Report generation with `report()`

    3. Data Export:
       - To files with `to_csv()`, `to_excel()`, etc.
       - To other formats with `to_pandas()`, `to_dicts()`, etc.

    4. Analysis:
       - SQL queries with `sql()`
       - Aggregation with `tally()`
       - Tree-based exploration with `tree()`

    This mixin is designed for fluent method chaining, allowing complex data manipulation
    pipelines to be built in an expressive and readable way.
    """

    pass


class ResultsOperationsMixin(DataOperationsBase):
    """
    Mixin providing data operations for Results objects.

    This mixin adapts DatasetOperationsMixin methods to work with Results objects.
    When a method is called on a Results object, it's automatically converted to
    a Dataset first via the to_dataset decorator applied in __init_subclass__.

    This allows Results objects to have the same data manipulation capabilities
    as Dataset objects without duplicating code.
    """

    def __init_subclass__(cls, **kwargs):
        """
        Automatically decorate all methods from DatasetOperationsMixin.

        This hook runs when a class inherits from ResultsOperationsMixin,
        applying the to_dataset decorator to all methods from DatasetOperationsMixin.
        """
        super().__init_subclass__(**kwargs)
        decorate_methods_from_mixin(cls, DatasetOperationsMixin)


class ScenarioListOperationsMixin(DataOperationsBase):
    """
    Mixin providing data operations for ScenarioList objects.

    This mixin adapts DatasetOperationsMixin methods to work with ScenarioList objects.
    ScenarioList objects are converted to Dataset objects before method execution
    via the to_dataset decorator applied in __init_subclass__.
    """

    def kl_divergence(
        self,
        group_field: str,
        value_field: str,
        from_group: Optional[str] = None,
        to_group: Optional[str] = None,
        bins: Optional[Union[int, str]] = None,
        base: float = 2.0,
        laplace_smooth: float = 1e-10,
    ) -> Union[float, dict]:
        """
        Compute KL divergence between distributions defined by groups.

        Measures how much one probability distribution diverges from another.
        Useful for comparing distributions across experimental conditions, agent
        personas, prompt variations, etc.

        Parameters:
            group_field: Field that defines the groups (e.g., 'condition', 'persona')
            value_field: Field containing values to compare distributions of
            from_group: The reference group (P in KL(P||Q)). If None, compute all pairs.
            to_group: The comparison group (Q in KL(P||Q)). Required if from_group specified.
            bins: For continuous data - number of bins or 'auto' (default: None = categorical)
            base: Logarithm base (2=bits, e=nats, 10=dits, default: 2)
            laplace_smooth: Small value to avoid log(0) (default: 1e-10)

        Returns:
            float: KL divergence value if from_group and to_group specified
            dict: All pairwise KL divergences if groups not specified

        Examples:
            >>> from edsl.scenarios import Scenario, ScenarioList
            >>> sl = ScenarioList([
            ...     Scenario({'condition': 'control', 'response': 'positive'}),
            ...     Scenario({'condition': 'control', 'response': 'positive'}),
            ...     Scenario({'condition': 'control', 'response': 'neutral'}),
            ...     Scenario({'condition': 'treatment', 'response': 'negative'}),
            ...     Scenario({'condition': 'treatment', 'response': 'neutral'}),
            ... ])
            >>> # Compare two specific groups
            >>> kl = sl.kl_divergence('condition', 'response', 'control', 'treatment')  # doctest: +SKIP
            >>> # Get all pairwise comparisons
            >>> kl_all = sl.kl_divergence('condition', 'response')  # doctest: +SKIP

        Notes:
            - KL divergence is asymmetric: KL(P||Q) ≠ KL(Q||P)
            - KL(P||Q) measures how much P diverges from Q
            - For categorical data, leave bins=None
            - For continuous data, set bins to number or 'auto'
            - Use base=2 for bits, base=e for nats
        """
        from .kl_divergence import kl_divergence

        return kl_divergence(
            self,
            group_field=group_field,
            value_field=value_field,
            from_group=from_group,
            to_group=to_group,
            bins=bins,
            base=base,
            laplace_smooth=laplace_smooth,
        )

    def __init_subclass__(cls, **kwargs):
        """
        Automatically decorate all methods from DatasetOperationsMixin.

        This hook runs when a class inherits from ScenarioListOperationsMixin,
        applying the to_dataset decorator to all methods from DatasetOperationsMixin.
        """
        super().__init_subclass__(**kwargs)
        decorate_methods_from_mixin(cls, DatasetOperationsMixin)


class AgentListOperationsMixin(DataOperationsBase):
    """
    Mixin providing data operations for AgentList objects.

    This mixin adapts DatasetOperationsMixin methods to work with AgentList objects.
    AgentList objects are converted to Dataset objects before method execution
    via the to_dataset decorator applied in __init_subclass__.
    """

    def __init_subclass__(cls, **kwargs):
        """
        Automatically decorate all methods from DatasetOperationsMixin.

        This hook runs when a class inherits from AgentListOperationsMixin,
        applying the to_dataset decorator to all methods from DatasetOperationsMixin.
        """
        super().__init_subclass__(**kwargs)
        decorate_methods_from_mixin(cls, DatasetOperationsMixin)


if __name__ == "__main__":
    import doctest

    doctest.testmod(optionflags=doctest.ELLIPSIS)
