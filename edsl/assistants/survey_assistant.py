#!/usr/bin/env python3
# /// script
# requires-python = ">=3.9"
# dependencies = [
#     "claude-agent-sdk",
#     "anyio>=3.6.0",
#     "typer>=0.9.0",
#     "rich>=13.0.0",
#     "edsl>=0.1.0",
#     "PyYAML>=6.0.0",
#     "pandas>=2.0.0",
#     "python-dotenv>=1.0.0",
# ]
# ///
"""
EDSL Survey Generator using Claude Agent SDK
Processes any document type and creates EDSL survey objects via YAML

Supported formats with smart handling:
- QSF: Pre-parsed to extract Qualtrics question structure
- CSV/TSV: Sampled (first 20 rows) for large files
- XLSX/XLS: Sampled (first 20 rows) for large files
- JSON: Sent as-is
- Text files: Sent as-is

Usage:
    uv run edsl_survey_generator.py file1.txt file2.csv file3.qsf file4.xlsx
"""

import anyio
import asyncio
import gzip
import json
import os
import sys
import io
import logging
import re
from html import unescape
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import yaml
import typer
from dotenv import load_dotenv
import csv
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich.table import Table
from rich.live import Live
from rich.text import Text
from rich.spinner import Spinner
from rich.columns import Columns

# Import for accessing package data
try:
    from importlib.resources import files
except ImportError:
    # Python < 3.9 fallback
    from importlib_resources import files

# Load environment variables from .env file
load_dotenv()


# Lazy import function for claude_agent_sdk
def _check_claude_sdk():
    """Check and import claude_agent_sdk when needed."""
    try:
        from claude_agent_sdk import (
            AssistantMessage,
            ClaudeAgentOptions,
            ResultMessage,
            TextBlock,
            ToolUseBlock,
            UserMessage,
            query,
        )

        return True
    except ImportError:
        print("\n" + "=" * 70)
        print("ERROR: claude-agent-sdk is not installed")
        print("=" * 70)
        print("\nTo install it, run:")
        print("  pip install claude-agent-sdk")
        print("\nOr if using uv:")
        print("  uv pip install claude-agent-sdk")
        print("\nFor more information, visit:")
        print("  https://github.com/anthropics/claude-agent-sdk-python")
        print("=" * 70 + "\n")
        raise


from edsl import QuestionBase, Survey
from edsl.questions import *  # Import all question types

# Import helper modules
from edsl.assistants.survey_assistant_helpers import (
    SurveyGenerator,
    ColumnMapper,
    read_pdf_file,
)

app = typer.Typer()
console = Console()


class SuppressOutput:
    """Context manager to suppress stdout temporarily."""

    def __enter__(self):
        self._original_stdout = sys.stdout
        sys.stdout = io.StringIO()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        sys.stdout = self._original_stdout


def sanitize_column_name(column_name: str) -> str:
    """Convert a data column name to a valid Python identifier for question_name.

    This uses the same logic that we instruct Claude to use, ensuring consistency.

    Args:
        column_name: Original column name from data file

    Returns:
        Sanitized name suitable for use as a Python identifier

    Examples:
        >>> sanitize_column_name("Reflective motivation 1 RM1")
        'reflective_motivation_1_rm1'
        >>> sanitize_column_name("Automatic motivation -average")
        'automatic_motivation_average'
        >>> sanitize_column_name("age")
        'age'
    """
    import re

    # Convert to lowercase
    name = column_name.lower()

    # Replace spaces and dashes with underscores
    name = name.replace(" ", "_").replace("-", "_")

    # Remove any remaining special characters except underscores and alphanumerics
    name = re.sub(r"[^a-z0-9_]", "", name)

    # Remove consecutive underscores
    name = re.sub(r"_+", "_", name)

    # Remove leading/trailing underscores
    name = name.strip("_")

    # Ensure it doesn't start with a number (prepend 'q_' if it does)
    if name and name[0].isdigit():
        name = "q_" + name

    return name


def sample_csv_file(
    csv_content: str, max_rows: int = 10, delimiter: str = ","
) -> Dict[str, Any]:
    """Sample a CSV/TSV file to provide structure without sending all data.

    Args:
        csv_content: Raw CSV/TSV file content
        max_rows: Maximum number of data rows to include (default: 10)
        delimiter: Field delimiter (default: ',')

    Returns:
        Dictionary with CSV structure and sample data
    """
    try:
        lines = csv_content.strip().split("\n")
        total_rows = len(lines) - 1  # Exclude header

        # Parse CSV/TSV
        csv_reader = csv.reader(io.StringIO(csv_content), delimiter=delimiter)
        rows = list(csv_reader)

        if not rows:
            return {"error": "Empty file"}

        header = rows[0]
        data_rows = rows[1:]

        # Sample data
        sample_rows = data_rows[:max_rows]

        return {
            "columns": header,
            "column_count": len(header),
            "total_rows": len(data_rows),
            "sample_rows": sample_rows,
            "sampled": len(data_rows) > max_rows,
            "delimiter": delimiter,
        }
    except Exception as e:
        return {"error": f"Failed to parse file: {e}"}


def sample_stata_spss_file(
    file_path: Path, file_type: str, max_rows: int = 10
) -> Dict[str, Any]:
    """Sample a Stata (.dta) or SPSS (.sav) file to provide structure without sending all data.

    Args:
        file_path: Path to file
        file_type: Either 'stata' or 'spss'
        max_rows: Maximum number of data rows to include (default: 10)

    Returns:
        Dictionary with file structure and sample data
    """
    try:
        try:
            import pyreadstat
        except ImportError:
            return {
                "error": f"pyreadstat is required for reading {file_type.upper()} files. Install with: pip install pyreadstat"
            }

        # Read the file
        if file_type == "stata":
            df, meta = pyreadstat.read_dta(str(file_path))
        elif file_type == "spss":
            df, meta = pyreadstat.read_sav(str(file_path))
        else:
            return {"error": f"Unknown file type: {file_type}"}

        # Sample the data
        total_rows = len(df)
        df_sample = df.head(max_rows)

        # Get column info
        columns = df.columns.tolist()
        column_labels = (
            meta.column_names_to_labels
            if hasattr(meta, "column_names_to_labels")
            else {}
        )
        value_labels = meta.value_labels if hasattr(meta, "value_labels") else {}

        # Convert to list of lists for display
        sample_rows = df_sample.values.tolist()

        return {
            "columns": columns,
            "column_count": len(columns),
            "total_rows": total_rows,
            "sample_rows": sample_rows,
            "sampled": total_rows > max_rows,
            "column_labels": column_labels,
            "value_labels": value_labels,
            "variable_count": meta.number_columns
            if hasattr(meta, "number_columns")
            else len(columns),
        }
    except Exception as e:
        return {"error": f"Failed to parse {file_type.upper()} file: {e}"}


def sample_xlsx_file(file_path: Path, max_rows: int = 10) -> Dict[str, Any]:
    """Sample an XLSX file to provide structure without sending all data.

    Args:
        file_path: Path to XLSX file
        max_rows: Maximum number of data rows to include (default: 10)

    Returns:
        Dictionary with XLSX structure and sample data
    """
    try:
        try:
            import openpyxl

            use_openpyxl = True
        except ImportError:
            use_openpyxl = False

        if use_openpyxl:
            # Use openpyxl (lighter weight)
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            sheet = wb.active

            rows = list(sheet.iter_rows(values_only=True))
            wb.close()
        else:
            # Fall back to pandas
            try:
                import pandas as pd

                df = pd.read_excel(file_path, nrows=max_rows + 1)  # +1 for header
                header = df.columns.tolist()
                data_rows = df.values.tolist()
                rows = [header] + data_rows
            except ImportError:
                return {
                    "error": "Neither openpyxl nor pandas is available for reading XLSX files"
                }

        if not rows:
            return {"error": "Empty XLSX file"}

        header = [str(cell) if cell is not None else "" for cell in rows[0]]
        data_rows = [
            [str(cell) if cell is not None else "" for cell in row] for row in rows[1:]
        ]

        # Sample data
        sample_rows = data_rows[:max_rows]

        return {
            "columns": header,
            "column_count": len(header),
            "total_rows": len(data_rows),
            "sample_rows": sample_rows,
            "sampled": len(data_rows) > max_rows,
            "sheet_name": sheet.title if use_openpyxl else "Sheet1",
        }
    except Exception as e:
        return {"error": f"Failed to parse XLSX: {e}"}


def read_docx_file(file_path: Path) -> Dict[str, Any]:
    """Read text content from a DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        Dictionary with content or error
    """
    try:
        try:
            from docx import Document
        except ImportError:
            return {
                "error": "python-docx is required for reading DOCX files. Install with: pip install python-docx"
            }

        doc = Document(str(file_path))

        # Extract all text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)

        content = "\n\n".join(full_text)

        return {
            "content": content,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "characters": len(content),
        }
    except Exception as e:
        return {"error": f"Failed to read DOCX file: {e}"}


def parse_qsf_file(qsf_content: str) -> Dict[str, Any]:
    """Parse a QSF (Qualtrics Survey Format) file and extract questions.

    Args:
        qsf_content: Raw QSF file content (JSON string)

    Returns:
        Dictionary with parsed survey structure including questions
    """

    def strip_html(text: str) -> str:
        """Remove HTML tags and decode entities."""
        if not text:
            return ""
        # Remove HTML tags
        text = re.sub(r"<[^>]+>", "", text)
        # Decode HTML entities
        text = unescape(text)
        # Clean up whitespace
        text = " ".join(text.split())
        return text

    try:
        data = json.loads(qsf_content)
    except json.JSONDecodeError as e:
        return {"error": f"Invalid JSON: {e}"}

    parsed = {
        "survey_name": data.get("SurveyEntry", {}).get("SurveyName", "Unknown"),
        "questions": [],
    }

    # Extract questions from SurveyElements
    if "SurveyElements" in data:
        for element in data["SurveyElements"]:
            if element.get("Element") == "SQ":  # Survey Question
                payload = element.get("Payload", {})

                question = {
                    "question_id": payload.get("QuestionID", ""),
                    "question_text": strip_html(payload.get("QuestionText", "")),
                    "question_type": payload.get("QuestionType", ""),
                    "selector": payload.get("Selector", ""),
                    "sub_selector": payload.get("SubSelector", ""),
                    "choices": {},
                    "answers": {},
                    "validation": payload.get("Validation", {}),
                }

                # Extract choices
                if "Choices" in payload and isinstance(payload["Choices"], dict):
                    for choice_id, choice_data in payload["Choices"].items():
                        if isinstance(choice_data, dict):
                            question["choices"][choice_id] = strip_html(
                                choice_data.get("Display", "")
                            )

                # Extract answers (for matrix questions)
                if "Answers" in payload and isinstance(payload["Answers"], dict):
                    for ans_id, ans_data in payload["Answers"].items():
                        if isinstance(ans_data, dict):
                            question["answers"][ans_id] = strip_html(
                                ans_data.get("Display", "")
                            )

                parsed["questions"].append(question)

    return parsed


# Check for Anthropic API key
def check_api_key():
    """Check if Anthropic API key is available."""
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        console.print("\n[bold red]WARNING: No Anthropic API key found![/bold red]")
        console.print(
            "\n[yellow]The SurveyAssistant requires an Anthropic API key to function.[/yellow]"
        )
        console.print("\nPlease set your API key in one of these ways:")
        console.print("  1. Create a .env file with: ANTHROPIC_API_KEY=your_key_here")
        console.print(
            "  2. Export it in your shell: export ANTHROPIC_API_KEY=your_key_here"
        )
        console.print("  3. Set it in your environment variables")
        console.print(
            "\nGet your API key from: [link]https://console.anthropic.com/[/link]\n"
        )
        return False
    return True


class SurveyAssistant:
    """Generates EDSL surveys from uploaded documents using Claude.

    Args:
        *file_paths: File paths or directory paths to process. If a directory is provided,
                    all files in that directory will be loaded (non-recursive).
        spec_file: Optional path to custom EDSL specification file.
                  If None (default), uses the bundled EDSL_SPECIFICATION.md
                  from the package data, which works regardless of current directory.
        verbose_logs: If True, prints detailed debug logs during generation
        max_turns: Maximum number of turns for Claude's agent loop (default: 30)

    Examples:
        >>> # Load specific files
        >>> sa = SurveyAssistant('file1.docx', 'file2.xlsx')

        >>> # Load all files from a directory
        >>> sa = SurveyAssistant('examples/pew')

        >>> # Mix files and directories
        >>> sa = SurveyAssistant('file1.docx', 'examples/pew', 'file2.csv')
    """

    def __init__(
        self,
        *file_paths,
        spec_file: Optional[Path] = None,
        verbose_logs: bool = False,
        max_turns: int = 30,
    ):
        # Lazy-load claude_agent_sdk when actually instantiating
        _check_claude_sdk()

        # Check for API key
        if not check_api_key():
            raise ValueError(
                "Anthropic API key is required. Please set ANTHROPIC_API_KEY environment variable."
            )

        self.spec_file = spec_file
        self.edsl_spec = self._load_specification()
        self.output_dir = Path("edsl_output")
        self.output_dir.mkdir(exist_ok=True)
        self._survey = None  # Will store the generated survey
        self._results = None  # Will store the generated results
        self._console = Console()  # Console for progress updates
        self.max_turns = max_turns  # Maximum turns for Claude agent

        # Setup logging first (before using it)
        self.verbose_logs = verbose_logs
        self.logger = logging.getLogger(f"SurveyAssistant.{id(self)}")
        if verbose_logs:
            self.logger.setLevel(logging.DEBUG)
            if not self.logger.handlers:
                handler = logging.StreamHandler()
                handler.setFormatter(logging.Formatter("[%(levelname)s] %(message)s"))
                self.logger.addHandler(handler)
        else:
            self.logger.setLevel(logging.WARNING)

        # Expand directories into file lists
        expanded_paths = []
        for fp in file_paths:
            path = Path(fp)
            if path.is_dir():
                # Add all files in the directory (non-recursive)
                dir_files = sorted([f for f in path.iterdir() if f.is_file()])
                expanded_paths.extend(dir_files)
                self.logger.info(
                    f"Expanded directory '{fp}' → {len(dir_files)} file(s)"
                )
                if self.verbose_logs:
                    for f in dir_files:
                        self.logger.debug(f"  - {f.name}")
            else:
                expanded_paths.append(path)

        self.file_paths = expanded_paths

        # Initialize helper modules
        self.column_mapper = ColumnMapper(self._console, self.logger)
        self.survey_generator = SurveyGenerator(
            self.edsl_spec, self.output_dir, self._console, self.logger, self.max_turns
        )

    def _load_specification(self) -> str:
        """Load EDSL specification from markdown file.

        If spec_file is provided, loads from that path.
        Otherwise, loads from the package data (bundled EDSL_SPECIFICATION.md).
        """
        if self.spec_file is not None:
            # User provided a custom spec file
            if not self.spec_file.exists():
                console.print(
                    f"[yellow]Warning: {self.spec_file} not found. Using bundled EDSL specification.[/yellow]"
                )
            else:
                with open(self.spec_file, "r") as f:
                    return f.read()

        # Load from package data
        try:
            package_files = files("edsl.assistants")
            spec_path = package_files / "EDSL_SPECIFICATION.md"
            return spec_path.read_text(encoding="utf-8")
        except Exception as e:
            console.print(
                f"[yellow]Warning: Could not load bundled EDSL specification: {e}[/yellow]"
            )
            return ""

    def _create_status_display(
        self, emoji: str, message: str, spinner: bool = False, style: str = "cyan"
    ):
        """Create a status display with optional spinner for Live display.

        Args:
            emoji: Emoji to display before the message
            message: The message to display
            spinner: Whether to show a spinner
            style: Rich style to apply

        Returns:
            Columns with spinner and text if spinner=True, otherwise just Text
        """
        text = Text(f"{emoji} {message}", style=style)
        if spinner:
            return Columns([Spinner("dots", style=style), text], padding=(0, 1))
        return text

    @property
    def survey(self) -> Optional[Survey]:
        """Get the generated Survey object. Generates it if not already created.

        Returns:
            The Survey object, generating it if necessary. Returns None if generation fails.
        """
        if self._survey is None:
            result = self.generate_survey()
            if not result.get("success", False):
                return None
        return self._survey

    @property
    def results(self):
        """Get the last generated Results object.

        Returns:
            The Results object from the last call to generate_results(), or None if not yet generated.

        Example:
            >>> sa = SurveyAssistant('data.dta')
            >>> sa.generate_survey()
            >>> sa.generate_results()  # Forgot to capture return value
            >>> results = sa.results  # Can still access it here
        """
        return self._results

    def generate_survey(
        self,
        file_paths: Optional[List[Path]] = None,
        verbose: bool = False,
        debug: bool = False,
    ) -> Dict[str, Any]:
        """Process multiple files and generate EDSL survey.

        Args:
            file_paths: Optional list of file paths. If not provided, uses the file paths
                       passed to the constructor.
            verbose: If True, returns detailed results dictionary. If False, returns simplified summary.
            debug: If True, streams Claude's responses to stdout in real-time for debugging.

        Returns:
            Dictionary with generation results and status. If verbose=False, returns simplified summary.
        """
        # The async function handles KeyboardInterrupt internally
        return anyio.run(self._generate_survey_async, file_paths, verbose, debug)

    async def _generate_survey_async(
        self,
        file_paths: Optional[List[Path]] = None,
        verbose: bool = False,
        debug: bool = False,
    ) -> Dict[str, Any]:
        """Async implementation of survey generation.

        Args:
            file_paths: Optional list of file paths to process
            verbose: If True, returns detailed results dictionary
            debug: If True, streams Claude's responses to stdout
        """
        # Use provided file_paths or fall back to instance file_paths
        if file_paths is None:
            file_paths = self.file_paths

        if not file_paths:
            return {
                "error": "No files provided. Pass file paths to constructor or to generate_survey()"
            }

        # Validate all files exist BEFORE starting any processing
        missing_files = []
        for file_path in file_paths:
            if not file_path.exists():
                missing_files.append(str(file_path))

        if missing_files:
            error_msg = f"File(s) not found: {', '.join(missing_files)}"
            self._console.print(f"\n[red]Error:[/red] {error_msg}")
            raise FileNotFoundError(error_msg)

        # Check if we have a structured data file (DTA, SAV) - use Conjure directly
        structured_formats = [".dta", ".sav"]
        structured_files = [
            f for f in file_paths if f.suffix.lower() in structured_formats
        ]

        if len(structured_files) >= 1:
            # We have at least one structured data file - use Conjure for cleaner parsing
            # If multiple, use the first one
            data_file = structured_files[0]

            if len(file_paths) > 1:
                other_files = [f for f in file_paths if f != data_file]
                self._console.print(
                    f"\n[bold cyan]📊 Detected structured data file: {data_file.name}[/bold cyan]"
                )
                self._console.print(
                    f"[dim]Using EDSL Conjure for native parsing (other files will be ignored)[/dim]"
                )
                if other_files:
                    self._console.print(
                        f"[dim]Additional files found (for reference): {', '.join(f.name for f in other_files[:3])}{'...' if len(other_files) > 3 else ''}[/dim]\n"
                    )
            else:
                self._console.print(
                    f"\n[bold cyan]📊 Detected structured data file: {data_file.name}[/bold cyan]"
                )
                self._console.print(
                    f"[dim]Using EDSL Conjure for native parsing (faster & more reliable)[/dim]\n"
                )

            try:
                from edsl.conjure import Conjure

                self._console.print(
                    "[yellow]⏳[/yellow] Loading data file with Conjure..."
                )
                conjure_data = Conjure(str(data_file))

                self._console.print(
                    "[yellow]⏳[/yellow] Generating survey from data structure..."
                )
                survey = conjure_data.to_survey(verbose=debug)

                # Save the survey
                self._survey = survey
                output_file = self.output_dir / "survey.json.gz"
                survey.save(str(output_file))

                # Also save as YAML for human readability
                yaml_file = self.output_dir / "survey.yaml"
                with open(yaml_file, "w") as f:
                    yaml.dump(survey.to_dict(), f, default_flow_style=False)

                self._console.print(f"\n[green]✅ Survey generation complete![/green]")
                self._console.print(f"[dim]Saved to {output_file}[/dim]")
                self._console.print(
                    f"[green]✓[/green] Generated survey with {len(survey.questions)} questions"
                )
                self._console.print(f"[dim]Output: {self.output_dir}/[/dim]\n")

                return {
                    "success": True,
                    "output_dir": str(self.output_dir),
                    "output_file": str(output_file),
                    "survey_yaml_file": str(yaml_file),
                    "questions_count": len(survey.questions),
                    "method": "conjure",
                }

            except ImportError:
                self._console.print(
                    "[yellow]⚠️  Conjure not available, falling back to Claude parsing[/yellow]\n"
                )
                # Fall through to Claude-based parsing
            except Exception as e:
                self._console.print(f"[yellow]⚠️  Conjure parsing failed: {e}[/yellow]")
                self._console.print(
                    "[dim]Falling back to Claude-based parsing...[/dim]\n"
                )
                # Fall through to Claude-based parsing

        # Initialize Live display for progress updates
        live = Live(Text(""), console=self._console, refresh_per_second=10)
        live.start()

        try:
            # Read all files
            file_contents = {}
            for i, file_path in enumerate(file_paths, 1):
                live.update(
                    self._create_status_display(
                        "📖",
                        f"Reading files... ({i}/{len(file_paths)}) {file_path.name}",
                        spinner=True,
                    )
                )

                # Try to read file directly first if it exists locally
                # Skip binary formats that need special parsing
                file_ext = file_path.suffix.lower().lstrip(".")
                binary_formats = ["dta", "sav", "xlsx", "xls", "docx", "doc", "pdf"]

                if file_path.exists() and file_ext not in binary_formats:
                    try:
                        with open(file_path, "r", encoding="utf-8") as f:
                            content_str = f.read()
                        file_contents[str(file_path)] = content_str
                        self.logger.debug(
                            f"Successfully read {file_path.name} directly: {len(content_str)} characters"
                        )
                        if debug:
                            self._console.print(
                                f"[green]✓[/green] Read {file_path.name} directly ({len(content_str)} chars)"
                            )
                    except Exception as e:
                        self.logger.warning(
                            f"Failed to read {file_path.name} directly: {e}, trying Claude..."
                        )
                        # Fall back to Claude if direct read fails
                        content = await self._read_file(file_path)
                        if "error" not in content:
                            file_contents[str(file_path)] = content["content"]
                            self.logger.debug(
                                f"Successfully read {file_path.name} via Claude: {len(content['content'])} characters"
                            )
                        else:
                            self.logger.warning(
                                f"Failed to read {file_path.name}: {content.get('error')}"
                            )
                elif file_path.exists() and file_ext in binary_formats:
                    # For binary formats, handle based on type
                    if file_ext in ["dta", "sav"]:
                        # Test Stata/SPSS file parsing
                        file_type = "stata" if file_ext == "dta" else "spss"
                        test_result = sample_stata_spss_file(
                            file_path, file_type, max_rows=1
                        )
                        if "error" in test_result:
                            live.stop()
                            error_msg = (
                                f"Cannot read {file_path.name}: {test_result['error']}"
                            )
                            self.logger.error(error_msg)
                            self._console.print(f"\n[red]Error:[/red] {error_msg}")
                            raise FileNotFoundError(error_msg)
                        # File is valid, use placeholder - will be parsed in _generate_survey_yaml
                        file_contents[
                            str(file_path)
                        ] = f"[Binary file: {file_path.name}]"
                        self.logger.debug(
                            f"Marked {file_path.name} as binary format for special parsing"
                        )
                        if debug:
                            self._console.print(
                                f"[green]✓[/green] Detected {file_path.name} as {file_ext.upper()} (binary format)"
                            )
                    elif file_ext in ["xlsx", "xls"]:
                        # Test Excel file parsing
                        test_result = sample_xlsx_file(file_path, max_rows=1)
                        if "error" in test_result:
                            live.stop()
                            error_msg = (
                                f"Cannot read {file_path.name}: {test_result['error']}"
                            )
                            self.logger.error(error_msg)
                            self._console.print(f"\n[red]Error:[/red] {error_msg}")
                            raise FileNotFoundError(error_msg)
                        # File is valid, use placeholder - will be parsed in _generate_survey_yaml
                        file_contents[
                            str(file_path)
                        ] = f"[Binary file: {file_path.name}]"
                        self.logger.debug(
                            f"Marked {file_path.name} as binary format for special parsing"
                        )
                        if debug:
                            self._console.print(
                                f"[green]✓[/green] Detected {file_path.name} as {file_ext.upper()} (binary format)"
                            )
                    elif file_ext == "docx":
                        # Read DOCX directly using python-docx
                        docx_result = read_docx_file(file_path)
                        if "error" in docx_result:
                            # Fall back to Claude if python-docx not available
                            if debug:
                                self._console.print(
                                    f"[yellow]→[/yellow] python-docx not available, using Claude for {file_path.name}"
                                )
                            content = await self.survey_generator.read_file(file_path)
                            if "error" not in content:
                                file_contents[str(file_path)] = content["content"]
                                if debug:
                                    self._console.print(
                                        f"[green]✓[/green] Read {file_path.name} via Claude ({len(content['content'])} chars)"
                                    )
                            else:
                                self.logger.warning(
                                    f"Failed to read {file_path.name}: {content.get('error')}"
                                )
                                if debug:
                                    self._console.print(
                                        f"[red]✗[/red] Failed to read {file_path.name}"
                                    )
                        else:
                            file_contents[str(file_path)] = docx_result["content"]
                            self.logger.debug(
                                f"Successfully read {file_path.name} directly: {docx_result['characters']} characters, "
                                f"{docx_result['paragraphs']} paragraphs, {docx_result['tables']} tables"
                            )
                            if debug:
                                self._console.print(
                                    f"[green]✓[/green] Read {file_path.name} directly ({docx_result['characters']} chars, "
                                    f"{docx_result['paragraphs']} paragraphs, {docx_result['tables']} tables)"
                                )
                    elif file_ext == "pdf":
                        # For PDF, use native Python parser with truncation
                        if debug:
                            self._console.print(
                                f"[yellow]→[/yellow] Reading {file_path.name} with native PDF parser"
                            )
                        pdf_result = read_pdf_file(
                            file_path, max_chars=50000, max_pages=50
                        )
                        if "error" not in pdf_result:
                            file_contents[str(file_path)] = pdf_result["content"]
                            truncation_note = (
                                " [TRUNCATED]"
                                if pdf_result.get("truncated", False)
                                else ""
                            )
                            self.logger.debug(
                                f"Successfully read {file_path.name}: {pdf_result['characters']} chars, "
                                f"{pdf_result['pages']}/{pdf_result['total_pages']} pages{truncation_note}"
                            )
                            if debug:
                                self._console.print(
                                    f"[green]✓[/green] Read {file_path.name} ({pdf_result['characters']} chars, "
                                    f"{pdf_result['pages']}/{pdf_result['total_pages']} pages{truncation_note})"
                                )
                        else:
                            # Fall back to Claude if native parser fails
                            if debug:
                                self._console.print(
                                    f"[yellow]→[/yellow] Native PDF parser failed, falling back to Claude for {file_path.name}"
                                )
                            content = await self._read_file(file_path)
                            if "error" not in content:
                                file_contents[str(file_path)] = content["content"]
                                if debug:
                                    self._console.print(
                                        f"[green]✓[/green] Read {file_path.name} via Claude ({len(content['content'])} chars)"
                                    )
                            else:
                                self.logger.warning(
                                    f"Failed to read {file_path.name}: {content.get('error')}"
                                )
                                if debug:
                                    self._console.print(
                                        f"[red]✗[/red] Failed to read {file_path.name}"
                                    )
                    elif file_ext == "doc":
                        # For DOC (old Word), use Claude's tools
                        if debug:
                            self._console.print(
                                f"[yellow]→[/yellow] Using Claude tools to read {file_path.name} (DOC)"
                            )
                        content = await self._read_file(file_path)
                        if "error" not in content:
                            file_contents[str(file_path)] = content["content"]
                            self.logger.debug(
                                f"Successfully read {file_path.name} via Claude: {len(content['content'])} characters"
                            )
                            if debug:
                                self._console.print(
                                    f"[green]✓[/green] Read {file_path.name} via Claude ({len(content['content'])} chars)"
                                )
                        else:
                            self.logger.warning(
                                f"Failed to read {file_path.name}: {content.get('error')}"
                            )
                            if debug:
                                self._console.print(
                                    f"[red]✗[/red] Failed to read {file_path.name}: {content.get('error')}"
                                )
                    else:
                        # Unknown binary format - try Claude's tools
                        if debug:
                            self._console.print(
                                f"[yellow]→[/yellow] Unknown binary format {file_ext.upper()}, using Claude tools"
                            )
                        content = await self._read_file(file_path)
                        if "error" not in content:
                            file_contents[str(file_path)] = content["content"]
                            self.logger.debug(
                                f"Successfully read {file_path.name} via Claude: {len(content['content'])} characters"
                            )
                        else:
                            self.logger.warning(
                                f"Failed to read {file_path.name}: {content.get('error')}"
                            )
                else:
                    # File doesn't exist locally, try Claude
                    content = await self._read_file(file_path)
                    if "error" not in content:
                        file_contents[str(file_path)] = content["content"]
                        self.logger.debug(
                            f"Successfully read {file_path.name}: {len(content['content'])} characters"
                        )
                    else:
                        self.logger.warning(
                            f"Failed to read {file_path.name}: {content.get('error')}"
                        )

            if not file_contents:
                live.stop()
                self.logger.error("No valid files could be read")
                return {"error": "No valid files could be read"}

            self.logger.debug(f"Total files read: {len(file_contents)}")

            # Extract data column names if tabular data files are present
            data_columns = await self.survey_generator.extract_data_columns(file_paths)
            if data_columns:
                self._console.print(
                    f"[dim]Found data file with {len(data_columns)} columns - will use for question naming[/dim]"
                )
                if debug:
                    self._console.print(
                        f"[dim]Columns: {', '.join(data_columns[:10])}{'...' if len(data_columns) > 10 else ''}[/dim]"
                    )

            # Check if we need to batch process (for large surveys)
            BATCH_SIZE = 100  # Process 100 questions at a time
            use_batching = data_columns and len(data_columns) > BATCH_SIZE

            if use_batching:
                self._console.print(
                    f"[yellow]⚠️  Large survey detected ({len(data_columns)} questions)[/yellow]"
                )
                self._console.print(
                    f"[dim]Processing in batches of {BATCH_SIZE} to avoid token limits...[/dim]\n"
                )

                # Split data_columns into batches
                batches = [
                    data_columns[i : i + BATCH_SIZE]
                    for i in range(0, len(data_columns), BATCH_SIZE)
                ]

                # For batching, filter out large files to reduce context per batch
                # Only keep the first 10,000 chars of each file
                filtered_contents = {}
                for filename, content in file_contents.items():
                    if len(content) > 10000:
                        filtered_contents[filename] = (
                            content[:10000]
                            + f"\n\n[TRUNCATED - showing first 10,000 of {len(content)} characters]"
                        )
                    else:
                        filtered_contents[filename] = content

                all_questions = []
                for batch_idx, batch_columns in enumerate(batches, 1):
                    if debug:
                        live.stop()
                        self._console.print(
                            f"\n[bold cyan]🤖 Batch {batch_idx}/{len(batches)}: Processing {len(batch_columns)} questions...[/bold cyan]"
                        )
                        self._console.print("[dim]" + "─" * 70 + "[/dim]")
                    else:
                        live.update(
                            self._create_status_display(
                                "🤖",
                                f"Generating questions (batch {batch_idx}/{len(batches)})...",
                                spinner=True,
                            )
                        )

                    batch_result = await self.survey_generator.generate_survey_yaml(
                        filtered_contents, data_columns=batch_columns, debug=debug
                    )

                    if "error" in batch_result:
                        if not debug:
                            live.stop()
                        self._console.print(
                            f"\n[red]Error in batch {batch_idx}:[/red] {batch_result['error']}"
                        )
                        # Continue with other batches even if one fails
                        continue

                    if "questions" in batch_result:
                        all_questions.extend(batch_result["questions"])
                        if not debug:
                            self._console.print(
                                f"[dim]Batch {batch_idx}/{len(batches)}: "
                                f"Generated {len(batch_result['questions'])} questions[/dim]"
                            )

                    if debug:
                        self._console.print("\n[dim]" + "─" * 70 + "[/dim]")
                        self._console.print(f"[dim]Batch {batch_idx} complete.[/dim]\n")
                        live.start()

                # Combine all batch results
                yaml_questions = {"questions": all_questions}

                if not debug:
                    live.stop()
                self._console.print(
                    f"\n[green]✓[/green] Generated {len(all_questions)} questions across {len(batches)} batches\n"
                )
                if not debug:
                    live.start()
            else:
                # Generate YAML for survey questions (single call for small surveys)
                if debug:
                    live.stop()
                    self._console.print(
                        "\n[bold cyan]🤖 Generating survey questions with Claude...[/bold cyan]"
                    )
                    self._console.print("[dim]" + "─" * 70 + "[/dim]")
                else:
                    live.update(
                        self._create_status_display(
                            "🤖",
                            "Generating survey questions with Claude...",
                            spinner=True,
                        )
                    )

                yaml_questions = await self.survey_generator.generate_survey_yaml(
                    file_contents, data_columns=data_columns, debug=debug
                )

                if debug:
                    self._console.print("\n[dim]" + "─" * 70 + "[/dim]")
                    self._console.print("[dim]Claude response complete.[/dim]\n")
                    live.start()

            if "error" in yaml_questions:
                if not debug:
                    live.stop()
                self.logger.error(f"YAML generation error: {yaml_questions['error']}")

                # Special handling for content filtering errors
                if "content filtering" in yaml_questions["error"].lower():
                    # Error message already printed by _generate_survey_yaml
                    pass
                else:
                    self._console.print(
                        f"\n[red]Error:[/red] {yaml_questions['error']}"
                    )
                    if not debug:
                        self._console.print(f"[dim]Debug files saved to:[/dim]")
                        self._console.print(
                            f"  [dim]- {self.output_dir}/last_prompt.txt[/dim]"
                        )
                        self._console.print(
                            f"  [dim]- {self.output_dir}/last_claude_response.txt[/dim]\n"
                        )
                return yaml_questions

            self.logger.debug(
                f"Generated {len(yaml_questions.get('questions', []))} questions from Claude"
            )
            if self.verbose_logs and yaml_questions.get("questions"):
                for q in yaml_questions["questions"]:
                    self.logger.debug(f"  - {q.get('name', 'unnamed')}")

            # Process each question individually
            questions = []
            errors = []
            question_results = []
            total_questions = len(yaml_questions["questions"])

            for idx, question_data in enumerate(yaml_questions["questions"], 1):
                question_name = question_data.get("name", "unnamed")
                yaml_content = question_data.get("yaml", "")

                live.update(
                    self._create_status_display(
                        "⚙️",
                        f"Processing questions... ({idx}/{total_questions}) {question_name}",
                        spinner=False,
                    )
                )

                # Save individual YAML file
                yaml_file = self.output_dir / f"{question_name}.yaml"
                with open(yaml_file, "w") as f:
                    f.write(yaml_content)

                # Try to instantiate the question
                result = self._create_question_from_yaml(yaml_content, question_name)

                if result["success"]:
                    questions.append(result["question"])
                    question_results.append(
                        {
                            "name": question_name,
                            "type": result["question"].__class__.__name__,
                            "status": "✓ Success",
                            "file": yaml_file.name,
                        }
                    )
                else:
                    errors.append(
                        {
                            "question_name": question_name,
                            "error": result["error"],
                            "yaml_file": yaml_file.name,
                        }
                    )
                    question_results.append(
                        {
                            "name": question_name,
                            "type": "Unknown",
                            "status": f"✗ Failed: {result['error']}",
                            "file": yaml_file.name,
                        }
                    )

            # Create survey if we have any valid questions
            survey = None
            survey_yaml = None

            if questions:
                try:
                    live.update(
                        self._create_status_display(
                            "📝", "Creating survey object...", spinner=False
                        )
                    )
                    survey = Survey(questions=questions)
                    self._survey = survey  # Store the survey in the instance

                    # Save survey YAML
                    live.update(
                        self._create_status_display(
                            "💾", "Saving survey files...", spinner=False
                        )
                    )
                    survey_yaml = survey.to_yaml()
                    survey_yaml_file = self.output_dir / "survey.yaml"
                    with open(survey_yaml_file, "w") as f:
                        f.write(survey_yaml)

                    # Save survey as JSON.gz (suppress the "Saved to..." message)
                    output_file = self.output_dir / "survey.json.gz"
                    with SuppressOutput():
                        survey.save(str(output_file))

                except Exception as e:
                    live.stop()
                    self._console.print(f"[red]Error creating survey: {e}[/red]")
                    return {
                        "error": f"Failed to create survey: {e}",
                        "questions_processed": question_results,
                        "errors": errors,
                    }

            live.update(
                self._create_status_display(
                    "✅",
                    "Survey generation complete!",
                    spinner=False,
                    style="green bold",
                )
            )
            live.stop()

            # Print a concise summary
            self._console.print(
                f"\n[green]✓[/green] Generated survey with {len(questions)}/{len(questions) + len(errors)} questions"
            )
            if errors:
                self._console.print(
                    f"[yellow]⚠[/yellow]  {len(errors)} question(s) failed to process:"
                )
                for error in errors:
                    # Extract just the first line of the error message for brevity
                    error_msg = error["error"].split("\n")[0]
                    if error_msg.startswith("Failed to create question: "):
                        error_msg = error_msg.replace("Failed to create question: ", "")
                    self._console.print(
                        f"   [red]✗[/red] [bold]{error['question_name']}[/bold]: {error_msg}"
                    )
                    self._console.print(
                        f"      [dim]Check: {self.output_dir}/{error['yaml_file']}[/dim]"
                    )
            self._console.print(f"[dim]Output: {self.output_dir}/[/dim]\n")

            # Build detailed results
            detailed_results = {
                "success": len(questions) > 0,
                "output_dir": str(self.output_dir),
                "output_file": str(self.output_dir / "survey.json.gz")
                if survey
                else None,
                "survey_yaml_file": str(self.output_dir / "survey.yaml")
                if survey
                else None,
                "questions_processed": question_results,
                "errors": errors,
                "survey_summary": {
                    "total_questions": len(questions) + len(errors),
                    "successful_questions": len(questions),
                    "failed_questions": len(errors),
                    "question_types": list(
                        set([q.__class__.__name__ for q in questions])
                    )
                    if questions
                    else [],
                    "question_names": [q.question_name for q in questions]
                    if questions
                    else [],
                },
            }

            # Return simplified summary if not verbose
            if not verbose:
                return {
                    "success": len(questions) > 0,
                    "total_questions": len(questions) + len(errors),
                    "successful_questions": len(questions),
                    "failed_questions": len(errors),
                    "output_dir": str(self.output_dir),
                    "output_file": str(self.output_dir / "survey.json.gz")
                    if survey
                    else None,
                }

            return detailed_results
        except KeyboardInterrupt:
            live.stop()
            self._console.print(
                "\n[yellow]⚠ Survey generation interrupted by user[/yellow]"
            )
            return {"error": "Interrupted by user", "success": False}
        except Exception as e:
            live.stop()
            self._console.print(f"[red]Error during survey generation: {e}[/red]")
            raise
        finally:
            # Always ensure Live display is stopped
            try:
                live.stop()
            except:
                pass  # Already stopped

    async def _extract_data_columns(
        self, file_paths: List[Path]
    ) -> Optional[List[str]]:
        """Extract column names from tabular data files if present.

        Args:
            file_paths: List of all file paths

        Returns:
            List of column names from the first tabular data file, or None if none found
        """
        tabular_formats = ["csv", "xlsx", "xls", "dta", "sav"]

        for file_path in file_paths:
            file_ext = file_path.suffix.lower().lstrip(".")
            if file_ext not in tabular_formats or not file_path.exists():
                continue

            try:
                if file_ext == "csv":
                    import pandas as pd

                    df = pd.read_csv(file_path, nrows=1)
                    return df.columns.tolist()
                elif file_ext in ["xlsx", "xls"]:
                    import pandas as pd

                    df = pd.read_excel(file_path, nrows=1)
                    return df.columns.tolist()
                elif file_ext == "dta":
                    try:
                        import pyreadstat

                        df, meta = pyreadstat.read_dta(str(file_path), row_limit=1)
                        return df.columns.tolist()
                    except ImportError:
                        continue
                elif file_ext == "sav":
                    try:
                        import pyreadstat

                        df, meta = pyreadstat.read_sav(str(file_path), row_limit=1)
                        return df.columns.tolist()
                    except ImportError:
                        continue
            except Exception as e:
                self.logger.warning(
                    f"Failed to extract columns from {file_path.name}: {e}"
                )
                continue

        return None

    async def _read_file(self, file_path: Path) -> Dict[str, str]:
        """Read content from any file type."""
        from claude_agent_sdk import (
            AssistantMessage,
            ClaudeAgentOptions,
            TextBlock,
            query,
        )

        options = ClaudeAgentOptions(
            allowed_tools=["Read"],
            system_prompt="You are a file reading assistant. Extract all text content from any file type (PDF, DOC, TXT, CSV, JSON, QSF, etc.).",
            max_turns=5,  # Give Claude enough turns to read large files
        )

        prompt = f"Read the file at {file_path} and extract all content. For structured data (CSV, JSON, QSF), preserve the structure."

        content = ""
        async for message in query(prompt=prompt, options=options):
            if isinstance(message, AssistantMessage):
                for block in message.content:
                    if isinstance(block, TextBlock):
                        content += block.text

        return (
            {"content": content}
            if content
            else {"error": f"Could not read {file_path}"}
        )

    async def _generate_survey_yaml(
        self,
        file_contents: Dict[str, str],
        data_columns: Optional[List[str]] = None,
        debug: bool = False,
    ) -> Dict[str, Any]:
        """Generate YAML representation of EDSL survey from document contents.

        Args:
            file_contents: Dictionary mapping filenames to their contents
            data_columns: Optional list of column names from response data file
            debug: If True, streams Claude's response to stdout in real-time
        """
        from claude_agent_sdk import (
            AssistantMessage,
            ClaudeAgentOptions,
            ResultMessage,
            TextBlock,
            ToolUseBlock,
            UserMessage,
            query,
        )

        # Build data columns section for the prompt if available
        data_columns_section = ""
        if data_columns:
            # Show actual column → question_name conversions
            examples = []
            for col in data_columns[:5]:  # Show first 5 as examples
                sanitized = sanitize_column_name(col)
                examples.append(f'  "{col}" → question_name: "{sanitized}"')

            examples_text = "\n".join(examples)
            if len(data_columns) > 5:
                examples_text += f"\n  ... and {len(data_columns) - 5} more columns"

            data_columns_section = f"""
**CRITICAL - DATA FILE COLUMNS PROVIDED:**
A response data file is available with EXACTLY these {len(data_columns)} column names:
{', '.join(data_columns)}

**CRITICAL REQUIREMENT - EXACT COLUMN MAPPING:**
Create EXACTLY ONE question for EACH column above. Do NOT infer, skip, or create additional columns.

For the question_name, convert the EXACT column name to a valid Python identifier by:
1. Converting to lowercase
2. Replacing spaces with underscores
3. Replacing dashes/hyphens with underscores  
4. Removing other special characters
5. Removing consecutive underscores and leading/trailing underscores

**Here are the ACTUAL conversions for the columns in THIS data file:**
{examples_text}

**IMPORTANT:** Preserve ALL parts of the column name including numbers! 
If the data has "RM1" and "RM3" but no "RM2", create questions for RM1 and RM3 only - do NOT create RM2!

Create exactly {len(data_columns)} questions, one for each column listed above.
"""

        system_prompt = f"""You are an expert in EDSL (Expected Parrot Domain Specific Language) for creating surveys.

**IMPORTANT CONTEXT:**
You are performing ACADEMIC RESEARCH ANALYSIS by parsing and converting existing survey instruments 
into a structured format. You are NOT conducting these surveys yourself or asking these questions 
to real people. Your role is purely technical: analyze the survey structure and convert it to EDSL 
format for research purposes. The content may include sensitive topics (medical, demographic, etc.) 
but you are only documenting what exists in the research instrument, not creating new content.

{self.edsl_spec}

Your task is to:
1. Analyze the provided documents (including CSV/TSV/Excel data, JSON/QSF survey files, Stata (.dta), SPSS (.sav), etc.)
2. For QSF files: The file will be pre-parsed for you with structured question data including:
   - question_text: The question text (HTML stripped)
   - question_type: Qualtrics question type (MC, TE, Matrix, etc.)
   - selector: Question selector (SAVR, FORM, Likert, etc.)
   - choices: Dictionary of answer choices
   - answers: Dictionary of answers (for matrix questions)
3. For tabular data files (CSV/TSV/Excel/Stata/SPSS): Large files are sampled (first 20 rows) to show structure. 
   For Stata and SPSS files, variable labels and value labels are also provided when available.
4. Map Qualtrics question types to appropriate EDSL question types
5. Generate YAML representations of EDSL questions
6. Output each question's YAML separately
{data_columns_section}
CRITICAL: Use SHORT STRING question types, NOT class names!

Valid question_type values (use these EXACT strings):
- "multiple_choice" - for single selection questions
- "checkbox" - for multiple selection questions  
- "free_text" - for open-ended text responses
- "linear_scale" - for numeric scales (e.g., 0-10)
- "numerical" - for numeric input
- "matrix" - for grid/matrix questions
- "rank" - for ranking items
- "top_k" - for selecting top K items
- "likert_five" - for 5-point Likert scales
- "yes_no" - for binary yes/no questions

WRONG: question_type: QuestionMultipleChoice
RIGHT: question_type: multiple_choice

WRONG: question_type: QuestionFreeText  
RIGHT: question_type: free_text

Important rules:
- Each question must have: question_name, question_text, question_type
- question_name must be a valid Python identifier (lowercase, underscores, no spaces)
- question_type MUST be one of the short string values listed above (lowercase with underscores)
- Output each question in a separate YAML block

Common Qualtrics to EDSL mappings:
- MC (Multiple Choice) with SAVR selector → "multiple_choice"
- MC with MAVR/MAHR selector → "checkbox"
- TE (Text Entry) with FORM/SL selector → "free_text"
- Matrix with Likert → "matrix"
- Slider → "linear_scale" or "numerical"
- Rank → "rank" or "top_k"

Format your response as:
QUESTION_NAME: [the question_name]
```yaml
question_name: [name]
question_text: [text]
question_type: [short_string_type]
# ... other fields as needed
```

EXAMPLE - Multiple choice question:
QUESTION_NAME: favorite_color
```yaml
question_name: favorite_color
question_text: What is your favorite color?
question_type: multiple_choice
question_options:
  - Red
  - Blue
  - Green
```

EXAMPLE - Linear scale question:
QUESTION_NAME: satisfaction
```yaml
question_name: satisfaction
question_text: How satisfied are you?
question_type: linear_scale
question_options: [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, 10]
option_labels:
  0: Not at all satisfied
  10: Completely satisfied
```

EXAMPLE - Free text question:
QUESTION_NAME: feedback
```yaml
question_name: feedback
question_text: Please provide any additional comments
question_type: free_text
```

Repeat for each question. ALWAYS use the short string question types (multiple_choice, free_text, linear_scale, etc.)!

NOTE: The EDSL spec above shows Python class names (QuestionMultipleChoice, QuestionFreeText, etc.) but you MUST use the short string equivalents in YAML (multiple_choice, free_text, etc.).

You do not need to use any tools - all necessary data is provided in the prompt above."""

        options = ClaudeAgentOptions(
            system_prompt=system_prompt, max_turns=self.max_turns
        )

        # Build the prompt with all file contents
        prompt = (
            "Please analyze these existing research survey documents and generate YAML for an EDSL survey.\n\n"
            "CONTEXT: You are parsing and converting an existing survey instrument that has already been "
            "used in research. You are NOT creating new questions or conducting a survey - you are simply "
            "converting the technical format of existing questions for analysis purposes.\n\n"
            "Documents to analyze:\n\n"
        )
        for filename, content in file_contents.items():
            # Check if this is a binary file placeholder
            if content.startswith("[Binary file:"):
                # This is a binary file - parse it now
                file_path = Path(filename)
                file_ext = file_path.suffix.lower().lstrip(".")

                # Process based on type
                if file_ext == "dta":
                    stata_sample = sample_stata_spss_file(
                        file_path, "stata", max_rows=20
                    )
                    if "error" in stata_sample:
                        prompt += f"=== {file_path.name} (Stata - {stata_sample['error']}) ===\n\n"
                    else:
                        prompt += f"=== {file_path.name} (Stata - Sampled) ===\n"
                        prompt += f"Variables: {stata_sample['variable_count']} | Total Rows: {stata_sample['total_rows']}\n"
                        prompt += (
                            f"Variable Names: {', '.join(stata_sample['columns'])}\n"
                        )

                        if stata_sample["column_labels"]:
                            prompt += f"\nVariable Labels:\n"
                            for var, label in list(
                                stata_sample["column_labels"].items()
                            )[:10]:
                                prompt += f"  {var}: {label}\n"

                        prompt += "\n"
                        if stata_sample["sampled"]:
                            prompt += f"Sample Data (first 20 of {stata_sample['total_rows']} rows):\n"
                        else:
                            prompt += (
                                f"Complete Data ({stata_sample['total_rows']} rows):\n"
                            )

                        stata_sample_text = ",".join(stata_sample["columns"]) + "\n"
                        for row in stata_sample["sample_rows"]:
                            stata_sample_text += (
                                ",".join(str(cell) for cell in row) + "\n"
                            )

                        prompt += stata_sample_text + "\n"

                        if debug:
                            self._console.print(
                                f"[green]✓[/green] Sampled Stata: {stata_sample['variable_count']} variables, {stata_sample['total_rows']} total rows"
                            )
                    continue

                elif file_ext == "sav":
                    spss_sample = sample_stata_spss_file(file_path, "spss", max_rows=20)
                    if "error" in spss_sample:
                        prompt += f"=== {file_path.name} (SPSS - {spss_sample['error']}) ===\n\n"
                    else:
                        prompt += f"=== {file_path.name} (SPSS - Sampled) ===\n"
                        prompt += f"Variables: {spss_sample['variable_count']} | Total Rows: {spss_sample['total_rows']}\n"
                        prompt += (
                            f"Variable Names: {', '.join(spss_sample['columns'])}\n"
                        )

                        if spss_sample["column_labels"]:
                            prompt += f"\nVariable Labels:\n"
                            for var, label in list(
                                spss_sample["column_labels"].items()
                            )[:10]:
                                prompt += f"  {var}: {label}\n"

                        prompt += "\n"
                        if spss_sample["sampled"]:
                            prompt += f"Sample Data (first 20 of {spss_sample['total_rows']} rows):\n"
                        else:
                            prompt += (
                                f"Complete Data ({spss_sample['total_rows']} rows):\n"
                            )

                        spss_sample_text = ",".join(spss_sample["columns"]) + "\n"
                        for row in spss_sample["sample_rows"]:
                            spss_sample_text += (
                                ",".join(str(cell) for cell in row) + "\n"
                            )

                        prompt += spss_sample_text + "\n"

                        if debug:
                            self._console.print(
                                f"[green]✓[/green] Sampled SPSS: {spss_sample['variable_count']} variables, {spss_sample['total_rows']} total rows"
                            )
                    continue

                elif file_ext in ["xlsx", "xls"]:
                    xlsx_sample = sample_xlsx_file(file_path, max_rows=20)
                    if "error" in xlsx_sample:
                        prompt += f"=== {file_path.name} (Excel - {xlsx_sample['error']}) ===\n"
                        prompt += "Could not parse Excel file. Please convert to CSV or provide in another format.\n\n"
                    else:
                        prompt += f"=== {file_path.name} (Excel - Sampled) ===\n"
                        prompt += f"Sheet: {xlsx_sample.get('sheet_name', 'Unknown')}\n"
                        prompt += f"Columns: {xlsx_sample['column_count']} | Total Rows: {xlsx_sample['total_rows']}\n"
                        prompt += (
                            f"Column Headers: {', '.join(xlsx_sample['columns'])}\n\n"
                        )

                        if xlsx_sample["sampled"]:
                            prompt += f"Sample Data (first 20 of {xlsx_sample['total_rows']} rows):\n"
                        else:
                            prompt += (
                                f"Complete Data ({xlsx_sample['total_rows']} rows):\n"
                            )

                        xlsx_sample_text = ",".join(xlsx_sample["columns"]) + "\n"
                        for row in xlsx_sample["sample_rows"]:
                            xlsx_sample_text += (
                                ",".join(str(cell) for cell in row) + "\n"
                            )

                        prompt += xlsx_sample_text + "\n"

                        if debug:
                            self._console.print(
                                f"[green]✓[/green] Sampled Excel: {xlsx_sample['column_count']} columns, {xlsx_sample['total_rows']} total rows"
                            )
                    continue

            # Check if file appears to be QSF
            file_ext = filename.lower().split(".")[-1] if "." in filename else ""

            if file_ext == "qsf":
                # Parse QSF file and provide structured data
                parsed_qsf = parse_qsf_file(content)

                if "error" in parsed_qsf:
                    # Parsing failed, send raw content to Claude
                    self.logger.warning(f"QSF parsing failed: {parsed_qsf['error']}")
                    if debug:
                        self._console.print(
                            f"[yellow]⚠[/yellow] QSF parsing failed: {parsed_qsf['error']}"
                        )
                        self._console.print(
                            f"[dim]Sending raw file content to Claude instead[/dim]"
                        )

                    prompt += f"=== {filename} (Qualtrics QSF - Raw JSON, parsing failed) ===\n"
                    prompt += f"Note: Automatic parsing failed with error: {parsed_qsf['error']}\n"
                    prompt += f"Please parse this QSF JSON manually to extract survey questions.\n\n"
                    prompt += f"{content}\n\n"
                else:
                    prompt += f"=== {filename} (Qualtrics QSF - Pre-parsed) ===\n"
                    prompt += f"Survey Name: {parsed_qsf['survey_name']}\n"
                    prompt += f"Total Questions: {len(parsed_qsf['questions'])}\n\n"
                    prompt += "Parsed Questions:\n"
                    prompt += json.dumps(parsed_qsf["questions"], indent=2)
                    prompt += "\n\n"

                    if debug:
                        self._console.print(
                            f"[green]✓[/green] Pre-parsed QSF: {len(parsed_qsf['questions'])} questions found"
                        )

            elif file_ext == "csv":
                # Sample CSV file for large files
                csv_sample = sample_csv_file(content, max_rows=20, delimiter=",")

                if "error" in csv_sample:
                    # Sampling failed, send raw content
                    prompt += (
                        f"=== {filename} (CSV - sampling failed) ===\n{content}\n\n"
                    )
                else:
                    prompt += f"=== {filename} (CSV - Sampled) ===\n"
                    prompt += f"Columns: {csv_sample['column_count']} | Total Rows: {csv_sample['total_rows']}\n"
                    prompt += f"Column Headers: {', '.join(csv_sample['columns'])}\n\n"

                    if csv_sample["sampled"]:
                        prompt += f"Sample Data (first 20 of {csv_sample['total_rows']} rows):\n"
                    else:
                        prompt += f"Complete Data ({csv_sample['total_rows']} rows):\n"

                    # Reconstruct CSV for sample
                    csv_sample_text = ",".join(csv_sample["columns"]) + "\n"
                    for row in csv_sample["sample_rows"]:
                        csv_sample_text += ",".join(str(cell) for cell in row) + "\n"

                    prompt += csv_sample_text + "\n"

                    if debug:
                        self._console.print(
                            f"[green]✓[/green] Sampled CSV: {csv_sample['column_count']} columns, {csv_sample['total_rows']} total rows"
                        )

            elif file_ext == "tsv":
                # Sample TSV file for large files
                tsv_sample = sample_csv_file(content, max_rows=20, delimiter="\t")

                if "error" in tsv_sample:
                    # Sampling failed, send raw content
                    prompt += (
                        f"=== {filename} (TSV - sampling failed) ===\n{content}\n\n"
                    )
                else:
                    prompt += f"=== {filename} (TSV - Sampled) ===\n"
                    prompt += f"Columns: {tsv_sample['column_count']} | Total Rows: {tsv_sample['total_rows']}\n"
                    prompt += f"Column Headers: {', '.join(tsv_sample['columns'])}\n\n"

                    if tsv_sample["sampled"]:
                        prompt += f"Sample Data (first 20 of {tsv_sample['total_rows']} rows):\n"
                    else:
                        prompt += f"Complete Data ({tsv_sample['total_rows']} rows):\n"

                    # Reconstruct TSV for sample
                    tsv_sample_text = "\t".join(tsv_sample["columns"]) + "\n"
                    for row in tsv_sample["sample_rows"]:
                        tsv_sample_text += "\t".join(str(cell) for cell in row) + "\n"

                    prompt += tsv_sample_text + "\n"

                    if debug:
                        self._console.print(
                            f"[green]✓[/green] Sampled TSV: {tsv_sample['column_count']} columns, {tsv_sample['total_rows']} total rows"
                        )

            elif file_ext in ["xlsx", "xls"]:
                # Sample Excel file
                xlsx_sample = sample_xlsx_file(Path(filename), max_rows=20)

                if "error" in xlsx_sample:
                    prompt += f"=== {filename} (Excel - {xlsx_sample['error']}) ===\n"
                    prompt += "Could not parse Excel file. Please convert to CSV or provide in another format.\n\n"
                else:
                    prompt += f"=== {filename} (Excel - Sampled) ===\n"
                    prompt += f"Sheet: {xlsx_sample.get('sheet_name', 'Unknown')}\n"
                    prompt += f"Columns: {xlsx_sample['column_count']} | Total Rows: {xlsx_sample['total_rows']}\n"
                    prompt += f"Column Headers: {', '.join(xlsx_sample['columns'])}\n\n"

                    if xlsx_sample["sampled"]:
                        prompt += f"Sample Data (first 20 of {xlsx_sample['total_rows']} rows):\n"
                    else:
                        prompt += f"Complete Data ({xlsx_sample['total_rows']} rows):\n"

                    # Reconstruct as CSV for sample
                    xlsx_sample_text = ",".join(xlsx_sample["columns"]) + "\n"
                    for row in xlsx_sample["sample_rows"]:
                        xlsx_sample_text += ",".join(str(cell) for cell in row) + "\n"

                    prompt += xlsx_sample_text + "\n"

                    if debug:
                        self._console.print(
                            f"[green]✓[/green] Sampled Excel: {xlsx_sample['column_count']} columns, {xlsx_sample['total_rows']} total rows"
                        )

            elif file_ext == "dta":
                # Sample Stata file
                stata_sample = sample_stata_spss_file(
                    Path(filename), "stata", max_rows=20
                )

                if "error" in stata_sample:
                    prompt += (
                        f"=== {filename} (Stata - {stata_sample['error']}) ===\n\n"
                    )
                else:
                    prompt += f"=== {filename} (Stata - Sampled) ===\n"
                    prompt += f"Variables: {stata_sample['variable_count']} | Total Rows: {stata_sample['total_rows']}\n"
                    prompt += f"Variable Names: {', '.join(stata_sample['columns'])}\n"

                    # Show variable labels if available
                    if stata_sample["column_labels"]:
                        prompt += f"\nVariable Labels:\n"
                        for var, label in list(stata_sample["column_labels"].items())[
                            :10
                        ]:
                            prompt += f"  {var}: {label}\n"

                    prompt += "\n"
                    if stata_sample["sampled"]:
                        prompt += f"Sample Data (first 20 of {stata_sample['total_rows']} rows):\n"
                    else:
                        prompt += (
                            f"Complete Data ({stata_sample['total_rows']} rows):\n"
                        )

                    # Reconstruct as CSV for sample
                    stata_sample_text = ",".join(stata_sample["columns"]) + "\n"
                    for row in stata_sample["sample_rows"]:
                        stata_sample_text += ",".join(str(cell) for cell in row) + "\n"

                    prompt += stata_sample_text + "\n"

                    if debug:
                        self._console.print(
                            f"[green]✓[/green] Sampled Stata: {stata_sample['variable_count']} variables, {stata_sample['total_rows']} total rows"
                        )

            elif file_ext == "sav":
                # Sample SPSS file
                spss_sample = sample_stata_spss_file(
                    Path(filename), "spss", max_rows=20
                )

                if "error" in spss_sample:
                    prompt += f"=== {filename} (SPSS - {spss_sample['error']}) ===\n\n"
                else:
                    prompt += f"=== {filename} (SPSS - Sampled) ===\n"
                    prompt += f"Variables: {spss_sample['variable_count']} | Total Rows: {spss_sample['total_rows']}\n"
                    prompt += f"Variable Names: {', '.join(spss_sample['columns'])}\n"

                    # Show variable labels if available
                    if spss_sample["column_labels"]:
                        prompt += f"\nVariable Labels:\n"
                        for var, label in list(spss_sample["column_labels"].items())[
                            :10
                        ]:
                            prompt += f"  {var}: {label}\n"

                    prompt += "\n"
                    if spss_sample["sampled"]:
                        prompt += f"Sample Data (first 20 of {spss_sample['total_rows']} rows):\n"
                    else:
                        prompt += f"Complete Data ({spss_sample['total_rows']} rows):\n"

                    # Reconstruct as CSV for sample
                    spss_sample_text = ",".join(spss_sample["columns"]) + "\n"
                    for row in spss_sample["sample_rows"]:
                        spss_sample_text += ",".join(str(cell) for cell in row) + "\n"

                    prompt += spss_sample_text + "\n"

                    if debug:
                        self._console.print(
                            f"[green]✓[/green] Sampled SPSS: {spss_sample['variable_count']} variables, {spss_sample['total_rows']} total rows"
                        )

            elif file_ext == "json" or (
                content.strip().startswith("{") or content.strip().startswith("[")
            ):
                prompt += f"=== {filename} (JSON format) ===\n{content}\n\n"
            else:
                prompt += f"=== {filename} ===\n{content}\n\n"

        # Save prompt for debugging (only if not in debug mode)
        if not debug:
            prompt_file = self.output_dir / "last_prompt.txt"
            with open(prompt_file, "w") as f:
                f.write(prompt)
            self.logger.debug(f"Saved prompt to {prompt_file}")
        else:
            # In debug mode, show the prompt
            self._console.print("\n[bold yellow]📄 File Contents Summary:[/bold yellow]")
            self._console.print("[dim]" + "─" * 70 + "[/dim]")
            # Show info about each file
            for filename, content in file_contents.items():
                # Handle binary file placeholders
                if content.startswith("[Binary file:"):
                    file_path = Path(filename)
                    file_ext = file_path.suffix.lower().lstrip(".")

                    if file_ext == "dta":
                        stata_sample = sample_stata_spss_file(
                            file_path, "stata", max_rows=20
                        )
                        if "error" not in stata_sample:
                            self._console.print(
                                f"[cyan]{file_path.name}[/cyan] (Stata)"
                            )
                            self._console.print(
                                f"  Variables: {stata_sample['variable_count']}"
                            )
                            self._console.print(f"  Rows: {stata_sample['total_rows']}")
                            if stata_sample["sampled"]:
                                self._console.print(
                                    f"  [dim](showing first 20 rows)[/dim]"
                                )
                        else:
                            self._console.print(
                                f"[cyan]{file_path.name}[/cyan] (Stata - {stata_sample['error']})"
                            )
                    elif file_ext == "sav":
                        spss_sample = sample_stata_spss_file(
                            file_path, "spss", max_rows=20
                        )
                        if "error" not in spss_sample:
                            self._console.print(f"[cyan]{file_path.name}[/cyan] (SPSS)")
                            self._console.print(
                                f"  Variables: {spss_sample['variable_count']}"
                            )
                            self._console.print(f"  Rows: {spss_sample['total_rows']}")
                            if spss_sample["sampled"]:
                                self._console.print(
                                    f"  [dim](showing first 20 rows)[/dim]"
                                )
                        else:
                            self._console.print(
                                f"[cyan]{file_path.name}[/cyan] (SPSS - {spss_sample['error']})"
                            )
                    elif file_ext in ["xlsx", "xls"]:
                        xlsx_sample = sample_xlsx_file(file_path, max_rows=20)
                        if "error" not in xlsx_sample:
                            self._console.print(
                                f"[cyan]{file_path.name}[/cyan] (Excel)"
                            )
                            self._console.print(
                                f"  Sheet: {xlsx_sample.get('sheet_name', 'Unknown')}"
                            )
                            self._console.print(
                                f"  Columns: {xlsx_sample['column_count']}"
                            )
                            self._console.print(f"  Rows: {xlsx_sample['total_rows']}")
                            if xlsx_sample["sampled"]:
                                self._console.print(
                                    f"  [dim](showing first 20 rows)[/dim]"
                                )
                        else:
                            self._console.print(
                                f"[cyan]{file_path.name}[/cyan] (Excel - {xlsx_sample['error']})"
                            )
                    continue

                file_ext = filename.lower().split(".")[-1] if "." in filename else ""
                if file_ext == "qsf":
                    parsed_qsf = parse_qsf_file(content)
                    if "error" not in parsed_qsf:
                        self._console.print(f"[cyan]{filename}[/cyan] (QSF)")
                        self._console.print(f"  Survey: {parsed_qsf['survey_name']}")
                        self._console.print(
                            f"  Questions: {len(parsed_qsf['questions'])}"
                        )
                    else:
                        self._console.print(
                            f"[cyan]{filename}[/cyan] (QSF - parse error)"
                        )
                elif file_ext == "csv":
                    csv_sample = sample_csv_file(content, max_rows=20, delimiter=",")
                    if "error" not in csv_sample:
                        self._console.print(f"[cyan]{filename}[/cyan] (CSV)")
                        self._console.print(f"  Columns: {csv_sample['column_count']}")
                        self._console.print(f"  Rows: {csv_sample['total_rows']}")
                        if csv_sample["sampled"]:
                            self._console.print(f"  [dim](showing first 20 rows)[/dim]")
                    else:
                        self._console.print(
                            f"[cyan]{filename}[/cyan] (CSV - parse error)"
                        )
                elif file_ext == "tsv":
                    tsv_sample = sample_csv_file(content, max_rows=20, delimiter="\t")
                    if "error" not in tsv_sample:
                        self._console.print(f"[cyan]{filename}[/cyan] (TSV)")
                        self._console.print(f"  Columns: {tsv_sample['column_count']}")
                        self._console.print(f"  Rows: {tsv_sample['total_rows']}")
                        if tsv_sample["sampled"]:
                            self._console.print(f"  [dim](showing first 20 rows)[/dim]")
                    else:
                        self._console.print(
                            f"[cyan]{filename}[/cyan] (TSV - parse error)"
                        )
                elif file_ext in ["xlsx", "xls"]:
                    xlsx_sample = sample_xlsx_file(Path(filename), max_rows=20)
                    if "error" not in xlsx_sample:
                        self._console.print(f"[cyan]{filename}[/cyan] (Excel)")
                        self._console.print(
                            f"  Sheet: {xlsx_sample.get('sheet_name', 'Unknown')}"
                        )
                        self._console.print(f"  Columns: {xlsx_sample['column_count']}")
                        self._console.print(f"  Rows: {xlsx_sample['total_rows']}")
                        if xlsx_sample["sampled"]:
                            self._console.print(f"  [dim](showing first 20 rows)[/dim]")
                    else:
                        self._console.print(
                            f"[cyan]{filename}[/cyan] (Excel - {xlsx_sample['error']})"
                        )
                elif file_ext == "dta":
                    stata_sample = sample_stata_spss_file(
                        Path(filename), "stata", max_rows=20
                    )
                    if "error" not in stata_sample:
                        self._console.print(f"[cyan]{filename}[/cyan] (Stata)")
                        self._console.print(
                            f"  Variables: {stata_sample['variable_count']}"
                        )
                        self._console.print(f"  Rows: {stata_sample['total_rows']}")
                        if stata_sample["sampled"]:
                            self._console.print(f"  [dim](showing first 20 rows)[/dim]")
                    else:
                        self._console.print(
                            f"[cyan]{filename}[/cyan] (Stata - {stata_sample['error']})"
                        )
                elif file_ext == "sav":
                    spss_sample = sample_stata_spss_file(
                        Path(filename), "spss", max_rows=20
                    )
                    if "error" not in spss_sample:
                        self._console.print(f"[cyan]{filename}[/cyan] (SPSS)")
                        self._console.print(
                            f"  Variables: {spss_sample['variable_count']}"
                        )
                        self._console.print(f"  Rows: {spss_sample['total_rows']}")
                        if spss_sample["sampled"]:
                            self._console.print(f"  [dim](showing first 20 rows)[/dim]")
                    else:
                        self._console.print(
                            f"[cyan]{filename}[/cyan] (SPSS - {spss_sample['error']})"
                        )
                else:
                    self._console.print(
                        f"[cyan]{filename}[/cyan] ({len(content)} chars)"
                    )
                    preview = content[:300] + ("..." if len(content) > 300 else "")
                    self._console.print(f"[dim]{preview}[/dim]")
                self._console.print()
            self._console.print("[dim]" + "─" * 70 + "[/dim]\n")

        questions = []
        all_text = []  # Collect all text for debugging

        # Retry logic for API overload errors
        max_retries = 3
        retry_delay = 2  # seconds

        for attempt in range(max_retries):
            try:
                async for message in query(prompt=prompt, options=options):
                    if debug:
                        # Show message type for debugging
                        print(f"\n{'='*70}", flush=True)
                        print(f"[Message: {type(message).__name__}]", flush=True)

                    # Handle ResultMessage (tool results)
                    if isinstance(message, ResultMessage):
                        if debug:
                            print(
                                f"[Tool result being sent back to Claude]", flush=True
                            )
                        continue

                    # Handle UserMessage (system messages back to Claude)
                    if isinstance(message, UserMessage):
                        if debug:
                            print(f"[System message to Claude]", flush=True)
                        continue

                    if isinstance(message, AssistantMessage):
                        if debug and message.content:
                            print(
                                f"[{len(message.content)} content block(s)]", flush=True
                            )

                        for block in message.content:
                            if isinstance(block, ToolUseBlock):
                                # Show tool usage details
                                if debug:
                                    print(f"\n[TOOL USE]:", flush=True)
                                    print(f"  Tool: {block.name}", flush=True)
                                    print(f"  ID: {block.id}", flush=True)
                                    if hasattr(block, "input") and block.input:
                                        print(
                                            f"  Input: {json.dumps(block.input, indent=2)}",
                                            flush=True,
                                        )
                                self.logger.debug(f"Tool use: {block.name}")

                            elif isinstance(block, TextBlock):
                                # Parse the response to extract individual questions
                                text = block.text
                                all_text.append(text)

                                # Stream to stdout if debug mode
                                if debug:
                                    print(f"\n[TEXT]:", flush=True)
                                    print(text, flush=True)

                                self.logger.debug(
                                    f"Claude response chunk: {text[:200]}..."
                                )
                                lines = text.split("\n")

                                current_name = None
                                in_yaml_block = False
                                yaml_lines = []

                                for line in lines:
                                    if line.startswith("QUESTION_NAME:"):
                                        # Save previous question if exists
                                        if current_name and yaml_lines:
                                            questions.append(
                                                {
                                                    "name": current_name,
                                                    "yaml": "\n".join(yaml_lines),
                                                }
                                            )
                                            yaml_lines = []

                                        current_name = line.split(":", 1)[1].strip()
                                        in_yaml_block = False
                                    elif line.strip() == "```yaml":
                                        in_yaml_block = True
                                    elif line.strip() == "```" and in_yaml_block:
                                        in_yaml_block = False
                                    elif in_yaml_block:
                                        yaml_lines.append(line)

                                # Don't forget the last question
                                if current_name and yaml_lines:
                                    questions.append(
                                        {
                                            "name": current_name,
                                            "yaml": "\n".join(yaml_lines),
                                        }
                                    )

                            else:
                                # Unknown block type
                                if debug:
                                    print(
                                        f"\n[UNKNOWN BLOCK TYPE: {type(block).__name__}]",
                                        flush=True,
                                    )
                                    print(f"  {repr(block)}", flush=True)

                # If we got here, the query succeeded - break out of retry loop
                break

            except KeyboardInterrupt:
                # Propagate interrupt to outer handler for proper cleanup
                if debug:
                    print("\n[yellow]⚠ Interrupted by user[/yellow]", flush=True)
                raise
            except Exception as e:
                error_msg = str(e)

                # Check for API overload error (500 errors with "Overloaded" message)
                if "overloaded" in error_msg.lower() or (
                    "api_error" in error_msg.lower() and "500" in error_msg
                ):
                    if attempt < max_retries - 1:
                        wait_time = retry_delay * (
                            2**attempt
                        )  # Exponential backoff: 2s, 4s, 8s
                        self._console.print(
                            f"\n[yellow]⚠ API Overloaded - Retrying in {wait_time}s (attempt {attempt + 1}/{max_retries})[/yellow]"
                        )
                        await asyncio.sleep(wait_time)
                        continue  # Retry
                    else:
                        self._console.print(
                            f"\n[red]✗ API Overloaded - Max retries ({max_retries}) exceeded[/red]"
                        )
                        return {
                            "error": "API overloaded after multiple retries. Please try again later."
                        }

                # Check for content filtering error
                if (
                    "content filtering policy" in error_msg.lower()
                    or "output blocked" in error_msg.lower()
                ):
                    self.logger.error(
                        "Claude API blocked output due to content filtering policy"
                    )

                    # Save the prompt that caused the error for debugging
                    if not debug:
                        prompt_file = self.output_dir / "last_prompt.txt"
                        with open(prompt_file, "w") as f:
                            f.write(prompt)
                        self.logger.debug(f"Saved prompt to {prompt_file}")

                    if not debug:
                        self._console.print(f"\n[red]⚠ Content Filtering Error[/red]")
                        self._console.print(
                            "Claude's API blocked the response due to content filtering policy.\n"
                            "This can happen when survey content contains sensitive topics.\n\n"
                            "[yellow]Suggestions:[/yellow]\n"
                            "1. Try removing or sanitizing sensitive content from the input files\n"
                            "2. If using a DOCX file, check for potentially sensitive text\n"
                            "3. Try processing smaller sections of the survey separately\n"
                            f"4. Check {prompt_file} to see what triggered the filter\n"
                        )
                    return {"error": "Content filtering policy violation"}

                # Log other exceptions but don't catch KeyboardInterrupt
                self.logger.error(f"Error during Claude query: {e}")
                raise

        # Save Claude's full response for debugging (only if not in debug mode)
        if all_text and not debug:
            response_file = self.output_dir / "last_claude_response.txt"
            with open(response_file, "w") as f:
                f.write("\n".join(all_text))
            self.logger.debug(f"Saved Claude response to {response_file}")

        if not questions:
            self.logger.error("No questions generated from Claude response")
            if not debug:
                self.logger.error(
                    f"Check {self.output_dir}/last_prompt.txt and {self.output_dir}/last_claude_response.txt for details"
                )
            if self.verbose_logs and all_text:
                self.logger.debug(f"Full Claude response:\n{''.join(all_text)}")
            return {"error": "No questions generated"}

        self.logger.debug(
            f"Successfully parsed {len(questions)} questions from Claude response"
        )
        return {"questions": questions}

    def _create_question_from_yaml(
        self, yaml_str: str, question_name: str
    ) -> Dict[str, Any]:
        """Create EDSL Question from YAML representation."""
        try:
            # First validate the YAML
            yaml_data = yaml.safe_load(yaml_str)
            if not yaml_data or "question_type" not in yaml_data:
                return {
                    "success": False,
                    "error": "Missing required field 'question_type'",
                }

            # Create question using from_yaml
            question = QuestionBase.from_yaml(yaml_str)

            return {"success": True, "question": question}

        except yaml.YAMLError as e:
            return {"success": False, "error": f"Invalid YAML: {e}"}
        except Exception as e:
            return {"success": False, "error": f"Failed to create question: {str(e)}"}

    async def _suggest_column_mapping(
        self,
        missing_questions: List[str],
        available_columns: List[str],
        survey_questions: Dict[str, Any],
    ) -> Dict[str, str]:
        """Use Claude to suggest mappings from survey questions to data columns.

        Args:
            missing_questions: List of question names that weren't found in data
            available_columns: List of column names available in the data file
            survey_questions: Dict mapping question names to Question objects

        Returns:
            Dictionary mapping question names to suggested column names
        """
        # Build a prompt for Claude
        prompt = f"""I am analyzing existing research survey data and need help mapping question names to data columns.

This is for ACADEMIC RESEARCH ANALYSIS. I have a survey instrument and a data file from that survey,
but the question names don't exactly match the column names. I need to map them correctly for analysis.

Survey Questions (not found in data):
"""
        for qname in missing_questions[:20]:  # Limit to avoid huge prompts
            q = survey_questions.get(qname)
            if q:
                prompt += f'- {qname}: "{q.question_text[:100]}"\n'
            else:
                prompt += f"- {qname}\n"

        prompt += f"\nAvailable columns in data file:\n"
        prompt += ", ".join(available_columns[:50])  # Limit to avoid huge prompts

        prompt += """

Please suggest mappings from question names to column names. 
Only suggest mappings where you're confident there's a match.
Return your answer as a JSON object where keys are question names and values are column names.

Example format:
{
  "age": "respondent_age",
  "gender": "sex",
  "income": "household_income"
}

If you can't find a good match for a question, don't include it.
Return ONLY the JSON object, no other text.
"""

        # Use Claude to get suggestions
        options = ClaudeAgentOptions(
            system_prompt="You are a data mapping assistant helping with ACADEMIC RESEARCH ANALYSIS. "
            "You are analyzing existing survey data to map question names to data column names. "
            "You are NOT conducting surveys or asking questions to real people - you are performing "
            "technical data analysis on existing research materials.",
            max_turns=2,
        )

        try:
            response_text = ""
            async for message in query(prompt=prompt, options=options):
                if isinstance(message, AssistantMessage):
                    for block in message.content:
                        if isinstance(block, TextBlock):
                            response_text += block.text

            # Extract JSON from response
            import json

            # Try to find JSON in the response
            start = response_text.find("{")
            end = response_text.rfind("}") + 1

            if start >= 0 and end > start:
                json_str = response_text[start:end]
                mapping = json.loads(json_str)

                # Validate mappings
                valid_mapping = {}
                for qname, col in mapping.items():
                    if qname in missing_questions and col in available_columns:
                        valid_mapping[qname] = col

                return valid_mapping
            else:
                return {}

        except Exception as e:
            self.logger.warning(f"Failed to parse Claude's column mapping: {e}")
            return {}

    async def _build_value_mappings(
        self,
        df,
        question_to_column: Dict[str, str],
        survey_questions: Dict[str, Any],
        file_ext: str,
        meta=None,
    ) -> Dict[str, Dict]:
        """Build mappings from coded values (e.g., 1, 2, 3) to text labels.

        Args:
            df: DataFrame with response data
            question_to_column: Mapping of question names to data columns
            survey_questions: Dict of Question objects
            file_ext: File extension (to check for Stata/SPSS)
            meta: Metadata from pyreadstat (for Stata/SPSS files)

        Returns:
            Dictionary mapping question names to value mappings (e.g., {'gender': {'1': 'Male', '2': 'Female'}})
        """
        import pandas as pd

        value_mappings = {}

        # First, extract value labels from Stata/SPSS metadata if available
        metadata_labels = {}
        if file_ext in ["dta", "sav"] and meta is not None:
            if hasattr(meta, "value_labels") and meta.value_labels:
                metadata_labels = meta.value_labels
                self._console.print(
                    f"[dim]Found value labels in {file_ext.upper()} metadata for {len(metadata_labels)} variable(s)[/dim]"
                )

        # For each question, check if we need code mapping
        questions_needing_mapping = []

        for qname, col in question_to_column.items():
            question = survey_questions[qname]

            # Check if question has options (multiple choice, checkbox, etc.)
            question_options = None
            if hasattr(question, "question_options") and question.question_options:
                question_options = question.question_options

            if question_options:
                # Get unique values in the data for this column
                unique_values = df[col].dropna().unique()

                # Check if values are numeric codes that don't match options
                needs_mapping = False
                if len(unique_values) > 0:
                    # Check if all values are numeric or numeric strings
                    try:
                        numeric_values = [
                            float(v) for v in unique_values if pd.notna(v)
                        ]
                        # If we have numeric values and they're integers like 1, 2, 3...
                        if numeric_values and all(v == int(v) for v in numeric_values):
                            # Check if these codes match the question options
                            str_values = [str(int(v)) for v in numeric_values]
                            if not any(
                                str_val in question_options for str_val in str_values
                            ):
                                needs_mapping = True
                    except (ValueError, TypeError):
                        pass  # Not numeric, no mapping needed

                if needs_mapping:
                    # Check if we have metadata labels for this column
                    if col in metadata_labels:
                        # Use metadata labels
                        value_mappings[qname] = {
                            str(k): str(v) for k, v in metadata_labels[col].items()
                        }
                        self._console.print(
                            f"[green]✓[/green] Using {file_ext.upper()} metadata labels for {qname}"
                        )
                    else:
                        # Need to ask Claude for mapping
                        questions_needing_mapping.append(
                            {
                                "question_name": qname,
                                "question_text": question.question_text,
                                "question_options": question_options,
                                "data_values": sorted(
                                    [
                                        str(int(float(v)))
                                        for v in unique_values
                                        if pd.notna(v)
                                    ]
                                ),
                                "column": col,
                            }
                        )

        # Ask Claude to map values for questions that need it
        if questions_needing_mapping:
            self._console.print(
                f"\n[cyan]🤖 Asking Claude to map response codes to text values for {len(questions_needing_mapping)} question(s)...[/cyan]"
            )

            for q_info in questions_needing_mapping:
                mapping = await self._suggest_value_mapping(
                    q_info["question_name"],
                    q_info["question_text"],
                    q_info["question_options"],
                    q_info["data_values"],
                )
                if mapping:
                    value_mappings[q_info["question_name"]] = mapping
                    self._console.print(
                        f"[green]✓[/green] Mapped codes for {q_info['question_name']}: {mapping}"
                    )

        return value_mappings

    async def _suggest_value_mapping(
        self,
        question_name: str,
        question_text: str,
        question_options: list,
        data_values: list,
    ) -> Dict[str, str]:
        """Use Claude to suggest mapping from numeric codes to text values.

        Args:
            question_name: Name of the question
            question_text: Text of the question
            question_options: List of text options for the question
            data_values: List of numeric codes found in the data

        Returns:
            Dictionary mapping codes to text values (e.g., {'1': 'Male', '2': 'Female'})
        """
        prompt = f"""I am analyzing existing research survey data and need to decode numeric response codes.

This is for ACADEMIC RESEARCH ANALYSIS of an existing survey that has already been conducted.
I need to map the numeric codes found in the data file to their corresponding text values.

Question: {question_name}
Text: "{question_text}"

The original survey instrument has these answer options:
{chr(10).join(f"- {opt}" for opt in question_options)}

The data file contains these numeric codes as responses:
{', '.join(data_values)}

Based on common survey conventions and the question context, suggest a mapping from numeric codes to answer options.
Typically, codes like 1, 2, 3... map to options in order.

Return ONLY a JSON object mapping codes to options, like:
{{"1": "First option text", "2": "Second option text"}}

Return ONLY the JSON object, no other text.
"""

        options = ClaudeAgentOptions(
            system_prompt="You are a survey data mapping assistant helping with ACADEMIC RESEARCH ANALYSIS. "
            "You are analyzing existing survey response data to map numeric codes to text values. "
            "You are NOT conducting surveys or asking questions to real people - you are performing "
            "technical data analysis on existing research materials to decode response values.",
            max_turns=2,
        )

        try:
            response_text = ""
            async for message in query(prompt=prompt, options=options):
                # Only process AssistantMessage - skip SystemMessage, UserMessage, ResultMessage
                if isinstance(message, AssistantMessage):
                    for block in message.content:
                        if isinstance(block, TextBlock):
                            response_text += block.text

            # Extract JSON from response
            import json

            start = response_text.find("{")
            end = response_text.rfind("}") + 1

            if start >= 0 and end > start:
                json_str = response_text[start:end]
                mapping = json.loads(json_str)

                # Validate that all data values are mapped
                valid_mapping = {}
                for code in data_values:
                    if code in mapping:
                        valid_mapping[code] = mapping[code]

                return valid_mapping
            else:
                return {}

        except Exception as e:
            self.logger.warning(f"Failed to parse Claude's value mapping: {e}")
            return {}

    def generate_results(
        self,
        response_file: Optional[Path] = None,
        sample_size: Optional[int] = None,
        seed: str = "edsl",
        disable_remote_cache: bool = False,
        disable_remote_inference: bool = True,
        verbose: bool = False,
    ):
        """Generate EDSL Results from survey response data (sync wrapper)."""
        return anyio.run(
            self._generate_results_async,
            response_file,
            sample_size,
            seed,
            disable_remote_cache,
            disable_remote_inference,
            verbose,
        )

    async def _select_response_file_with_claude(
        self, file_paths: List[Path]
    ) -> Optional[Path]:
        """Use Claude to analyze files and select the best one for response data.

        Args:
            file_paths: List of available file paths

        Returns:
            Path to the recommended response data file, or None if Claude can't decide
        """
        from claude_agent_sdk import (
            AssistantMessage,
            ClaudeAgentOptions,
            TextBlock,
            query,
        )
        import pandas as pd

        # Only analyze tabular files
        tabular_formats = ["csv", "xlsx", "xls", "dta", "sav"]
        tabular_files = [
            fp for fp in file_paths if fp.suffix.lower().lstrip(".") in tabular_formats
        ]

        if not tabular_files:
            return None

        if len(tabular_files) == 1:
            # Only one file, use it
            return tabular_files[0]

        # Build file analysis for Claude
        file_info = []
        for fp in tabular_files:
            try:
                # Get sample data
                file_ext = fp.suffix.lower().lstrip(".")
                if file_ext == "csv":
                    df = pd.read_csv(fp, nrows=5)
                elif file_ext in ["xlsx", "xls"]:
                    df = pd.read_excel(fp, nrows=5)
                elif file_ext == "dta":
                    import pyreadstat

                    df, meta = pyreadstat.read_dta(str(fp), row_limit=5)
                elif file_ext == "sav":
                    import pyreadstat

                    df, meta = pyreadstat.read_sav(str(fp), row_limit=5)
                else:
                    continue

                info = {
                    "filename": fp.name,
                    "columns": list(df.columns),
                    "num_columns": len(df.columns),
                    "sample": df.head(3).to_string(max_cols=8, max_colwidth=30),
                }
                file_info.append(info)
            except Exception as e:
                self.logger.warning(f"Failed to analyze {fp.name}: {e}")
                continue

        if not file_info:
            return None

        # Ask Claude to analyze
        prompt = f"""I have {len(file_info)} data file(s) and need to identify which contains SURVEY RESPONSE DATA (actual responses from survey participants).

This is for ACADEMIC RESEARCH ANALYSIS. I need to generate results from survey responses, not from metadata/codebooks.

Files available:

"""
        for i, info in enumerate(file_info, 1):
            prompt += f"\nFile {i}: {info['filename']}\n"
            prompt += f"  Columns ({info['num_columns']}): {', '.join(info['columns'][:15])}\n"
            if len(info["columns"]) > 15:
                prompt += f"    ... and {len(info['columns']) - 15} more\n"
            prompt += f"  Sample data:\n{info['sample']}\n"

        prompt += """

IMPORTANT: Identify which file contains ACTUAL SURVEY RESPONSES where:
- Each ROW represents a survey respondent/participant
- Each COLUMN represents a survey question/variable (e.g., age, gender, income, attitudes, etc.)
- Values are actual responses from participants

DO NOT select files that are:
- Codebooks (metadata ABOUT variables, with columns like "Variable", "Label", "Values")
- Data dictionaries
- Survey instruments/questionnaires

Return ONLY the filename of the response data file, nothing else.
If multiple files have response data, pick the one that looks most complete.
If no file looks like response data, return "NONE".
"""

        options = ClaudeAgentOptions(
            system_prompt="You are a research data analyst helping identify survey response data files. "
            "You excel at distinguishing between actual survey responses and metadata files.",
            max_turns=2,
        )

        try:
            response_text = ""
            async for message in query(prompt=prompt, options=options):
                if isinstance(message, AssistantMessage):
                    for block in message.content:
                        if isinstance(block, TextBlock):
                            response_text += block.text

            # Extract filename from response
            response_text = response_text.strip()

            if "NONE" in response_text.upper():
                return None

            # Find the mentioned filename in our files
            for fp in tabular_files:
                if fp.name in response_text:
                    self._console.print(
                        f"[green]✓[/green] Claude identified response data: [cyan]{fp.name}[/cyan]"
                    )
                    return fp

            return None

        except Exception as e:
            self.logger.warning(f"Claude file selection failed: {e}")
            return None

    async def _generate_results_async(
        self,
        response_file: Optional[Path] = None,
        sample_size: Optional[int] = None,
        seed: str = "edsl",
        disable_remote_cache: bool = False,
        disable_remote_inference: bool = True,
        verbose: bool = False,
    ):
        """Generate EDSL Results from survey response data.

        This method creates agents from survey response data (CSV, Excel, Stata, or SPSS files)
        and runs them through the generated survey to produce EDSL Results.

        Args:
            response_file: Path to the file containing survey responses (CSV, .xlsx, .dta, or .sav).
                          If None, automatically selects the first tabular data file from constructor.
            sample_size: Optional number of respondents to sample. If None, uses all respondents.
                        If the file has >1000 rows, you'll be prompted unless sample_size is set.
            seed: Random seed for sampling (default: "edsl")
            disable_remote_cache: Whether to disable remote caching (default: False)
            disable_remote_inference: Whether to disable remote inference (default: True)
            verbose: If True, shows detailed progress information (default: False)

        Returns:
            EDSL Results object containing the survey responses as if answered by LLM agents

        Raises:
            ValueError: If no survey has been generated yet, or if response file format is not supported
            FileNotFoundError: If the response file cannot be read

        Example:
            >>> from edsl.assistants import SurveyAssistant
            >>> # Generate survey from documents, get results from data file
            >>> sa = SurveyAssistant('survey.docx', 'responses.xlsx')
            >>> sa.generate_survey()
            >>> results = sa.generate_results(sample_size=100)
            >>> results.select('answer.*').print()
        """
        from edsl.agents import Agent, AgentList
        from edsl.results import Results

        # Check if survey exists
        if self._survey is None:
            self._console.print("[red]Error:[/red] No survey has been generated yet.")
            self._console.print("Please run generate_survey() first.")
            raise ValueError("No survey generated. Call generate_survey() first.")

        # Determine which file to use
        if response_file is None:
            if not self.file_paths:
                raise ValueError(
                    "No response file specified and no files in constructor."
                )

            # Use Claude to analyze files and recommend which to use
            self._console.print(
                f"\n[cyan]🤖 Analyzing {len(self.file_paths)} file(s) to find response data...[/cyan]"
            )
            response_file = await self._select_response_file_with_claude(
                self.file_paths
            )

            if response_file is None:
                # Fallback to heuristic selection
                self._console.print(
                    "[yellow]Claude couldn't determine the best file, using heuristic selection...[/yellow]"
                )
                tabular_formats = ["csv", "xlsx", "xls", "dta", "sav"]
                tabular_files = [
                    fp
                    for fp in self.file_paths
                    if fp.suffix.lower().lstrip(".") in tabular_formats
                ]
                if tabular_files:
                    response_file = tabular_files[0]
                else:
                    response_file = self.file_paths[0]
                self._console.print(f"[dim]Selected: {response_file.name}[/dim]")
        else:
            response_file = Path(response_file)

        if not response_file.exists():
            raise FileNotFoundError(f"Response file not found: {response_file}")

        # Check file type
        file_ext = response_file.suffix.lower().lstrip(".")
        supported_formats = ["csv", "xlsx", "xls", "dta", "sav"]

        if file_ext not in supported_formats:
            raise ValueError(
                f"Unsupported file format: {file_ext}. "
                f"Supported formats: {', '.join(supported_formats)}"
            )

        self._console.print(
            f"\n[bold cyan]🔄 Generating Results from survey responses[/bold cyan]"
        )
        self._console.print(f"Response file: [cyan]{response_file.name}[/cyan]")
        self._console.print(
            f"Survey: [cyan]{len(self._survey.questions)} questions[/cyan]\n"
        )

        # For structured data files (DTA, SAV), try using Conjure first
        if file_ext in ["dta", "sav"]:
            try:
                from edsl.conjure import Conjure

                self._console.print(
                    f"[dim]Using EDSL Conjure for native {file_ext.upper()} parsing...[/dim]"
                )

                conjure_data = Conjure(str(response_file))

                # Get row count to inform user
                try:
                    import pandas as pd

                    if file_ext == "dta":
                        import pyreadstat

                        _, meta = pyreadstat.read_dta(
                            str(response_file), metadataonly=True
                        )
                        total_rows = meta.number_rows
                    elif file_ext == "sav":
                        import pyreadstat

                        _, meta = pyreadstat.read_sav(
                            str(response_file), metadataonly=True
                        )
                        total_rows = meta.number_rows

                    # Warn about large datasets if no sample_size specified
                    if sample_size is None and total_rows > 1000:
                        self._console.print(
                            f"\n[yellow]⚠️  Large dataset detected: {total_rows} respondents[/yellow]"
                        )
                        self._console.print(
                            f"[dim]Processing all rows. This may take a while...[/dim]"
                        )
                        self._console.print(
                            f"[dim]Tip: Use sample_size parameter to process a subset, e.g.:[/dim]"
                        )
                        self._console.print(
                            f"[dim]  sa.generate_results(sample_size=100)[/dim]\n"
                        )
                    elif sample_size:
                        self._console.print(
                            f"[dim]Sampling {sample_size} of {total_rows} total respondents[/dim]\n"
                        )
                except:
                    pass  # If we can't get row count, just continue

                # Use Conjure's built-in to_results method with verbose=True for progress
                self._console.print(
                    f"[yellow]⏳[/yellow] Generating results from {conjure_data.question_names.__len__()} questions...\n"
                )
                results = conjure_data.to_results(
                    sample_size=sample_size,
                    seed=seed,
                    disable_remote_cache=disable_remote_cache,
                    disable_remote_inference=disable_remote_inference,
                    verbose=True,  # Always show progress for long-running operations
                )

                self._console.print(f"\n[green]✅ Results generation complete![/green]")
                self._console.print(
                    f"[green]✓[/green] Generated results for {len(results)} respondent(s)"
                )

                # Cache results for later access
                self._results = results
                return results

            except ImportError:
                self._console.print(
                    "[yellow]⚠️  Conjure not available, using manual agent construction[/yellow]"
                )
                # Fall through to manual agent construction
            except Exception as e:
                self._console.print(f"[yellow]⚠️  Conjure failed: {e}[/yellow]")
                self._console.print(
                    "[dim]Falling back to manual agent construction...[/dim]\n"
                )
                # Fall through to manual agent construction

        # Import pandas once at the start (needed for all file types)
        import pandas as pd

        # Read the data using appropriate method
        meta = None  # Will store metadata from Stata/SPSS files
        try:
            if file_ext == "csv":
                df = pd.read_csv(response_file)
                total_rows = len(df)
                column_names = df.columns.tolist()
            elif file_ext in ["xlsx", "xls"]:
                df = pd.read_excel(response_file)
                total_rows = len(df)
                column_names = df.columns.tolist()
            elif file_ext == "dta":
                import pyreadstat

                df, meta = pyreadstat.read_dta(str(response_file))
                total_rows = len(df)
                column_names = df.columns.tolist()
            elif file_ext == "sav":
                import pyreadstat

                df, meta = pyreadstat.read_sav(str(response_file))
                total_rows = len(df)
                column_names = df.columns.tolist()
        except ImportError as e:
            if "pyreadstat" in str(e):
                raise ImportError(
                    "pyreadstat is required for Stata and SPSS files. "
                    "Install with: pip install pyreadstat"
                )
            elif "pandas" in str(e):
                raise ImportError(
                    "pandas is required for CSV and Excel files. "
                    "Install with: pip install pandas"
                )
            elif "openpyxl" in str(e) or "xlrd" in str(e):
                raise ImportError(
                    "openpyxl is required for Excel files. "
                    "Install with: pip install openpyxl"
                )
            raise
        except Exception as e:
            raise ValueError(f"Failed to read response file: {e}")

        # Check file size and warn/prompt if needed
        if total_rows > 1000 and sample_size is None:
            self._console.print(
                f"[yellow]⚠ Warning:[/yellow] Response file contains {total_rows} rows."
            )
            self._console.print(
                "This may take a long time to process. Consider using sample_size parameter."
            )
            self._console.print("Example: generate_results(..., sample_size=100)\n")
            # Don't auto-set sample_size, let the user decide

        # Check if this looks like a codebook file instead of response data
        codebook_indicators = [
            "variable",
            "variable_label",
            "value_labels",
            "values",
            "variablelabel",
            "valuelabels",
        ]
        is_likely_codebook = any(
            col.lower().replace("_", "").replace(" ", "") in codebook_indicators
            for col in column_names
        )

        # Strong codebook detection - if we have typical codebook columns and very few columns, reject it
        if is_likely_codebook and len(column_names) <= 10:
            has_variable_col = any(
                "variable" in col.lower() and col.lower() != "variables"
                for col in column_names
            )
            has_label_col = any("label" in col.lower() for col in column_names)

            if has_variable_col and has_label_col:
                # This is definitely a codebook - reject it
                error_msg = (
                    f"Cannot use codebook/data dictionary file for generating results.\n\n"
                    f"File: {response_file.name}\n"
                    f"Columns found: {', '.join(column_names)}\n\n"
                    f"This appears to be a CODEBOOK (metadata about variables), not actual survey response data.\n"
                    f"Response data should have columns with variable names (like 'age', 'gender', 'income', etc.) "
                    f"and rows representing individual survey respondents.\n\n"
                    f"Please provide the actual survey response data file."
                )
                self._console.print(f"\n[red]✗ Error:[/red] {error_msg}")
                raise ValueError(error_msg)

            # Warn but continue
            self._console.print(
                f"\n[yellow]⚠️  Warning:[/yellow] This file might be a CODEBOOK or DATA DICTIONARY."
            )
            self._console.print(f"   Columns: {', '.join(column_names)}")
            self._console.print(
                f"   If results look wrong, check if you're using the correct file.\n"
            )

        # Create mapping from survey question names to data columns
        survey_questions = {q.question_name: q for q in self._survey.questions}

        self._console.print(f"[dim]Mapping survey questions to data columns...[/dim]")
        question_to_column = {}
        missing_columns = []

        # Build reverse mapping: sanitized column name -> original column name
        sanitized_to_original = {}
        for col in column_names:
            sanitized = sanitize_column_name(col)
            sanitized_to_original[sanitized] = col

        for qname in survey_questions.keys():
            if qname in column_names:
                # Exact match
                question_to_column[qname] = qname
            elif qname in sanitized_to_original:
                # Question name matches sanitized version of a column
                question_to_column[qname] = sanitized_to_original[qname]
            else:
                # Try case-insensitive matching
                found = False
                for col in column_names:
                    if col.lower() == qname.lower():
                        question_to_column[qname] = col
                        found = True
                        break
                if not found:
                    missing_columns.append(qname)

        if missing_columns:
            self._console.print(
                f"\n[yellow]⚠ Column Mismatch:[/yellow] {len(missing_columns)} question(s) not found in data:"
            )
            for qname in missing_columns[:10]:
                self._console.print(f"  - {qname}")
            if len(missing_columns) > 10:
                self._console.print(f"  ... and {len(missing_columns) - 10} more")

            # Try to use Claude to map columns
            self._console.print(
                f"\n[cyan]🤖 Asking Claude to suggest column mappings...[/cyan]"
            )

            try:
                suggested_mapping = await self.column_mapper.suggest_column_mapping(
                    missing_columns, column_names, survey_questions, df
                )

                if suggested_mapping:
                    self._console.print(
                        f"[green]✓[/green] Claude suggested {len(suggested_mapping)} mapping(s):"
                    )
                    for qname, col in suggested_mapping.items():
                        self._console.print(f"  {qname} → {col}")
                        question_to_column[qname] = col
                        missing_columns.remove(qname)
                    self._console.print()
                else:
                    self._console.print(
                        f"[yellow]Claude could not suggest mappings[/yellow]\n"
                    )
            except Exception as e:
                self._console.print(
                    f"[yellow]Warning:[/yellow] Claude mapping failed: {e}\n"
                )

            # If still have missing columns, raise error
            if missing_columns:
                error_msg = (
                    f"{len(missing_columns)} survey question(s) have no matching columns in data file:\n"
                    f"Missing: {', '.join(missing_columns[:10])}\n"
                    f"Available columns: {', '.join(column_names[:20])}"
                )
                self._console.print(f"\n[red]Error:[/red] {error_msg}")

                # Give extra hint if this looks like a codebook
                if is_likely_codebook:
                    self._console.print(
                        f"\n[yellow]💡 Hint:[/yellow] The file appears to be a codebook/data dictionary."
                    )
                    self._console.print(
                        "   Try using the actual survey response data file instead."
                    )
                    self._console.print(
                        "   Response data should have columns matching your survey questions."
                    )

                raise ValueError(error_msg)

        # Build value mappings (e.g., 1 -> "Male", 2 -> "Female")
        self._console.print(f"\n[dim]Checking for response code mappings...[/dim]")
        value_mappings = await self.column_mapper.build_value_mappings(
            df, question_to_column, survey_questions, file_ext, meta
        )

        if value_mappings:
            self._console.print(
                f"[green]✓[/green] Built value mappings for {len(value_mappings)} question(s)\n"
            )
        else:
            self._console.print(f"[dim]No code mappings needed[/dim]\n")

        # Determine sample size
        if sample_size is not None:
            if sample_size > total_rows:
                self._console.print(
                    f"[yellow]Warning:[/yellow] Requested sample_size ({sample_size}) > total rows ({total_rows}). Using all rows."
                )
                sample_size = total_rows
            else:
                self._console.print(
                    f"Using sample of {sample_size} out of {total_rows} respondents"
                )
                import random

                random.seed(seed)
                df = df.sample(n=sample_size, random_state=seed)
        else:
            sample_size = total_rows
            self._console.print(f"Processing all {total_rows} respondents")

        # Create agents from the data
        self._console.print(
            f"\n[dim]Creating {sample_size} agents from survey data...[/dim]"
        )
        agents = []

        for idx, row in df.iterrows():
            # Build agent traits with _agent suffix and direct answering
            traits = {}
            for qname, col in question_to_column.items():
                value = row[col]
                # Convert to string, handle NaN/None
                if pd.isna(value):
                    value = None
                else:
                    value = str(value)

                    # Apply value mapping if available (e.g., 1 -> "Male")
                    if qname in value_mappings:
                        # Check if this is a numeric code that needs mapping
                        try:
                            # Convert to int string for lookup (e.g., 1.0 -> "1")
                            code = str(int(float(value)))
                            if code in value_mappings[qname]:
                                value = value_mappings[qname][code]
                        except (ValueError, TypeError):
                            pass  # Not a numeric code, use as-is

                traits[f"{qname}_agent"] = value

            # Create agent with direct question answering
            agent = Agent(traits=traits)

            # Add direct question answering method
            def make_answerer(agent_traits):
                def answer_func(self, question, scenario=None):
                    return agent_traits.get(question.question_name + "_agent", None)

                return answer_func

            agent.add_direct_question_answering_method(make_answerer(traits))
            agents.append(agent)

        agent_list = AgentList(agents)
        self._console.print(f"[green]✓[/green] Created {len(agent_list)} agents")

        # Run the survey
        self._console.print(
            f"\n[dim]Running survey with {len(agent_list)} agents...[/dim]"
        )

        try:
            results = self._survey.by(agent_list).run(
                disable_remote_cache=disable_remote_cache,
                disable_remote_inference=disable_remote_inference,
                verbose=verbose,
            )

            self._console.print(
                f"\n[bold green]✅ Results generated successfully![/bold green]"
            )
            self._console.print(f"Total responses: {len(results)}")

            # Cache results for later access
            self._results = results
            return results

        except Exception as e:
            self._console.print(f"\n[red]Error running survey:[/red] {e}")
            raise

    def to_survey(self) -> Optional[Survey]:
        """Return the generated Survey object.

        Returns:
            The Survey object if one has been generated, None otherwise.
            Note: Use the .survey property instead for automatic generation.
        """
        return self._survey


@app.command()
def generate(
    files: List[Path] = typer.Argument(
        ..., help="List of files relevant to survey generation"
    ),
    spec_file: Optional[Path] = typer.Option(
        None,
        "--spec",
        "-s",
        help="Path to custom EDSL specification markdown file (uses bundled spec by default)",
    ),
    output_dir: Optional[Path] = typer.Option(
        Path("edsl_output"),
        "--output-dir",
        "-o",
        help="Output directory for generated files",
    ),
):
    """Generate EDSL survey from document files."""
    console.print(f"[bold]EDSL Survey Generator[/bold]")

    # Check for API key early
    if not check_api_key():
        raise typer.Exit(1)

    console.print(f"Processing {len(files)} files...")

    # Validate files exist
    valid_files = []
    for file in files:
        if file.exists():
            valid_files.append(file)
            console.print(f"  ✓ {file}")
        else:
            console.print(f"  [red]✗ {file} (not found)[/red]")

    if not valid_files:
        console.print("[red]No valid files found![/red]")
        raise typer.Exit(1)

    # Run async generation
    async def run_generation():
        generator = SurveyAssistant(spec_file=spec_file)
        generator.output_dir = output_dir
        generator.output_dir.mkdir(exist_ok=True)

        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("Generating survey...", total=None)

            result = await generator._generate_survey_async(valid_files)

            progress.stop()

            # Display results table
            console.print(f"\n[bold]Generation Results[/bold]")

            table = Table(show_header=True, header_style="bold magenta")
            table.add_column("Question Name", style="cyan", no_wrap=True)
            table.add_column("Type", style="blue")
            table.add_column("Status", style="green")
            table.add_column("YAML File", style="dim")

            for q_result in result.get("questions_processed", []):
                status_style = "green" if "Success" in q_result["status"] else "red"
                table.add_row(
                    q_result["name"],
                    q_result["type"],
                    q_result["status"],
                    q_result["file"],
                )

            console.print(table)

            # Summary
            summary = result.get("survey_summary", {})
            console.print(f"\n[bold]Summary:[/bold]")
            console.print(
                f"  Total questions attempted: {summary.get('total_questions', 0)}"
            )
            console.print(
                f"  Successfully created: [green]{summary.get('successful_questions', 0)}[/green]"
            )
            console.print(f"  Failed: [red]{summary.get('failed_questions', 0)}[/red]")

            if result.get("success"):
                console.print(f"\n[green]✓ Survey generated successfully![/green]")
                console.print(f"Output directory: [bold]{result['output_dir']}[/bold]")
                console.print(f"  - Individual YAMLs: {result['output_dir']}/*.yaml")
                console.print(
                    f"  - Survey YAML: {result.get('survey_yaml_file', 'N/A')}"
                )
                console.print(f"  - Survey object: {result.get('output_file', 'N/A')}")
            else:
                console.print(f"\n[red]✗ Survey generation failed![/red]")
                if result.get("errors"):
                    console.print(f"\n[bold]Errors:[/bold]")
                    for error in result["errors"]:
                        console.print(f"  - {error['question_name']}: {error['error']}")
                        console.print(f"    Check: {error['yaml_file']}")

    anyio.run(run_generation)


@app.command()
def validate(
    yaml_file: Path = typer.Argument(..., help="YAML file to validate"),
):
    """Validate a single YAML question file."""
    if not yaml_file.exists():
        console.print(f"[red]File not found: {yaml_file}[/red]")
        raise typer.Exit(1)

    with open(yaml_file, "r") as f:
        yaml_content = f.read()

    console.print(f"[bold]Validating: {yaml_file}[/bold]\n")
    console.print("[dim]YAML Content:[/dim]")
    console.print(yaml_content)
    console.print("\n[dim]Validation Result:[/dim]")

    generator = SurveyAssistant()
    result = generator._create_question_from_yaml(yaml_content, yaml_file.stem)

    if result["success"]:
        question = result["question"]
        console.print(f"[green]✓ Valid EDSL question![/green]")
        console.print(f"  Type: {question.__class__.__name__}")
        console.print(f"  Name: {question.question_name}")
        console.print(f"  Text: {question.question_text}")
    else:
        console.print(f"[red]✗ Invalid: {result['error']}[/red]")


@app.command()
def example():
    """Show example usage and create sample files."""
    console.print("[bold]EDSL Survey Generator - Example Usage[/bold]\n")

    # Create example survey specification
    example_content = """Customer Satisfaction Survey

We need to measure satisfaction with our new product launch.

Questions needed:
1. Overall satisfaction rating (5-point scale)
2. Which features do you use? (multiple choice, select all that apply: Analytics, Reports, API)
3. Would you recommend to others? (0-10 scale)
4. Any additional feedback? (open text)
5. Would you renew? (Yes/No)
"""

    example_file = Path("example_survey_spec.txt")
    with open(example_file, "w") as f:
        f.write(example_content)

    # Create example CSV with responses
    csv_content = """question_name,response,respondent_id
satisfaction,4,001
features_used,Analytics,001
recommend,8,001
satisfaction,5,002
features_used,Reports,002
recommend,10,002
"""

    csv_file = Path("example_responses.csv")
    with open(csv_file, "w") as f:
        f.write(csv_content)

    console.print(f"Created example files:")
    console.print(f"  - {example_file}")
    console.print(f"  - {csv_file}")
    console.print(f"\nExample command:")
    console.print(
        f"  [bold]uv run edsl_survey_generator.py {example_file} {csv_file}[/bold]"
    )


if __name__ == "__main__":
    app()
