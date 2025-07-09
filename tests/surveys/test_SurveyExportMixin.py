import pytest
from docx import Document
from docx.document import Document as DocumentClass
from edsl.surveys import Survey


class TestDocxMethod:
    @pytest.fixture
    def example_survey(self):
        return Survey.example()

    # def test_docx_method_without_filename(self, example_survey):
    #     result = example_survey.docx()

    #     print(f"Type of result: {type(result)}")
    #     print(f"Result: {result}")

    #     assert result is not None, "docx method returned None"
    #     # Since docx() now returns a FileStore, we need to load the document
    #     from edsl import FileStore
    #     assert isinstance(
    #         result, FileStore
    #     ), f"Expected FileStore object, got {type(result)}"

    #     # Load the document from the FileStore to check content
    #     doc = Document(result.path)
        
    #     # Check if the result has expected Document methods/attributes
    #     assert hasattr(
    #         doc, "add_paragraph"
    #     ), "Document doesn't have 'add_paragraph' method"
    #     assert hasattr(
    #         doc, "add_heading"
    #     ), "Document doesn't have 'add_heading' method"

    #     # Check document content
    #     assert len(doc.paragraphs) > 0, "Document has no paragraphs"
    #     assert (
    #         "EDSL Survey" in doc.paragraphs[0].text
    #     ), "Document doesn't start with expected title"

    #     # Check if all questions are present
    #     question_count = len(example_survey._questions)
    #     for i in range(question_count):
    #         question = example_survey._questions[i]
    #         question_text = f"Question {i + 1} ({question.question_name})"
    #         assert any(
    #             question_text in p.text for p in doc.paragraphs
    #         ), f"Question {i+1} not found in document"

    # def test_docx_method_with_filename(self, example_survey, tmp_path):
    #     filename = tmp_path / "test_survey.docx"
    #     result = example_survey.docx(filename=str(filename))

    #     # Now docx() returns a FileStore object even when filename is provided
    #     from edsl import FileStore
    #     assert isinstance(result, FileStore), f"Expected FileStore object, got {type(result)}"
    #     assert filename.exists()

    #     # Verify the content of the created document
    #     doc = Document(filename)
    #     assert len(doc.paragraphs) > 0
    #     assert "EDSL Survey" in doc.paragraphs[0].text

    #     # Check if all questions are present
    #     question_count = len(example_survey._questions)
    #     for i in range(question_count):
    #         question = example_survey._questions[i]
    #         question_text = f"Question {i + 1} ({question.question_name})"
    #         assert any(question_text in p.text for p in doc.paragraphs)

    # def test_docx_method_content(self, example_survey):
    #     result = example_survey.docx()
        
    #     # Load the document from the FileStore to check content
    #     doc = Document(result.path)

    #     # Check the structure and content of the document
    #     paragraphs = doc.paragraphs
    #     assert "EDSL Survey" in paragraphs[0].text

    #     current_paragraph = (
    #         2  # Start from the first question (after title and blank line)
    #     )
    #     for i, question in enumerate(example_survey._questions):
    #         # Check question header
    #         assert (
    #             f"Question {i + 1} ({question.question_name})"
    #             in paragraphs[current_paragraph].text
    #         )
    #         assert question.question_type in paragraphs[current_paragraph].text
    #         current_paragraph += 1

    #         # Check question text
    #         assert question.question_text in paragraphs[current_paragraph].text
    #         current_paragraph += 1

    #         # Check options
    #         if question.question_type == "linear_scale":
    #             for key, value in question.option_labels.items():
    #                 assert f"{key}: {value}" in paragraphs[current_paragraph].text
    #                 current_paragraph += 1
    #         elif hasattr(question, "question_options"):
    #             for option in question.question_options:
    #                 assert option in paragraphs[current_paragraph].text
    #                 current_paragraph += 1

    #     # Ensure we've checked all paragraphs
    #     assert current_paragraph == len(paragraphs)
